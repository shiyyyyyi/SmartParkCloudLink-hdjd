# -*- coding: utf-8 -*-
"""Generate the SRS Word document from the local requirement template copy.

This script reads the copied template base under the generated-documents folder
and writes only the final SRS deliverable. It does not touch files under
交付物/巩固式解密模板.
"""
from pathlib import Path
from copy import deepcopy
import os
import re
import shutil

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"D:\实训\SmartParkCloudLink-hdjd")
DOC_DIR = ROOT / "文档（项目相关的作业所需提交文档生成在这里）"
SRS_DIR = DOC_DIR / "03_需求分析"
TEMPLATE = SRS_DIR / "_template_base.docx"
OUTPUT = SRS_DIR / "19组-程晓洋-需求分析.docx"
MARKDOWN_OUTPUT = SRS_DIR / "功能需求规格说明书_SRS_V1.0.md"
DIAGRAM_DIR = SRS_DIR / "_generated_use_case_diagrams"

PROJECT_NAME = "AI-driven城市智慧停车管理与诱导系统"
GROUP = "第19组"
DATE = "2026.06.02"
HEADING_SIZE = {1: 22, 2: 16, 3: 16}

INTEGRATION_NOTES = {
    "M1": "检测结果同步给城市停车一张图、违停取证、车场运营后台和监管平台；低置信度识别结果必须保留人工复核记录。",
    "M2": "联网接入数据是城市停车一张图、预约导航、动态定价、监管统计和大数据分析的基础数据源。",
    "M3": "地图展示依赖路内检测和停车场接入数据，并向预约、导航、诱导发布和监管看板提供统一停车资源视图。",
    "M4": "诱导发布读取实时空位、价格、拥堵和设备状态，发布结果回写日志，便于运营后台和监管平台追溯。",
    "M5": "预约与导航联动车位锁定、订单生成、无感支付、反向寻车和车主移动端消息提醒。",
    "M6": "车牌识别结果与预约、订单、支付状态和车辆信息关联，是无感入出场和费用结算的关键闭环。",
    "M7": "分析结果沉淀为运营报表、动态定价依据、监管决策依据和商圈营销效果评估数据。",
    "M8": "定价策略读取供需、时段、区域和月卡/营销约束，输出给预约、支付、诱导屏和车主端展示。",
    "M9": "共享车位发布后进入统一车位池，参与地图展示、预约锁定、订单结算和收益统计。",
    "M10": "充电车位状态同时影响车位预约、车辆类型筛选、设备运维和订单计费。",
    "M11": "违停证据来源于视觉检测和人工复核，处理结果同步城市监管平台并留存审计日志。",
    "M12": "月卡权益影响预约权限、入场识别、订单折扣和运营后台收入统计。",
    "M13": "运营后台统一承接车场、车位、订单、设备、价格、营销和统计数据，是管理员的主要工作台。",
    "M14": "移动端承载车主侧搜索、预约、支付、寻车、消息和个人记录查询等核心体验。",
    "M15": "反向寻车依赖入场识别、车位定位、停车场区域数据和移动端导航展示。",
    "M16": "设备运维状态影响检测、道闸、诱导屏和充电桩业务可用性，并向运营后台输出告警。",
    "M17": "监管平台汇总路内外停车数据、违停取证、收费统计和供需分析，形成城市级监管视角。",
    "M18": "营销活动与停车订单、商户核销、用户画像和商圈客流分析关联，输出优惠与效果统计。",
}

EXCEPTION_NOTES = {
    "M1": "摄像头离线、车牌遮挡、夜间识别失败或车位线被遮挡时，系统标记为待复核，不直接写入最终处罚或结算结果。",
    "M2": "停车场接口超时、字段缺失、空位数异常跳变或密钥失效时，保留最近一次有效数据并生成接入告警。",
    "M3": "地图服务不可用时，以列表和区域分组方式降级展示；空位数据超过有效期时需显示数据延迟提示。",
    "M4": "诱导屏离线、发布失败或内容审核未通过时，不覆盖上一条有效发布内容，并记录失败原因。",
    "M5": "预约锁定超时、车位被人工占用、用户取消或车辆未按时到场时，系统释放锁定并同步订单状态。",
    "M6": "车牌未识别、无支付账户、支付失败或人工放行时，系统进入异常处理流程并保留人工操作日志。",
    "M7": "统计样本不足、异常订单过多或数据同步延迟时，分析结果需标注数据口径和更新时间。",
    "M8": "价格策略冲突、活动优惠叠加异常或监管限价触发时，以监管规则和基础价格为优先。",
    "M9": "共享时段冲突、车位主下架、预约用户爽约或收益结算失败时，系统需回滚共享状态。",
    "M10": "充电桩故障、车辆类型不匹配、充电时长异常或充电车位被普通车辆占用时，系统提示并进入运维处理。",
    "M11": "证据图片缺失、车牌不清晰、位置不明确或复核驳回时，违停记录不得进入处罚状态。",
    "M12": "月卡过期、车牌不一致、适用车场不匹配或余额异常时，按临停规则处理并提示用户。",
    "M13": "管理员误操作、订单冲突、价格发布失败或统计导出异常时，系统需要二次确认和操作日志。",
    "M14": "网络中断、定位失败、支付跳转失败或预约状态变化时，移动端需给出明确状态提示。",
    "M15": "定位点缺失、楼层信息不完整或路线不可达时，系统提供文字位置提示和人工问询入口。",
    "M16": "设备心跳丢失、固件版本不一致、告警重复上报或维护任务超期时，系统合并告警并保留处理记录。",
    "M17": "监管数据延迟、跨区域统计口径不一致或敏感数据展示风险时，系统采用汇总展示和数据脱敏策略。",
    "M18": "优惠券过期、核销失败、订单不满足优惠条件或商户活动下架时，系统需提示原因并保留核销记录。",
}


MODULES = [
    {
        "id": "M1",
        "name": "路内停车位视觉检测",
        "source": "课题要求(1)",
        "priority": "P1",
        "owner": "丁梓钊",
        "actor": "路内运维人员/交管部门",
        "use_cases": ["配置检测点位", "查看车位占用", "识别车牌信息", "识别违停状态", "人工复核结果"],
        "description": "在路内停车位部署AI视频桩或高位视频相机，实时识别车位空闲、占用、违停及车牌信息。实训阶段以模拟检测数据和人工复核流程呈现。",
        "inputs": "摄像头编号、车位编号、检测时间、模拟图片/视频帧、车牌号、车位状态、复核结果。",
        "processing": "系统接收模拟视觉检测结果，按车位编号更新状态；对低置信度或违停结果进入人工复核流程；复核通过后同步至城市停车一张图和监管平台。",
        "outputs": "车位状态、车牌识别结果、违停标记、检测时间、复核状态和检测准确率统计。",
        "rules": "单摄像头覆盖8-12个车位；状态包括空闲、占用、违停、异常；模拟识别结果需保留人工修正入口。"
    },
    {
        "id": "M2",
        "name": "停车场数据联网接入",
        "source": "课题要求(2)",
        "priority": "P0",
        "owner": "程晓洋",
        "actor": "系统管理员/停车场管理员",
        "use_cases": ["配置接入停车场", "同步空余车位", "校验数据格式", "查看接入状态", "处理离线告警"],
        "description": "对接商业停车场、住宅小区、机关单位等路外停车管理系统，通过标准API汇聚实时空余车位、价格和运营状态。",
        "inputs": "停车场基础信息、API地址、接口密钥、空余车位数、收费标准、最后更新时间。",
        "processing": "系统按标准数据格式接收或模拟接收停车场数据，完成字段校验、状态转换、更新时间记录和异常告警。",
        "outputs": "接入状态面板、标准化停车场数据、空位统计、离线告警和同步日志。",
        "rules": "接入数据必须包含停车场ID、总车位数、空余车位数和更新时间；超过5分钟未更新标记为数据延迟。"
    },
    {
        "id": "M3",
        "name": "城市停车一张图",
        "source": "课题要求(3)",
        "priority": "P0",
        "owner": "程子浩",
        "actor": "私家车主/交管部门",
        "use_cases": ["查看车场分布", "按区域筛选", "按价格筛选", "按距离排序", "查看空位热度"],
        "description": "在GIS地图或模拟地图上展示全市停车场实时车位信息，支持按区域、价格、距离、空余车位等维度分类展示。",
        "inputs": "用户目的地、当前位置、区域筛选条件、价格范围、停车场坐标、空余车位数。",
        "processing": "系统汇总联网接入与视觉检测数据，按区域和距离生成停车资源视图；实训阶段可用列表、坐标点和区域热度模拟GIS展示。",
        "outputs": "停车场地图/列表、空余车位、价格、距离、拥挤程度、推荐排序结果。",
        "rules": "空余车位为0时不可预约；空余车位少于5个显示车位紧张；默认按距离和空余数综合排序。"
    },
    {
        "id": "M4",
        "name": "停车诱导信息发布",
        "source": "课题要求(4)",
        "priority": "P1",
        "owner": "程子浩",
        "actor": "交管部门/停车场管理员",
        "use_cases": ["维护诱导屏", "生成诱导内容", "模拟发布信息", "查看发布状态", "调整诱导策略"],
        "description": "通过区域诱导屏、路段诱导屏、停车场入口诱导屏分级发布车位引导信息，将车辆引导至合适停车场。",
        "inputs": "诱导屏位置、关联停车场、显示内容、空位阈值、发布时间和发布状态。",
        "processing": "系统根据实时空位和区域热度生成诱导内容，管理员可编辑并模拟发布；发布记录用于后续统计和审计。",
        "outputs": "诱导屏内容、发布状态、发布时间、关联停车场和诱导策略记录。",
        "rules": "诱导内容必须与实时空位状态一致；停车场满位时不得继续诱导车辆驶入。"
    },
    {
        "id": "M5",
        "name": "智能导航与车位预约",
        "source": "课题要求(5)",
        "priority": "P0",
        "owner": "程子浩",
        "actor": "私家车主",
        "use_cases": ["输入目的地", "查看推荐方案", "预约车位", "锁定15分钟", "到场确认", "取消预约"],
        "description": "用户通过APP/小程序输入目的地后，系统推荐周边最优停车方案，支持在线预约车位并锁定15分钟。",
        "inputs": "目的地、当前位置、停车偏好、车位ID、预计到达时间、预约取消/确认指令。",
        "processing": "系统按距离、价格、空余数和预约可用性排序；预约后将车位状态置为锁定，超时未确认自动释放，到场确认后转为占用。",
        "outputs": "推荐停车方案、预约单、倒计时、车位锁定状态、导航提示和预约历史。",
        "rules": "同一用户同一时间只能有一个活跃预约；车位锁定15分钟；取消或超时后必须释放车位。"
    },
    {
        "id": "M6",
        "name": "车牌识别与无感支付",
        "source": "课题要求(6)",
        "priority": "P0",
        "owner": "程子浩",
        "actor": "私家车主/停车场管理员",
        "use_cases": ["模拟车牌入场", "生成停车订单", "自动计费", "确认无感支付", "释放车位"],
        "description": "停车场出入口部署AI车牌识别相机，入场自动识别车牌并抬杆，出场自动计费并从绑定账户扣款。实训阶段使用模拟识别与模拟支付。",
        "inputs": "车牌号、入场时间、出场时间、停车场费率、绑定账户状态、支付确认。",
        "processing": "系统根据车牌号生成入场订单，出场时计算停车时长和费用，模拟扣费成功后完成订单并释放车位。",
        "outputs": "车牌识别结果、停车订单、费用明细、支付状态、车位释放结果。",
        "rules": "不足1小时按1小时计费；首15分钟可免费；支付成功后订单状态为已完成。"
    },
    {
        "id": "M7",
        "name": "停车行为大数据分析",
        "source": "课题要求(7)",
        "priority": "P1",
        "owner": "程晓洋",
        "actor": "停车场管理员/交管部门",
        "use_cases": ["查看周转率", "分析平均时长", "查看高峰饱和度", "分析区域热度", "导出分析结果"],
        "description": "基于车位使用数据和用户停车行为，分析车位周转率、平均停车时长、高峰饱和度等指标。",
        "inputs": "停车订单、车位状态变化、预约记录、区域信息、统计时间范围。",
        "processing": "系统按日、周、月聚合停车行为数据，计算周转率、平均时长、饱和度和区域热力趋势。",
        "outputs": "统计图表、周转率排行、时长分布、高峰时段热力和导出报表。",
        "rules": "统计指标应支持按停车场、区域和时间范围筛选；复杂统计响应时间不超过1秒。"
    },
    {
        "id": "M8",
        "name": "动态定价策略引擎",
        "source": "课题要求(8)",
        "priority": "P1",
        "owner": "程晓洋",
        "actor": "停车场管理员/交管部门",
        "use_cases": ["配置价格规则", "设置时段系数", "模拟价格计算", "查看调价记录", "发布价格说明"],
        "description": "基于区域热度、时段、利用率、天气、周边活动等因素动态调整停车收费标准。",
        "inputs": "基础费率、时段、区域热度、车位利用率、天气/活动标签和调价上限。",
        "processing": "系统根据规则计算动态价格；实训阶段采用规则引擎模拟强化学习策略，并保留价格公示信息。",
        "outputs": "动态费率、调价原因、调价记录、价格公示内容和模拟收益对比。",
        "rules": "高峰期价格浮动不超过固定价格200%；价格调整必须可解释、可追溯、可公示。"
    },
    {
        "id": "M9",
        "name": "共享停车管理平台",
        "source": "课题要求(9)",
        "priority": "P1",
        "owner": "丁梓钊",
        "actor": "车位主/私家车主",
        "use_cases": ["发布共享车位", "设置共享时段", "浏览共享车位", "预约共享车位", "查看收益"],
        "description": "推动机关单位、商业楼宇、住宅小区闲置车位错时共享，车位主发布时段和价格，用户在线预约。",
        "inputs": "车位位置、可共享时段、共享价格、车位主信息、预约用户信息。",
        "processing": "系统校验时段冲突，发布可预约共享车位；用户预约后锁定时段并生成共享订单。",
        "outputs": "共享车位列表、预约记录、收益统计、时段占用状态。",
        "rules": "同一共享车位同一时段不得重复预约；车位主可下架未被预约的共享时段。"
    },
    {
        "id": "M10",
        "name": "充电车位智能管理",
        "source": "课题要求(10)",
        "priority": "P2",
        "owner": "丁梓钊",
        "actor": "新能源车主/停车场管理员",
        "use_cases": ["查看充电车位", "预约充电车位", "更新充电状态", "识别异常占用", "处理故障告警"],
        "description": "针对新能源充电车位，集成充电桩状态监测，支持充电+停车一体化预约，对燃油车占用充电车位行为告警。",
        "inputs": "充电车位编号、充电桩状态、车辆类型、预约时间、故障信息。",
        "processing": "系统区分普通车位和充电车位，展示空闲、充电中、故障状态；异常占用时生成告警。",
        "outputs": "充电车位状态、预约结果、故障告警、异常占用记录。",
        "rules": "故障充电桩不可预约；燃油车占用充电车位应标记为异常。"
    },
    {
        "id": "M11",
        "name": "违停自动抓拍与取证",
        "source": "课题要求(11)",
        "priority": "P2",
        "owner": "丁梓钊",
        "actor": "交管部门/停车场管理员",
        "use_cases": ["配置禁停区域", "模拟违停抓拍", "生成取证记录", "审核证据", "推送处理结果"],
        "description": "对禁停区域、消防通道、公交专用道等违停行为，AI摄像机自动抓拍取证并生成标准化证据。",
        "inputs": "禁停区域、车牌号、违停类型、抓拍时间、远景/近景/车牌特写描述。",
        "processing": "系统生成违停记录，关联车牌和位置；人工审核后进入待处理或已处理状态。",
        "outputs": "违停证据记录、审核结果、处理状态、统计报表。",
        "rules": "违停证据至少包含时间、地点、车牌、类型和图片描述；消防通道违停优先级最高。"
    },
    {
        "id": "M12",
        "name": "长期停车月卡管理",
        "source": "课题要求(12)",
        "priority": "P2",
        "owner": "丁梓钊",
        "actor": "私家车主/停车场管理员",
        "use_cases": ["申请月卡", "审批月卡", "续费月卡", "绑定车辆", "查看到期提醒"],
        "description": "支持用户在线申请和续费月卡、季卡、年卡，系统自动计算有效期限和费用，到期前自动提醒。",
        "inputs": "用户信息、车牌号、车场ID、卡类型、有效期、支付状态。",
        "processing": "系统根据卡类型计算有效期和费用，管理员审批后生效，到期前生成提醒。",
        "outputs": "月卡申请单、审批结果、有效期、续费记录和到期提醒。",
        "rules": "同一车辆在同一车场同一时间只能绑定一张有效月卡；到期前7天提醒。"
    },
    {
        "id": "M13",
        "name": "车场运营管理后台",
        "source": "课题要求(13)",
        "priority": "P0",
        "owner": "程晓洋",
        "actor": "停车场管理员",
        "use_cases": ["查看运营概览", "管理车位状态", "查看收入统计", "查看进出场流量", "处理设备故障", "远程开闸"],
        "description": "停车场管理员通过Web端或移动端查看实时车位状态、收入统计、进出场流量和设备运行状态，支持远程操作。",
        "inputs": "管理员账号、车场ID、车位状态、订单数据、设备状态、操作原因。",
        "processing": "系统校验管理员权限，聚合车位、订单、收入、设备数据，提供车位释放、远程开闸和故障报修入口。",
        "outputs": "运营概览、车位列表、收入报表、进出场流水、设备状态和操作日志。",
        "rules": "管理员只能管理自己负责的车场；关键操作需二次确认并记录日志。"
    },
    {
        "id": "M14",
        "name": "车主移动端APP/小程序",
        "source": "课题要求(14)",
        "priority": "P0",
        "owner": "丁梓钊",
        "actor": "私家车主",
        "use_cases": ["注册登录", "搜索停车场", "预约车位", "在线支付", "查询停车记录", "申请电子发票"],
        "description": "为车主提供停车场搜索、车位预约、在线支付、停车记录查询、电子发票、反向寻车导航等全流程停车服务。",
        "inputs": "用户账号、目的地、车牌号、预约信息、支付确认、发票信息。",
        "processing": "移动端整合搜索、预约、支付、记录、发票和寻车入口；调用后端统一API完成业务闭环。",
        "outputs": "车主首页、停车场列表、预约记录、订单支付状态、停车记录和发票申请结果。",
        "rules": "车主端支持移动端响应式布局；未登录用户只能浏览基础信息，预约和支付需登录。"
    },
    {
        "id": "M15",
        "name": "反向寻车导航",
        "source": "课题要求(15)",
        "priority": "P1",
        "owner": "程子浩",
        "actor": "私家车主",
        "use_cases": ["输入车牌号", "定位当前车位", "查看步行路线", "查看楼层区域", "结束寻车"],
        "description": "用户在大型停车场返回取车时，在APP中输入车牌号或车位号，系统规划从当前位置到停车位的步行路线。",
        "inputs": "车牌号、车位号、当前入口/楼层、停车订单状态。",
        "processing": "系统查询进行中停车订单，定位车位楼层、区域和编号，生成文字路线和简化示意图。",
        "outputs": "车位位置、楼层区域、步行指引、寻车结果提示。",
        "rules": "反向寻车仅对进行中停车记录有效；未找到车辆时提示检查车牌号或车位号。"
    },
    {
        "id": "M16",
        "name": "设备运维管理平台",
        "source": "课题要求(16)",
        "priority": "P2",
        "owner": "程晓洋",
        "actor": "设备运维人员",
        "use_cases": ["查看设备列表", "监控在线状态", "登记故障", "远程诊断", "查看运行日志"],
        "description": "对全市停车场设备和路内检测设备集中管理，实时监控在线状态、故障告警、固件版本和运行日志。",
        "inputs": "设备编号、设备类型、所属车场、在线状态、故障类型、固件版本。",
        "processing": "系统汇总摄像头、道闸、诱导屏、充电桩等设备状态，生成故障告警并支持状态更新。",
        "outputs": "设备台账、在线率、故障告警、运行日志和运维处理记录。",
        "rules": "故障设备应进入告警列表；运维处理需记录处理人、时间和结果。"
    },
    {
        "id": "M17",
        "name": "城市停车监管平台",
        "source": "课题要求(17)",
        "priority": "P1",
        "owner": "程晓洋",
        "actor": "城市交管部门",
        "use_cases": ["查看监管大屏", "查看供需分析", "查看收费统计", "查看投诉热力", "评估停车政策"],
        "description": "为城市交通管理部门提供全市停车态势监控大屏，展示供需分析、收费统计、周转率排名、投诉热力图等。",
        "inputs": "全市停车数据、区域信息、收费统计、投诉数据、政策评估时间范围。",
        "processing": "系统聚合多车场和路内停车数据，形成监管指标和区域对比，为政策评估提供量化依据。",
        "outputs": "监管大屏、供需图表、收费统计、周转率排行、投诉热力和政策评估报告。",
        "rules": "监管平台展示汇总数据，不暴露用户敏感信息；指标按区域和时间维度可筛选。"
    },
    {
        "id": "M18",
        "name": "商圈停车联合营销",
        "source": "课题要求(18)",
        "priority": "P2",
        "owner": "程子浩",
        "actor": "商圈运营人员/停车场管理员",
        "use_cases": ["创建营销活动", "配置优惠规则", "绑定合作商户", "核销停车券", "查看活动效果"],
        "description": "与周边商圈、酒店、医院等合作，支持消费免停车费、积分兑换停车券、节假日停车优惠等营销活动配置。",
        "inputs": "活动名称、合作商户、优惠规则、有效期、核销记录、适用停车场。",
        "processing": "系统创建营销活动并绑定停车优惠规则，用户满足条件后核销停车券或减免停车费。",
        "outputs": "活动列表、优惠规则、停车券、核销记录、活动效果统计。",
        "rules": "活动必须设置有效期和适用范围；同一订单不可重复叠加同类优惠。"
    },
]


def get_font(size, bold=False):
    candidates = [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]
    for path in candidates:
        if path and os.path.exists(path):
            try:
                return ImageFont.truetype(path, size)
            except OSError:
                pass
    return ImageFont.load_default()


def wrap_text(draw, text, font, max_width):
    lines, line = [], ""
    for ch in text:
        trial = line + ch
        if draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
            line = trial
        else:
            if line:
                lines.append(line)
            line = ch
    if line:
        lines.append(line)
    return lines


def center_text(draw, box, text, font, fill=(31, 53, 82), max_width=None):
    x1, y1, x2, y2 = box
    max_width = max_width or (x2 - x1 - 24)
    lines = wrap_text(draw, text, font, max_width)
    line_height = font.size + 6
    total_height = len(lines) * line_height
    y = y1 + (y2 - y1 - total_height) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        x = x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, font=font, fill=fill)
        y += line_height


def build_use_case_diagram(module):
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    path = DIAGRAM_DIR / f"{module['id']}_use_case.png"
    width, height = 1500, 820
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = get_font(34, bold=True)
    label_font = get_font(25)
    small_font = get_font(21)
    line_color = (60, 92, 130)
    border = (73, 120, 172)
    fill = (238, 246, 255)
    actor_color = (40, 55, 70)

    draw.text((60, 38), f"{module['id']} {module['name']} 用例图", font=title_font, fill=(24, 68, 120))

    ax, ay = 190, 255
    draw.ellipse((ax - 35, ay - 85, ax + 35, ay - 15), outline=actor_color, width=5)
    draw.line((ax, ay - 15, ax, ay + 110), fill=actor_color, width=5)
    draw.line((ax - 70, ay + 30, ax + 70, ay + 30), fill=actor_color, width=5)
    draw.line((ax, ay + 110, ax - 65, ay + 205), fill=actor_color, width=5)
    draw.line((ax, ay + 110, ax + 65, ay + 205), fill=actor_color, width=5)
    center_text(draw, (60, 515, 330, 610), module["actor"], small_font, fill=actor_color)

    sys_box = (410, 135, 1430, 760)
    draw.rounded_rectangle(sys_box, radius=28, outline=border, width=4, fill=(250, 253, 255))
    draw.text((450, 160), "AI-driven城市智慧停车管理与诱导系统", font=small_font, fill=(78, 92, 108))

    cases = module["use_cases"]
    cols = 2 if len(cases) <= 4 else 3
    rows = (len(cases) + cols - 1) // cols
    oval_w = 275
    oval_h = 108 if rows <= 2 else 96
    start_x = 495
    gap_x = (830 - cols * oval_w) // max(cols - 1, 1) if cols > 1 else 0
    start_y = 245
    gap_y = 78 if rows <= 2 else 42

    centers = []
    for index, use_case in enumerate(cases):
        col = index % cols
        row = index // cols
        x1 = start_x + col * (oval_w + gap_x)
        y1 = start_y + row * (oval_h + gap_y)
        x2, y2 = x1 + oval_w, y1 + oval_h
        draw.ellipse((x1, y1, x2, y2), fill=fill, outline=border, width=4)
        center_text(draw, (x1 + 14, y1 + 8, x2 - 14, y2 - 8), use_case, label_font, max_width=oval_w - 50)
        centers.append((x1 + oval_w // 2, y1 + oval_h // 2))

    for cx, cy in centers:
        draw.line((300, 360, cx - oval_w // 2 + 12, cy), fill=line_color, width=3)

    image.save(path, quality=95)
    return path


def draw_box(draw, box, text, font, fill=(244, 248, 252), outline=(73, 120, 172), text_fill=(31, 53, 82)):
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=3)
    center_text(draw, box, text, font, fill=text_fill, max_width=box[2] - box[0] - 26)


def build_function_structure_diagram():
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    path = DIAGRAM_DIR / "system_function_structure.png"
    width, height = 1600, 980
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = get_font(36, bold=True)
    group_font = get_font(26, bold=True)
    item_font = get_font(21)
    draw.text((70, 42), "AI-driven城市智慧停车管理与诱导系统功能结构图", font=title_font, fill=(24, 68, 120))
    root_box = (410, 110, 1190, 175)
    draw_box(draw, root_box, "智慧停车管理与诱导系统", group_font, fill=(230, 242, 255))

    groups = [
        ("车主服务", ["城市停车一张图", "智能导航与车位预约", "无感支付", "移动端APP/小程序", "反向寻车", "停车联合营销"]),
        ("车场运营", ["停车场联网接入", "停车诱导发布", "动态定价", "月卡管理", "运营后台", "共享停车"]),
        ("城市监管", ["路内视觉检测", "违停抓拍取证", "停车大数据分析", "城市监管平台"]),
        ("设备与能源", ["充电车位管理", "设备运维平台", "摄像头/道闸/诱导屏模拟"]),
    ]
    start_x, start_y = 85, 260
    group_w, group_h = 340, 72
    item_h, item_gap = 58, 18
    for idx, (group, items) in enumerate(groups):
        x = start_x + idx * 380
        draw.line((800, 175, x + group_w // 2, start_y), fill=(84, 110, 136), width=3)
        draw_box(draw, (x, start_y, x + group_w, start_y + group_h), group, group_font, fill=(238, 246, 255))
        y = start_y + group_h + 34
        for item in items:
            draw.line((x + group_w // 2, y - 34, x + group_w // 2, y), fill=(125, 145, 165), width=2)
            draw_box(draw, (x + 18, y, x + group_w - 18, y + item_h), item, item_font, fill=(252, 254, 255), outline=(136, 166, 196))
            y += item_h + item_gap
    image.save(path, quality=95)
    return path


def build_system_use_case_diagram():
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    path = DIAGRAM_DIR / "system_overall_use_case.png"
    width, height = 1600, 940
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = get_font(36, bold=True)
    actor_font = get_font(23, bold=True)
    use_font = get_font(22)
    draw.text((70, 42), "系统总体用例图", font=title_font, fill=(24, 68, 120))
    sys_box = (360, 120, 1260, 850)
    draw.rounded_rectangle(sys_box, radius=28, outline=(73, 120, 172), width=4, fill=(250, 253, 255))
    draw.text((445, 145), "AI-driven城市智慧停车管理与诱导系统", font=actor_font, fill=(78, 92, 108))

    actors = [
        ("私家车主", (155, 260), ["查找车位", "预约导航", "支付寻车"]),
        ("停车场管理员", (155, 610), ["运营管理", "价格营销", "订单处理"]),
        ("城市交管部门", (1415, 260), ["监管态势", "违停审核", "政策评估"]),
        ("设备运维人员", (1415, 610), ["设备监控", "告警处理", "维护记录"]),
    ]
    use_cases = [
        ("查看停车一张图", 470, 240), ("预约车位并导航", 760, 240), ("车牌识别无感支付", 1010, 240),
        ("接入停车场数据", 470, 410), ("发布诱导与价格", 760, 410), ("管理共享/月卡/营销", 1010, 410),
        ("分析停车行为", 470, 590), ("违停抓拍取证", 760, 590), ("城市停车监管", 1010, 590),
        ("设备运维管理", 760, 735),
    ]
    centers = []
    for text, cx, cy in use_cases:
        box = (cx - 120, cy - 48, cx + 120, cy + 48)
        draw.ellipse(box, fill=(238, 246, 255), outline=(73, 120, 172), width=3)
        center_text(draw, box, text, use_font)
        centers.append((cx, cy))
    for actor, (ax, ay), labels in actors:
        draw.ellipse((ax - 30, ay - 75, ax + 30, ay - 15), outline=(40, 55, 70), width=5)
        draw.line((ax, ay - 15, ax, ay + 95), fill=(40, 55, 70), width=5)
        draw.line((ax - 58, ay + 24, ax + 58, ay + 24), fill=(40, 55, 70), width=5)
        draw.line((ax, ay + 95, ax - 55, ay + 175), fill=(40, 55, 70), width=5)
        draw.line((ax, ay + 95, ax + 55, ay + 175), fill=(40, 55, 70), width=5)
        center_text(draw, (ax - 120, ay + 190, ax + 120, ay + 245), actor, actor_font, fill=(40, 55, 70))
        for label in labels:
            target = next((c for c in centers if label[:2] in "".join(str(c))), None)
        if ax < 800:
            line_targets = centers[:6] if ay < 400 else centers[3:10]
            start = (ax + 75, ay + 40)
        else:
            line_targets = centers[6:] if ay < 400 else centers[8:]
            start = (ax - 75, ay + 40)
        for cx, cy in line_targets[:4]:
            draw.line((start[0], start[1], cx, cy), fill=(120, 145, 170), width=2)
    image.save(path, quality=95)
    return path


def build_er_diagram():
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    path = DIAGRAM_DIR / "core_er_diagram.png"
    width, height = 1600, 940
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = get_font(36, bold=True)
    entity_font = get_font(23, bold=True)
    small_font = get_font(19)
    draw.text((70, 42), "核心E-R关系图", font=title_font, fill=(24, 68, 120))
    entities = {
        "users 用户": (90, 140, 310, 215),
        "vehicles 车辆": (405, 140, 625, 215),
        "reservations 预约": (720, 140, 940, 215),
        "parking_orders 订单": (1035, 140, 1255, 215),
        "parking_lots 停车场": (405, 390, 625, 465),
        "parking_spots 车位": (720, 390, 940, 465),
        "devices 设备": (1035, 390, 1255, 465),
        "shared_spots 共享车位": (90, 650, 310, 725),
        "violations 违停证据": (405, 650, 625, 725),
        "marketing_campaigns 营销": (720, 650, 980, 725),
    }
    for name, box in entities.items():
        draw_box(draw, box, name, entity_font, fill=(238, 246, 255))
    relations = [
        ("users 用户", "vehicles 车辆", "拥有 1:N"),
        ("users 用户", "reservations 预约", "发起 1:N"),
        ("reservations 预约", "parking_orders 订单", "生成 1:1"),
        ("parking_lots 停车场", "parking_spots 车位", "包含 1:N"),
        ("parking_spots 车位", "reservations 预约", "被预约 1:N"),
        ("parking_spots 车位", "parking_orders 订单", "产生订单 1:N"),
        ("parking_lots 停车场", "devices 设备", "部署 1:N"),
        ("parking_spots 车位", "shared_spots 共享车位", "共享 1:0..N"),
        ("parking_spots 车位", "violations 违停证据", "关联 1:N"),
        ("parking_orders 订单", "marketing_campaigns 营销", "核销 N:1"),
    ]
    def center(box):
        return ((box[0] + box[2]) // 2, (box[1] + box[3]) // 2)
    for left, right, label in relations:
        x1, y1 = center(entities[left])
        x2, y2 = center(entities[right])
        draw.line((x1, y1, x2, y2), fill=(100, 126, 152), width=3)
        mx, my = (x1 + x2) // 2, (y1 + y2) // 2
        draw.rectangle((mx - 72, my - 17, mx + 72, my + 17), fill="white")
        center_text(draw, (mx - 74, my - 20, mx + 74, my + 20), label, small_font, fill=(60, 78, 96))
    add_note = "说明：实训阶段采用SQLite保存核心业务对象，外部AI、GIS、支付、硬件数据通过模拟字段和预留接口接入。"
    center_text(draw, (120, 820, 1480, 885), add_note, small_font, fill=(68, 82, 96))
    image.save(path, quality=95)
    return path


def set_run_black(run, size=None, bold=None):
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.color.rgb = RGBColor(0, 0, 0)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_para(doc, text="", style=None, size=10.5, bold=None, align=None):
    paragraph = doc.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    if text:
        run = paragraph.add_run(text)
        set_run_black(run, size=size, bold=bold)
    return paragraph


def add_heading_black(doc, text, level=1):
    paragraph = doc.add_heading("", level=level)
    run = paragraph.add_run(text)
    set_run_black(run, size=HEADING_SIZE.get(level, 10.5), bold=True)
    return paragraph


def set_table_black(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_black(run, size=9.5, bold=run.bold)


def add_table_black(doc, headers, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for col, header in enumerate(headers):
        cell = table.cell(0, col)
        cell.text = str(header)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_black(run, size=9.5, bold=True)
    for row_index, row in enumerate(rows, start=1):
        for col, value in enumerate(row):
            table.cell(row_index, col).text = str(value)
    set_table_black(table)
    add_para(doc)
    return table


def configure_clean_doc(doc):
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Pt(81.1)
    section.right_margin = Pt(64.3)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.line_spacing = None
    normal.paragraph_format.space_after = None
    for style_name, size in [("Heading 1", 22), ("Heading 2", 16), ("Heading 3", 16)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)


def add_clean_picture(doc, path, caption):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(6.2))
    cap = add_para(doc, caption, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    return cap


def add_black_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    set_run_black(run, size=10.5)
    return paragraph


def add_section_page_break(doc):
    doc.add_page_break()


def add_main_module_detail(doc, module, index):
    add_heading_black(doc, f"3.2.{index} {module['id']} {module['name']}", level=3)
    add_clean_picture(doc, build_use_case_diagram(module), f"图3-2-{index} {module['id']} {module['name']}用例图")
    add_para(doc, f"功能说明：{module['description']}")
    add_para(doc, f"参与角色：{module['actor']}。")
    add_para(doc, f"主要用例：{'、'.join(module['use_cases'])}。")
    add_para(doc, f"业务流程：用户或管理员首先进入“{module['name']}”功能入口，系统依据当前角色加载可操作数据；随后执行{'、'.join(module['use_cases'][:3])}等核心操作，操作结果写入业务表并同步到相关统计、告警或消息模块。流程结束后，系统保留操作时间、操作人、状态变化和异常原因，便于后续测试、审计和问题复盘。")
    add_para(doc, f"输入：{module['inputs']}")
    add_para(doc, f"处理：{module['processing']}")
    add_para(doc, f"输出：{module['outputs']}")
    add_para(doc, f"业务规则：{module['rules']}")
    add_para(doc, f"接口与数据联动：{INTEGRATION_NOTES[module['id']]}")
    add_para(doc, f"异常处理与边界：{EXCEPTION_NOTES[module['id']]}")
    add_para(doc, f"权限与审计：仅授权角色可新增、修改、发布或审核该模块关键数据；车主只能访问本人车辆、订单、预约和优惠信息；管理员、运维人员和监管人员的关键操作必须写入操作日志。")
    add_para(doc, f"验收标准：测试用例TC-{module['id']}至少覆盖正常新增/查询/更新流程、异常输入、权限拦截、状态流转、数据同步和页面提示；该模块优先级为{module['priority']}，需求负责人为{module['owner']}。验收时需检查页面展示、后端接口返回、数据库记录和关联模块状态是否一致。")


def add_data_dictionary(doc):
    add_heading_black(doc, "3.3 数据字典", level=2)
    add_para(doc, "本节在正文中给出核心数据字典，保证设计、编码和测试可直接依据。字段命名以当前项目实现为准，真实AI、GIS、支付和硬件能力采用模拟字段或预留接口承载。")
    add_heading_black(doc, "3.3.1 数据字典", level=3)
    add_para(doc, "核心数据字典覆盖用户、车辆、停车场、车位、预约、订单、共享车位、违停证据、设备和营销活动等对象。后续详细设计和编码阶段应保持字段名称、含义和状态枚举一致。")
    tables = [
        ("users（用户表）", [("id", "INTEGER", "主键"), ("username", "VARCHAR(50)", "用户名"), ("password_hash", "VARCHAR(255)", "密码哈希"), ("phone", "VARCHAR(20)", "手机号"), ("role", "VARCHAR(20)", "用户角色：车主、管理员、监管、运维等")]),
        ("vehicles（车辆表）", [("id", "INTEGER", "主键"), ("user_id", "INTEGER", "所属用户"), ("plate_number", "VARCHAR(20)", "车牌号"), ("vehicle_type", "VARCHAR(20)", "车辆类型：燃油车/新能源车")]),
        ("parking_lots（停车场表）", [("id", "INTEGER", "主键"), ("name", "VARCHAR(100)", "停车场名称"), ("address", "VARCHAR(255)", "地址"), ("region", "VARCHAR(50)", "所属区域"), ("total_spots", "INTEGER", "总车位数"), ("rate_per_hour", "DECIMAL", "基础费率"), ("sync_status", "VARCHAR(20)", "联网接入状态")]),
        ("parking_spots（车位表）", [("id", "INTEGER", "主键"), ("lot_id", "INTEGER", "所属停车场"), ("spot_number", "VARCHAR(20)", "车位编号"), ("status", "VARCHAR(20)", "free/locked/occupied/disabled"), ("spot_type", "VARCHAR(20)", "normal/charging/shared/roadside"), ("last_detected_at", "DATETIME", "最后检测时间")]),
        ("reservations（预约表）", [("id", "INTEGER", "主键"), ("user_id", "INTEGER", "预约用户"), ("spot_id", "INTEGER", "预约车位"), ("status", "VARCHAR(20)", "active/confirmed/cancelled/expired"), ("created_at", "DATETIME", "预约创建时间"), ("expire_at", "DATETIME", "15分钟锁定过期时间")]),
        ("parking_orders（停车订单表）", [("id", "INTEGER", "主键"), ("user_id", "INTEGER", "用户"), ("spot_id", "INTEGER", "车位"), ("plate_number", "VARCHAR(20)", "车牌号"), ("entry_time", "DATETIME", "入场时间"), ("exit_time", "DATETIME", "出场时间"), ("amount", "DECIMAL", "费用"), ("pay_status", "VARCHAR(20)", "支付状态"), ("status", "VARCHAR(20)", "订单状态")]),
        ("shared_spots（共享车位表）", [("id", "INTEGER", "主键"), ("owner_id", "INTEGER", "车位主"), ("spot_id", "INTEGER", "共享车位"), ("available_start", "DATETIME", "可共享开始时间"), ("available_end", "DATETIME", "可共享结束时间"), ("price", "DECIMAL", "共享价格"), ("status", "VARCHAR(20)", "发布/预约/下架状态")]),
        ("violations（违停证据表）", [("id", "INTEGER", "主键"), ("plate_number", "VARCHAR(20)", "车牌号"), ("location", "VARCHAR(255)", "违停位置"), ("violation_type", "VARCHAR(50)", "违停类型"), ("evidence_desc", "TEXT", "证据描述"), ("review_status", "VARCHAR(20)", "审核状态"), ("status", "VARCHAR(20)", "处理状态")]),
        ("devices（设备表）", [("id", "INTEGER", "主键"), ("lot_id", "INTEGER", "所属停车场"), ("device_type", "VARCHAR(30)", "摄像头/道闸/诱导屏/充电桩"), ("status", "VARCHAR(20)", "在线/离线/故障"), ("firmware_version", "VARCHAR(50)", "固件版本"), ("last_heartbeat", "DATETIME", "最后心跳")]),
        ("marketing_campaigns（营销活动表）", [("id", "INTEGER", "主键"), ("name", "VARCHAR(100)", "活动名称"), ("merchant", "VARCHAR(100)", "合作商户"), ("rule_desc", "TEXT", "优惠规则"), ("valid_from", "DATETIME", "开始时间"), ("valid_to", "DATETIME", "结束时间"), ("status", "VARCHAR(20)", "活动状态")]),
    ]
    for idx, (title, rows) in enumerate(tables, start=1):
        add_para(doc, f"表3-3-{idx} {title}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_table_black(doc, ["字段", "类型", "说明"], rows)
    add_heading_black(doc, "3.3.2 E-R关系图", level=3)
    add_para(doc, "用户与车辆、预约、订单、月卡、共享车位存在一对多关系；停车场与车位、设备、诱导屏、订单存在一对多关系；预约确认后生成停车订单；订单完成后参与收入统计、行为分析和监管汇总；违停证据、设备告警和营销核销分别关联停车场、车位或订单。")
    add_clean_picture(doc, build_er_diagram(), "图3-3 核心E-R关系图")


def add_front_matter(doc):
    title = add_para(doc, f"{PROJECT_NAME}需求规格说明书", size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    title.paragraph_format.space_after = Pt(18)
    add_table_black(
        doc,
        ["项目", "内容"],
        [
            ("文档编号", "SmartPark-SRS-001"),
            ("版本号", "V1.0"),
            ("小组", GROUP),
            ("编写人", "程晓洋"),
            ("项目成员", "程晓洋、程子浩、丁梓钊"),
            ("编写日期", DATE),
            ("适用范围", "需求分析、系统设计、编码实现、测试验收与课程提交"),
        ],
    )
    add_heading_black(doc, "文档修订记录", level=1)
    add_table_black(
        doc,
        ["日期", "版本", "变更编号", "章节", "修改内容", "修改人"],
        [
            ("2026.06.01", "V0.1", "N/A", "全部", "按项目要求18个一级功能模块建立需求规格框架", "程子浩（RA）"),
            ("2026.06.01", "V0.2", "N/A", "3.2", "补充每个模块的输入、处理、输出、业务规则和用例图", "丁梓钊（PM）"),
            ("2026.06.02", "V0.3", "IR-01~IR-08", "全部", "依据评审意见统一18模块口径并补充非功能需求", "全员评审"),
            ("2026.06.02", "V1.0", "N/A", "全部", "正式版发布，前置需求追踪矩阵、评审记录和数据字典", "程晓洋（PM）"),
        ],
    )
    add_heading_black(doc, "关键字", level=1)
    add_para(doc, "智慧停车；车位预约；城市停车一张图；无感支付；停车诱导；数据联网接入；停车监管；共享停车")
    add_heading_black(doc, "摘要", level=1)
    add_para(doc, f"本文档描述{PROJECT_NAME}的业务背景、18个一级功能模块、数据字典、接口约束、性能要求、质量属性、需求分级和需求评审结论。文档严格以项目要求.txt中的主体要求为依据，避免将注册登录、基础CRUD、订单管理等实现支撑项误列为独立一级模块。")
    add_heading_black(doc, "术语与缩略语", level=1)
    add_table_black(
        doc,
        ["缩写", "英文全称", "中文解释"],
        [
            ("SRS", "Software Requirement Specification", "软件需求规格说明书"),
            ("UC", "Use Case", "用例"),
            ("RTM", "Requirement Traceability Matrix", "需求追踪矩阵"),
            ("BR", "Business Rule", "业务规则"),
            ("RA", "Requirement Analyst", "需求分析师"),
            ("GIS", "Geographic Information System", "地理信息系统"),
        ],
    )


def add_main_content(doc):
    add_section_page_break(doc)
    add_heading_black(doc, "1 简介", level=1)
    add_para(doc, "本章说明本文档的编写目的、需求范围和读者对象，为后续功能需求、数据需求和非功能需求提供阅读边界。")
    add_heading_black(doc, "1.1 目的", level=2)
    add_para(doc, f"本文档用于明确{PROJECT_NAME}在实训阶段需要实现和模拟呈现的需求范围，作为后续概要设计、详细设计、编码实现、测试用例编写和课程验收的统一依据。文档以项目要求.txt为最高优先级输入，所有一级功能模块均按原始18项课题要求展开。")
    add_para(doc, "本文档同时给出需求来源、优先级、责任人、输入输出、业务规则、异常边界和验收标准，便于开发阶段按模块拆分任务，也便于测试阶段按需求追踪矩阵逐项核验。")
    add_para(doc, "该说明书的预期读者包括项目组成员、指导教师、需求评审人员、开发人员、测试人员和后续维护人员。开发人员依据本文档拆分接口、数据库和前端页面；测试人员依据本文档编写功能测试、异常测试和验收测试；指导教师依据本文档检查项目是否贴合智慧停车课题要求。")
    add_heading_black(doc, "1.2 范围", level=2)
    add_para(doc, "本文档覆盖18个一级功能模块的功能需求、模块用例图、输入处理输出、业务规则、异常处理、接口联动、优先级、责任人、需求追踪矩阵、评审记录、核心数据字典、接口需求、性能需求、设计约束和质量属性。")
    add_para(doc, "实训阶段不包含真实AI模型训练、真实支付通道、真实GIS地图服务和真实硬件设备部署，但需要通过模拟数据、状态字段、接口预留和可视化界面完整表达业务闭环。涉及外部能力的功能不得删除，应明确模拟实现方式和后续扩展口径。")

    add_section_page_break(doc)
    add_heading_black(doc, "2 总体概述", level=1)
    add_para(doc, "本章从项目背景、运行环境、用户角色和实施假设四个方面概述系统整体情况，确保18个一级功能模块都服务于同一业务目标。")
    add_heading_black(doc, "2.1 软件概述", level=2)
    add_para(doc, "本节按照模板要求概述软件项目和产品环境。系统定位为面向车主、停车场运营方、城市监管方、设备运维方和商圈运营方的智慧停车管理与诱导平台。")
    add_heading_black(doc, "2.1.1 项目介绍", level=3)
    add_para(doc, "城市停车资源分布不均、路内外车位数据割裂、停车场信息不透明、车辆入出场效率低和监管统计滞后，是本项目要解决的核心问题。系统通过AI视觉检测、停车场数据接入、城市停车一张图、预约导航、无感支付、大数据分析和监管平台等能力，形成车主侧、运营侧和监管侧的停车业务闭环。")
    add_para(doc, "从车主视角看，系统需要降低找车位、排队入场、缴费离场和反向寻车的时间成本；从停车场视角看，系统需要提升车位周转率、设备可见性、价格配置能力和营销联动能力；从监管视角看，系统需要支撑城市停车资源统筹、违停治理、收费透明和停车政策评估。")
    add_heading_black(doc, "2.1.2 产品环境介绍", level=3)
    add_para(doc, "系统采用前后端分离B/S架构。前端建议使用Vue3、Element Plus、ECharts和Vite，后端使用FastAPI、SQLAlchemy和SQLite。实训阶段通过初始化数据、模拟接口和管理端录入方式呈现AI识别、地图、支付、设备接入等外部能力。")
    add_para(doc, "运行环境包括车主移动端页面、停车场运营管理后台、城市监管后台和设备状态模拟界面。数据库保存用户、车辆、停车场、车位、预约、订单、设备、违停、月卡、营销活动和统计指标等数据。前端页面以可操作流程为主，不能只做静态展示。")
    add_heading_black(doc, "2.2 软件功能", level=2)
    add_para(doc, "软件功能严格覆盖项目要求中的18个一级模块。核心功能包括路内停车位视觉检测、停车场数据联网接入、城市停车一张图、停车诱导信息发布、智能导航与车位预约、车牌识别与无感支付、停车行为大数据分析、动态定价、共享停车、充电车位管理、违停取证、月卡管理、运营后台、车主移动端、反向寻车、设备运维、城市监管和商圈联合营销。")
    add_para(doc, "这些功能不是孤立页面，而是组成三条业务主线：车主停车闭环、车场运营闭环和城市监管闭环。注册登录、权限控制、订单基础管理、图表展示和基础CRUD属于支撑能力，应嵌入对应一级模块。")
    add_clean_picture(doc, build_function_structure_diagram(), "图2-1 系统功能结构图")
    add_heading_black(doc, "2.3 用户特征", level=2)
    add_para(doc, "系统用户具有多角色、多场景特点。车主关注低成本、少等待、易操作；停车场管理员关注数据准确、订单清晰、设备可控；监管人员关注城市级供需、收费和违停治理；设备运维人员关注告警、心跳和维护记录；商圈运营人员关注活动配置、核销和效果评估。")
    add_table_black(
        doc,
        ["角色", "主要职责", "关注模块"],
        [
            ("私家车主", "搜索停车场、预约车位、导航、支付、寻车、查询记录", "M3、M5、M6、M10、M12、M14、M15、M18"),
            ("停车场管理员", "维护车场车位、查看订单收入、处理设备和运营异常", "M2、M4、M6、M8、M10、M13、M16、M18"),
            ("车位主", "发布共享车位、设置共享时段、查看收益", "M9"),
            ("城市交管部门", "查看停车态势、监管违停、评估政策", "M1、M3、M7、M11、M17"),
            ("设备运维人员", "维护摄像头、道闸、诱导屏、充电桩等设备状态", "M1、M4、M10、M16"),
            ("商圈运营人员", "配置商户优惠、核销停车券、评估活动效果", "M18"),
        ],
    )
    add_heading_black(doc, "2.4 假设和依赖关系", level=2)
    add_para(doc, "假设停车场、车位、订单、设备、营销和统计数据可以通过模拟数据初始化；AI视觉识别、GIS地图、支付扣款和硬件动作均以模拟结果、状态字段或预留接口呈现。系统依赖本地Python、Node.js、SQLite数据库和现代浏览器。")
    add_para(doc, "项目验收时以演示闭环为准：车主能完成查询、预约、入场、支付、寻车和活动使用；管理员能完成车场运营、价格配置、设备处理和营销管理；监管人员能查看停车态势、违停证据和统计分析。")

    add_section_page_break(doc)
    add_heading_black(doc, "3 具体需求", level=1)
    add_para(doc, "本章是需求规格的核心部分，按照模板结构描述系统用例、子功能模块、数据字典和E-R关系。18个一级功能模块均来自项目要求.txt，不新增与项目无关的模块。")
    add_heading_black(doc, "3.1 系统用例", level=2)
    add_para(doc, "系统总体用例分为车主停车闭环、车场运营闭环和城市监管闭环。车主侧流程为搜索目的地、查看停车一张图、预约车位、导航前往、车牌识别入场、出场无感支付、查询停车记录和反向寻车。运营侧流程为接入停车场数据、监控车位设备、发布诱导信息、配置动态定价、处理订单和营销活动。监管侧流程为查看全市停车态势、分析供需与收费、审核违停证据和评估停车政策。")
    add_para(doc, "系统总体用例的边界是停车业务本身。注册登录、权限管理、基础CRUD、订单增删改查、图表展示等属于支撑能力，必须服务于18个一级模块，不作为独立一级需求重复列出。")
    add_clean_picture(doc, build_system_use_case_diagram(), "图3-1 系统总体用例图")
    add_heading_black(doc, "3.2 子功能模块需求", level=2)
    add_para(doc, "本节严格按项目要求的18个一级功能模块展开。每个模块均包含用例图、功能说明、角色、输入、处理、输出、业务规则和验收要点，避免出现只有标题没有内容的空白模块。")
    for index, module in enumerate(MODULES, start=1):
        add_main_module_detail(doc, module, index)

    add_data_dictionary(doc)

    add_section_page_break(doc)
    add_heading_black(doc, "4 性能需求", level=1)
    add_para(doc, "性能需求以实训演示和本地开发环境为基准，重点保证常用页面响应及时、状态刷新可感知、统计分析不卡顿。涉及真实硬件、真实地图和真实支付的接口在本阶段采用模拟响应，因此验收时重点检查接口设计是否可替换、状态变更是否准确。")
    add_heading_black(doc, "4.1 时间性能需求", level=2)
    add_table_black(
        doc,
        ["指标", "要求", "说明"],
        [
            ("页面加载", "首次加载不超过2秒", "本地数据和常用页面需快速展示"),
            ("简单查询", "API响应不超过500ms", "停车场列表、车位状态、订单详情等"),
            ("复杂统计", "响应不超过1秒", "周转率、饱和度、区域热度等聚合指标"),
            ("状态刷新", "默认30秒", "车位、设备、诱导屏状态可配置刷新"),
            ("预约锁定", "15分钟", "超时未到场确认自动释放车位"),
            ("导出数据", "500条以内不超过3秒", "报表导出和统计导出"),
        ],
    )
    add_heading_black(doc, "4.2 系统开放性需求", level=2)
    add_para(doc, "系统应通过RESTful接口、统一数据模型和可替换模拟服务保持开放性。AI视觉、GIS地图、支付通道、设备接入、诱导屏发布等外部能力在实训阶段采用模拟接口，后续可用真实服务替换而不改变18个一级模块的业务边界。")
    add_heading_black(doc, "4.3 界面友好性需求", level=2)
    add_para(doc, "车主端页面应突出查找车位、预约、支付、寻车和优惠使用；管理端页面应突出表格筛选、状态标识、批量操作、图表统计和异常处理；监管端页面应突出全市态势、区域对比、违停审核和趋势分析。关键操作需有明确按钮、状态反馈和错误提示。")
    add_heading_black(doc, "4.4 系统可用性需求", level=2)
    add_para(doc, "系统在本地演示环境下应保证核心车主闭环、运营后台闭环和监管查看闭环可连续演示。外部能力不可用时，页面应以模拟数据或降级展示方式继续表达需求，不应出现空白页面或无法说明的中断流程。")
    add_heading_black(doc, "4.5 可管理性需求", level=2)
    add_para(doc, "管理员应能查看车场、车位、订单、设备、价格、营销、违停和统计数据的状态；运维人员应能查看设备告警和处理记录；监管人员应能查看汇总数据。关键配置、审核、发布和处理动作需要记录操作者、时间、对象和结果。")

    add_section_page_break(doc)
    add_heading_black(doc, "5 接口需求", level=1)
    add_para(doc, "本章定义用户界面、软件接口、硬件模拟接口和通信接口的基本要求，确保前端、后端、数据库和模拟外部能力能够协同工作。")
    add_heading_black(doc, "5.1 用户接口", level=2)
    add_para(doc, "车主端提供搜索、预约、支付、寻车、记录查询和发票申请界面；管理端提供车场、车位、订单、设备、违停、营销和运营统计界面；监管端提供停车态势、供需分析、收费统计、投诉热力和政策评估界面。")
    add_heading_black(doc, "5.2 软件接口", level=2)
    add_para(doc, "前后端采用RESTful JSON接口，统一响应格式为{code, msg, data}。后端通过SQLAlchemy访问SQLite。AI识别、地图、支付、设备接入和诱导屏发布在实训阶段以模拟接口呈现，后续可替换为真实服务。")
    add_heading_black(doc, "5.3 硬件接口", level=2)
    add_para(doc, "系统不直接连接真实硬件。摄像头、道闸、诱导屏、充电桩等设备通过设备表、运行日志和模拟状态面板管理。设备接口需保留设备编号、设备类型、运行状态、最后心跳、故障等级和处理记录。")
    add_heading_black(doc, "5.4 通讯接口", level=2)
    add_para(doc, "开发环境前端运行于localhost:5173，后端运行于localhost:8000，通信协议为HTTP/1.1。接口返回统一结构{code, msg, data}，其中code=0表示成功，code=400表示参数错误，code=401表示未登录或登录失效，code=403表示无权限，code=500表示服务端异常。未来可扩展WebSocket推送实时车位状态。")

    add_section_page_break(doc)
    add_heading_black(doc, "6 总体设计约束", level=1)
    add_heading_black(doc, "6.1 标准符合性", level=2)
    add_para(doc, "系统需求、设计和实现必须符合项目要求.txt中的18个一级模块，不得将注册登录、基础CRUD或订单管理误列为独立一级模块。接口命名、数据库字段和页面模块应保持一致，便于后续测试和文档交叉核验。")
    add_heading_black(doc, "6.2 硬件约束", level=2)
    add_para(doc, "实训环境不部署真实摄像头、道闸、诱导屏和充电桩。相关功能通过模拟数据、设备状态表、运行日志和预留接口呈现。硬件状态变化必须能在设备运维平台和运营后台中被查看。")
    add_heading_black(doc, "6.3 技术限制", level=2)
    add_para(doc, "技术栈约束：前端使用Vue3、Element Plus、Vite、ECharts；后端使用Python 3.10+、FastAPI、SQLAlchemy；数据库使用SQLite。安全约束：密码不得明文存储，关键操作需记录日志，监管平台展示汇总数据时不得暴露用户敏感信息。业务约束：停车位状态变化必须可追踪，预约超时必须释放，价格调整必须可解释、可公示。项目不得为了降低实现难度删除AI视觉、GIS地图、支付、硬件、监管、共享、充电、营销等课题功能。")

    add_section_page_break(doc)
    add_heading_black(doc, "7 软件质量特性", level=1)
    add_para(doc, "质量特性要求覆盖用户可用性、开发可维护性、系统安全性和演示稳定性。每个质量属性均需要能映射到具体设计或测试点，避免只写抽象口号。")
    add_heading_black(doc, "7.1 可靠性", level=2)
    add_para(doc, "核心业务状态必须可靠保存，包括车位状态、预约锁定、订单支付、违停审核、设备告警和营销核销。模拟接口失败时需要给出明确提示，并保留最近一次可用状态或异常记录。")
    add_heading_black(doc, "7.2 易用性", level=2)
    add_para(doc, "车主端操作路径应简洁，管理端应支持筛选、排序、状态标识和批量处理，监管端应支持区域、时间、类型维度的统计查看。所有关键按钮和状态变化都需要可理解的中文提示。")
    add_table_black(
        doc,
        ["质量属性", "要求"],
        [
            ("可用性", "支持Chrome、Edge、Firefox主流浏览器；异常场景给出明确提示"),
            ("可维护性", "前端组件化、后端分层、接口文档可查、数据库字段与文档同步"),
            ("可扩展性", "18个一级模块通过标准API和数据表解耦，可替换模拟AI、支付、地图和设备接口"),
            ("安全性", "登录鉴权、角色权限、密码哈希、关键操作审计、监管数据脱敏"),
            ("易用性", "车主端适配移动端操作，管理端支持筛选、排序、图表和二次确认"),
        ],
    )

    add_section_page_break(doc)
    add_heading_black(doc, "8 需求分级", level=1)
    add_para(doc, "P0 必须有：停车场数据联网接入、城市停车一张图、智能导航与车位预约、车牌识别与无感支付、车场运营管理后台、车主移动端APP/小程序。")
    add_para(doc, "P1 重要：路内停车位视觉检测、停车诱导信息发布、停车行为大数据分析、动态定价策略引擎、共享停车管理平台、反向寻车导航、城市停车监管平台。")
    add_para(doc, "P2 最好有：充电车位智能管理、违停自动抓拍与取证、长期停车月卡管理、设备运维管理平台、商圈停车联合营销。")
    add_para(doc, "分级说明：P0模块构成车主停车和车场运营的最小闭环，必须优先完成；P1模块提升系统智能化、监管分析和运营效率，建议在P0稳定后并行推进；P2模块可采用模拟数据和重点页面完成演示，但仍需在需求、设计和测试中保留完整边界。")

    add_section_page_break(doc)
    add_heading_black(doc, "9 附录", level=1)
    add_para(doc, "附录用于保存需求追踪、评审和RA细化记录。附录内容不替代正文需求，只作为需求来源、评审过程和验收核查的补充依据。")
    add_heading_black(doc, "9.1 需求追踪矩阵", level=2)
    add_para(doc, "本节用于说明每个一级需求的来源、优先级、负责人、参与角色和对应测试用例。")
    add_table_black(
        doc,
        ["需求ID", "一级模块名称", "需求来源", "优先级", "负责人", "关联角色", "测试用例"],
        [(m["id"], m["name"], m["source"], m["priority"], m["owner"], m["actor"], f"TC-{m['id']}") for m in MODULES],
    )
    add_para(doc, "说明：P0为必须完成的核心闭环模块，P1为重要增强模块，P2为可模拟呈现的扩展模块。注册登录、停车场CRUD、订单管理、支付模拟、统计图表等实现项归入对应一级模块，不再误拆为额外一级模块。")
    add_heading_black(doc, "9.2 需求评审记录", level=2)
    add_para(doc, "评审日期：2026年6月2日。评审方式：小组会议。评审角色：PM（程晓洋）、RA（程子浩）、产品经理（丁梓钊）。")
    add_table_black(
        doc,
        ["编号", "类型", "发现", "严重程度", "处理方式"],
        [
            ("IR-01", "遗漏", "需求规格需回到项目要求18个一级模块口径", "高", "已按原始项目功能表统一为M1-M18"),
            ("IR-02", "遗漏", "AI视觉、GIS地图、支付和硬件模块缺少实训实现策略", "高", "明确采用模拟数据、状态面板和接口预留方式"),
            ("IR-03", "不明确", "实时刷新未定义具体间隔", "低", "明确车位状态默认30秒刷新，可配置"),
            ("IR-04", "不明确", "车位预约锁定时间需明确", "中", "明确为15分钟，超时自动释放"),
            ("IR-05", "矛盾", "计费单位小时/分钟表述不一致", "低", "统一为按小时计费，不足1小时按1小时"),
            ("IR-06", "遗漏", "缺少模块用例图", "高", "为18个一级模块逐一补充用例图"),
            ("IR-07", "遗漏", "缺少接口统一返回格式", "中", "定义为{code, msg, data}"),
            ("IR-08", "遗漏", "缺少安全和权限约束", "中", "补充登录、角色权限、关键操作日志和监管数据脱敏要求"),
        ],
    )
    add_para(doc, "评审结论：通过。18个一级功能模块已全部覆盖，并补充模块用例图、需求追踪矩阵、数据字典、非功能需求和实训阶段实现边界。")
    add_heading_black(doc, "9.3 RA细化记录", level=2)
    add_para(doc, "本节用于记录需求分析过程中的关键步骤和决策，避免需求来源不清。")
    add_table_black(
        doc,
        ["步骤", "时间", "执行人", "内容"],
        [
            ("1", "2026.06.01上午", "程子浩（RA）", "逐项阅读项目要求中的18个功能模块，识别模糊点"),
            ("2", "2026.06.01下午", "程子浩（RA）", "将每个一级模块拆解为输入、处理、输出、业务规则和验收点"),
            ("3", "2026.06.01晚上", "丁梓钊（PM）", "从产品和用户体验视角补充用例与异常流程"),
            ("4", "2026.06.02上午", "全员", "需求评审，确认18模块口径和模拟实现边界"),
            ("5", "2026.06.02下午", "程晓洋（PM）", "汇总评审意见，生成SRS V1.0正式版"),
        ],
    )
    add_para(doc, "关键决策：18个一级模块严格对应项目要求；实训无法完整落地的AI视觉、GIS地图、真实支付和硬件接入模块，以模拟数据、状态面板和接口预留方式实现；编码阶段优先完成车主停车闭环、运营管理闭环和监管分析闭环。")


def force_black_doc(doc):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            set_run_black(run, bold=run.bold)
    for table in doc.tables:
        set_table_black(table)


def replace_paragraph_text(paragraph, text, size=10.5, bold=None):
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = text
    set_run_black(run, size=size, bold=bold)


def set_template_cell(table, row, col, text, bold=False):
    cell = table.cell(row, col)
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(text))
    set_run_black(run, size=9.5, bold=bold)


def fill_template_table(table, rows, header_rows=1):
    for row_index in range(header_rows, len(table.rows)):
        for col in range(len(table.columns)):
            set_template_cell(table, row_index, col, "")
    for offset, row_values in enumerate(rows, start=header_rows):
        if offset >= len(table.rows):
            table.add_row()
        for col, value in enumerate(row_values):
            set_template_cell(table, offset, col, value)
    set_table_black(table)


def iter_body_blocks(doc):
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def remove_blocks_after_table_count(doc, keep_table_count):
    table_count = 0
    remove = False
    for child in list(doc.element.body.iterchildren()):
        if remove:
            child.getparent().remove(child)
            continue
        if child.tag == qn("w:tbl"):
            table_count += 1
            if table_count == keep_table_count:
                remove = True


def insert_paragraph_after(paragraph, text, style=None, ppr=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    if ppr is not None:
        new_p.append(deepcopy(ppr))
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_paragraph.style = style
    run = new_paragraph.add_run(text)
    set_run_black(run, size=10.5)
    return new_paragraph


def fill_static_toc(doc):
    toc_lines = [
        ("1 简介", "toc 1", "1"),
        ("1.1 目的", "toc 2", "1"),
        ("1.2 范围", "toc 2", "1"),
        ("2 总体概述", "toc 1", "2"),
        ("2.1 软件概述", "toc 2", "2"),
        ("2.1.1 项目介绍", "toc 3", "2"),
        ("2.1.2 产品环境介绍", "toc 3", "2"),
        ("2.2 软件功能", "toc 2", "3"),
        ("2.3 用户特征", "toc 2", "3"),
        ("2.4 假设和依赖关系", "toc 2", "4"),
        ("3 具体需求", "toc 1", "4"),
        ("3.1 系统用例", "toc 2", "4"),
        ("3.2 子功能模块需求", "toc 2", "5"),
        ("3.2.1 M1 路内停车位视觉检测", "toc 3", "5"),
        ("3.2.2 M2 停车场数据联网接入", "toc 3", "6"),
        ("3.2.3 M3 城市停车一张图", "toc 3", "7"),
        ("3.2.4 M4 停车诱导信息发布", "toc 3", "8"),
        ("3.2.5 M5 智能导航与车位预约", "toc 3", "9"),
        ("3.2.6 M6 车牌识别与无感支付", "toc 3", "10"),
        ("3.2.7 M7 停车行为大数据分析", "toc 3", "11"),
        ("3.2.8 M8 动态定价策略引擎", "toc 3", "12"),
        ("3.2.9 M9 共享停车管理平台", "toc 3", "13"),
        ("3.2.10 M10 充电车位智能管理", "toc 3", "14"),
        ("3.2.11 M11 违停自动抓拍与取证", "toc 3", "15"),
        ("3.2.12 M12 长期停车月卡管理", "toc 3", "16"),
        ("3.2.13 M13 车场运营管理后台", "toc 3", "17"),
        ("3.2.14 M14 车主移动端APP/小程序", "toc 3", "18"),
        ("3.2.15 M15 反向寻车导航", "toc 3", "19"),
        ("3.2.16 M16 设备运维管理平台", "toc 3", "20"),
        ("3.2.17 M17 城市停车监管平台", "toc 3", "21"),
        ("3.2.18 M18 商圈停车联合营销", "toc 3", "22"),
        ("3.3 数据字典", "toc 2", "23"),
        ("3.3.1 数据字典", "toc 3", "23"),
        ("3.3.2 E-R关系图", "toc 3", "25"),
        ("4 性能需求", "toc 1", "26"),
        ("4.1 时间性能需求", "toc 2", "26"),
        ("4.2 系统开放性需求", "toc 2", "26"),
        ("4.3 界面友好性需求", "toc 2", "27"),
        ("4.4 系统可用性需求", "toc 2", "27"),
        ("4.5 可管理性需求", "toc 2", "27"),
        ("5 接口需求", "toc 1", "28"),
        ("5.1 用户接口", "toc 2", "28"),
        ("5.2 软件接口", "toc 2", "28"),
        ("5.3 硬件接口", "toc 2", "29"),
        ("5.4 通讯接口", "toc 2", "29"),
        ("6 总体设计约束", "toc 1", "30"),
        ("6.1 标准符合性", "toc 2", "30"),
        ("6.2 硬件约束", "toc 2", "30"),
        ("6.3 技术限制", "toc 2", "31"),
        ("7 软件质量特性", "toc 1", "32"),
        ("7.1 可靠性", "toc 2", "32"),
        ("7.2 易用性", "toc 2", "32"),
        ("8 需求分级", "toc 1", "33"),
        ("9 附录", "toc 1", "34"),
        ("9.1 需求追踪矩阵", "toc 2", "34"),
        ("9.2 需求评审记录", "toc 2", "35"),
        ("9.3 RA细化记录", "toc 2", "35"),
    ]
    paragraphs = list(doc.paragraphs)
    marker = None
    toc_ppr_by_style = {}
    for index, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() == "目录":
            marker = paragraph
            start = index + 1
            break
    if marker is None:
        return
    end = start
    while end < len(paragraphs) and not paragraphs[end].text.strip().startswith("Keywords"):
        style_name = paragraphs[end].style.name
        if style_name.startswith("toc ") and paragraphs[end]._p.pPr is not None:
            toc_ppr_by_style.setdefault(style_name, deepcopy(paragraphs[end]._p.pPr))
        end += 1
    for paragraph in paragraphs[start:end]:
        paragraph._p.getparent().remove(paragraph._p)
    previous = marker
    for title, style_name, page in toc_lines:
        line = f"{title}\t{page}"
        previous = insert_paragraph_after(
            previous,
            line,
            style=style_name,
            ppr=toc_ppr_by_style.get(style_name),
        )
    insert_paragraph_after(previous, "", style="Normal")


def fill_template_front_matter(doc):
    configure_clean_doc(doc)
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if "XX系统需求规格说明书" in text:
            replace_paragraph_text(paragraph, f"{PROJECT_NAME}需求规格说明书", size=18, bold=True)
        elif text == "关键字":
            replace_paragraph_text(paragraph, "智慧停车；车位预约；城市停车一张图；无感支付；停车诱导；数据联网接入；停车监管；共享停车", size=10.5)
        elif text == "摘要信息":
            replace_paragraph_text(paragraph, f"本文档描述{PROJECT_NAME}的业务背景、18个一级功能模块、数据字典、接口约束、性能要求、质量属性、需求分级和需求评审结论。文档严格以项目要求.txt中的主体要求为依据，避免将注册登录、基础CRUD、订单管理等实现支撑项误列为独立一级模块。", size=10.5)

    fill_static_toc(doc)

    # Cover metadata table.
    set_template_cell(doc.tables[0], 1, 0, PROJECT_NAME)
    set_template_cell(doc.tables[0], 1, 1, "内部")
    set_template_cell(doc.tables[0], 1, 2, "仅供收件方查阅")
    set_template_cell(doc.tables[0], 3, 0, "SmartPark-19-2026")
    set_template_cell(doc.tables[0], 3, 1, "V1.0")
    set_template_cell(doc.tables[0], 3, 2, "SmartPark-SRS-001")

    # Author/review/approval table.
    set_template_cell(doc.tables[1], 0, 1, "程晓洋 / 项目经理")
    set_template_cell(doc.tables[1], 0, 3, DATE)
    set_template_cell(doc.tables[1], 1, 1, "程子浩 / 需求分析师")
    set_template_cell(doc.tables[1], 1, 3, DATE)
    set_template_cell(doc.tables[1], 2, 1, "程晓洋 / 组长")
    set_template_cell(doc.tables[1], 2, 3, DATE)

    fill_template_table(
        doc.tables[2],
        [
            ("2026.06.01", "V0.1", "N/A", "全部", "按项目要求18个一级功能模块建立需求规格框架", "程子浩（RA）"),
            ("2026.06.01", "V0.2", "N/A", "3.2", "补充每个模块的输入、处理、输出、业务规则和用例图", "丁梓钊（PM）"),
            ("2026.06.02", "V0.3", "IR-01~IR-08", "全部", "依据评审意见统一18模块口径并补充非功能需求", "全员评审"),
            ("2026.06.02", "V1.0", "N/A", "全部", "正式版发布，含需求追踪矩阵、评审记录和数据字典", "程晓洋（PM）"),
        ],
        header_rows=1,
    )
    fill_template_table(
        doc.tables[3],
        [
            ("SRS", "Software Requirement Specification", "软件需求规格说明书"),
            ("UC", "Use Case", "用例"),
            ("RTM", "Requirement Traceability Matrix", "需求追踪矩阵"),
            ("BR", "Business Rule", "业务规则"),
            ("RA", "Requirement Analyst", "需求分析师"),
        ],
        header_rows=1,
    )


def build_clean_srs():
    print("[生成] 基于需求分析模板副本填充正文...")
    if DIAGRAM_DIR.exists():
        shutil.rmtree(DIAGRAM_DIR)
    if not TEMPLATE.exists():
        raise FileNotFoundError(f"缺少模板副本: {TEMPLATE}")
    doc = Document(TEMPLATE)
    fill_template_front_matter(doc)
    remove_blocks_after_table_count(doc, 4)
    add_main_content(doc)
    force_black_doc(doc)
    return doc


def add_md_table(lines, headers, rows):
    lines.append("| " + " | ".join(headers) + " |")
    lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
    for row in rows:
        lines.append("| " + " | ".join(str(value) for value in row) + " |")
    lines.append("")


def write_markdown_srs():
    lines = [
        "# AI-driven城市智慧停车管理与诱导系统需求规格说明书",
        "",
        "> 文档编号：SmartPark-SRS-001  ",
        "> 小组：第19组  ",
        "> 编写人：程晓洋、丁梓钊、程子浩  ",
        f"> 日期：{DATE}  ",
        "> 说明：本文档按需求规格说明书模板栏目编写，以 `项目要求.txt` 中18个一级功能模块为主体要求；注册登录、基础CRUD、订单管理和统计图表作为支撑能力归入对应模块。",
        "",
        "## 文档修订记录",
        "",
    ]
    add_md_table(
        lines,
        ["日期", "版本", "变更编号", "修改章节", "修改内容", "修改人"],
        [
            ("2026.06.01", "V0.1", "N/A", "全部", "按项目要求18个一级功能模块建立需求规格框架", "程子浩（RA）"),
            ("2026.06.01", "V0.2", "N/A", "3.2", "补充每个模块的输入、处理、输出、业务规则和用例图", "丁梓钊（PM）"),
            ("2026.06.02", "V0.3", "IR-01~IR-08", "全部", "依据评审意见统一18模块口径并补充非功能需求", "全员评审"),
            ("2026.06.02", "V1.0", "N/A", "全部", "正式版发布，含需求追踪矩阵、评审记录和数据字典", "程晓洋（PM）"),
        ],
    )
    lines.extend([
        "## 关键字",
        "",
        "智慧停车；车位预约；城市停车一张图；无感支付；停车诱导；数据联网接入；停车监管；共享停车",
        "",
        "## 摘要",
        "",
        f"本文档描述{PROJECT_NAME}的业务背景、18个一级功能模块、数据字典、接口约束、性能要求、质量属性、需求分级和需求评审结论。文档严格以项目要求.txt中的主体要求为依据，避免将注册登录、基础CRUD、订单管理等实现支撑项误列为独立一级模块。",
        "",
        "## 术语与缩略语",
        "",
    ])
    add_md_table(
        lines,
        ["术语", "英文/缩写", "说明"],
        [
            ("SRS", "Software Requirement Specification", "软件需求规格说明书"),
            ("UC", "Use Case", "用例"),
            ("RTM", "Requirement Traceability Matrix", "需求追踪矩阵"),
            ("BR", "Business Rule", "业务规则"),
            ("RA", "Requirement Analyst", "需求分析师"),
        ],
    )

    lines.extend([
        "## 1 简介",
        "",
        "### 1.1 目的",
        "",
        f"本文档用于明确{PROJECT_NAME}在实训阶段需要实现和模拟呈现的需求范围，作为后续概要设计、详细设计、编码实现、测试用例编写和课程验收的统一依据。文档以项目要求.txt为最高优先级输入，所有一级功能模块均按原始18项课题要求展开。",
        "",
        "本文档同时给出需求来源、优先级、责任人、输入输出、业务规则、异常边界和验收标准，便于开发阶段按模块拆分任务，也便于测试阶段按需求追踪矩阵逐项核验。",
        "",
        "### 1.2 范围",
        "",
        "本文档覆盖18个一级功能模块的功能需求、模块用例图、输入处理输出、业务规则、异常处理、接口联动、优先级、责任人、需求追踪矩阵、评审记录、核心数据字典、接口需求、性能需求、设计约束和质量属性。",
        "",
        "实训阶段不包含真实AI模型训练、真实支付通道、真实GIS地图服务和真实硬件设备部署，但需要通过模拟数据、状态字段、接口预留和可视化界面完整表达业务闭环。",
        "",
        "## 2 总体概述",
        "",
        "### 2.1 软件概述",
        "",
        "### 2.1.1 项目介绍",
        "",
        "城市停车资源分布不均、路内外车位数据割裂、停车场信息不透明、车辆入出场效率低和监管统计滞后，是本项目要解决的核心问题。系统通过AI视觉检测、停车场数据接入、城市停车一张图、预约导航、无感支付、大数据分析和监管平台等能力，形成车主侧、运营侧和监管侧的停车业务闭环。",
        "",
        "从车主视角看，系统需要降低找车位、排队入场、缴费离场和反向寻车的时间成本；从停车场视角看，系统需要提升车位周转率、设备可见性、价格配置能力和营销联动能力；从监管视角看，系统需要支撑城市停车资源统筹、违停治理、收费透明和停车政策评估。",
        "",
        "### 2.1.2 产品环境介绍",
        "",
        "系统采用前后端分离B/S架构。前端建议使用Vue3、Element Plus、ECharts和Vite，后端使用FastAPI、SQLAlchemy和SQLite。实训阶段通过初始化数据、模拟接口和管理端录入方式呈现AI识别、地图、支付、设备接入等外部能力。",
        "",
        "### 2.2 软件功能",
        "",
        "软件功能严格覆盖项目要求中的18个一级模块。核心功能包括路内停车位视觉检测、停车场数据联网接入、城市停车一张图、停车诱导信息发布、智能导航与车位预约、车牌识别与无感支付、停车行为大数据分析、动态定价、共享停车、充电车位管理、违停取证、月卡管理、运营后台、车主移动端、反向寻车、设备运维、城市监管和商圈联合营销。",
        "",
        "### 2.3 用户特征",
        "",
    ])
    add_md_table(
        lines,
        ["用户类型", "主要目标", "使用特点"],
        [
            ("私家车主", "查询车位、预约、支付、寻车", "关注操作简单、信息准确、少等待"),
            ("停车场管理员", "管理车场、车位、订单、设备和收入", "关注数据可筛选、异常可处理、报表可导出"),
            ("车位主", "发布共享车位并查看收益", "关注时段冲突、预约状态和收益记录"),
            ("城市交管部门", "监管停车态势、违停和收费", "关注汇总数据、区域趋势和政策评估"),
            ("设备运维人员", "维护摄像头、道闸、诱导屏和充电桩", "关注在线状态、故障等级和处理记录"),
            ("商圈运营人员", "配置优惠并核销停车券", "关注活动规则、适用范围和效果统计"),
        ],
    )
    lines.extend([
        "### 2.4 假设和依赖关系",
        "",
        "假设停车场、车位、订单、设备、营销和统计数据可以通过模拟数据初始化；AI视觉识别、GIS地图、支付扣款和硬件动作均以模拟结果、状态字段或预留接口呈现。系统依赖本地Python、Node.js、SQLite数据库和现代浏览器。",
        "",
        "项目验收时以演示闭环为准：车主能完成查询、预约、入场、支付、寻车和活动使用；管理员能完成车场运营、价格配置、设备处理和营销管理；监管人员能查看停车态势、违停证据和统计分析。",
        "",
        "## 3 具体需求",
        "",
        "### 3.1 系统用例",
        "",
        "系统总体用例分为车主停车闭环、车场运营闭环和城市监管闭环。车主侧流程为搜索目的地、查看停车一张图、预约车位、导航前往、车牌识别入场、出场无感支付、查询停车记录和反向寻车。运营侧流程为接入停车场数据、监控车位设备、发布诱导信息、配置动态定价、处理订单和营销活动。监管侧流程为查看全市停车态势、分析供需与收费、审核违停证据和评估停车政策。",
        "",
        "总体用例图、系统功能结构图和各模块用例图已生成到 `_generated_use_case_diagrams/`，并嵌入正式Word交付物。",
        "",
        "### 3.2 子功能模块需求",
        "",
    ])
    for index, module in enumerate(MODULES, start=1):
        lines.extend([
            f"#### 3.2.{index} {module['id']} {module['name']}",
            "",
            f"**功能说明：** {module['description']}",
            "",
            f"**参与角色：** {module['actor']}",
            "",
            f"**主要用例：** {'、'.join(module['use_cases'])}",
            "",
            f"**输入：** {module['inputs']}",
            "",
            f"**处理：** {module['processing']}",
            "",
            f"**输出：** {module['outputs']}",
            "",
            f"**业务规则：** {module['rules']}",
            "",
            f"**接口与数据联动：** {INTEGRATION_NOTES[module['id']]}",
            "",
            f"**异常处理与边界：** {EXCEPTION_NOTES[module['id']]}",
            "",
            f"**验收标准：** 测试用例TC-{module['id']}至少覆盖正常流程、异常输入、权限拦截、状态流转、数据同步和页面提示；优先级为{module['priority']}，需求负责人为{module['owner']}。",
            "",
        ])

    lines.extend([
        "### 3.3 数据字典",
        "",
        "### 3.3.1 数据字典",
        "",
        "核心数据字典覆盖用户、车辆、停车场、车位、预约、订单、共享车位、违停证据、设备和营销活动等对象。后续详细设计和编码阶段应保持字段名称、含义和状态枚举一致。",
        "",
    ])
    add_md_table(
        lines,
        ["数据对象", "用途", "核心字段"],
        [
            ("users", "用户与角色权限", "id, username, password_hash, phone, role"),
            ("vehicles", "车辆与车牌信息", "id, user_id, plate_number, vehicle_type"),
            ("parking_lots", "停车场基础信息和接入状态", "id, name, address, region, total_spots, rate_per_hour"),
            ("parking_spots", "车位编号、类型和状态", "id, lot_id, spot_number, status, spot_type"),
            ("reservations", "预约锁定、确认、取消和过期", "id, user_id, spot_id, status, expire_at"),
            ("parking_orders", "入出场、计费和支付订单", "id, user_id, spot_id, plate_number, entry_time, amount, status"),
            ("shared_spots", "共享车位和共享时段", "id, owner_id, spot_id, available_start, available_end, price"),
            ("charging_piles", "充电桩状态和故障", "id, spot_id, status, fault_type, last_update"),
            ("violations", "违停证据和审核处理", "id, plate_number, location, violation_type, review_status"),
            ("monthly_cards", "月卡申请、审批和续费", "id, user_id, plate_number, card_type, valid_to, status"),
            ("devices", "设备台账、在线状态和运维日志", "id, lot_id, device_type, status, firmware_version"),
            ("guide_screens", "诱导屏位置和发布内容", "id, location, related_lot_id, content, publish_status"),
            ("pricing_rules", "动态定价规则和调价记录", "id, lot_id, period, factor, reason, enabled"),
            ("marketing_campaigns", "商圈营销活动和优惠规则", "id, name, merchant, rule_desc, valid_to, status"),
            ("analytics_snapshots", "统计指标快照", "id, lot_id, metric_type, metric_value, snapshot_at"),
        ],
    )
    lines.extend([
        "### 3.3.2 E-R关系图",
        "",
        "用户与车辆、预约、订单、月卡、共享车位存在一对多关系；停车场与车位、设备、诱导屏、订单存在一对多关系；预约确认后生成停车订单；订单完成后参与收入统计、行为分析和监管汇总；违停证据、设备告警和营销核销分别关联停车场、车位或订单。E-R关系图已嵌入Word交付物。",
        "",
        "## 4 性能需求",
        "",
        "### 4.1 时间性能需求",
        "",
    ])
    add_md_table(
        lines,
        ["场景", "指标要求", "验收方式"],
        [
            ("车主端查询停车场", "常规查询响应不超过500ms", "输入目的地后查看列表/地图加载时间"),
            ("预约车位", "预约状态写入不超过500ms，倒计时误差不超过1秒", "检查车位从空闲到锁定的状态变化"),
            ("运营后台统计", "复杂统计聚合不超过1秒", "按停车场和时间范围筛选统计报表"),
            ("状态刷新", "车位、设备和诱导屏状态默认30秒刷新，可配置", "检查页面展示时间和后端更新时间"),
            ("模拟支付", "出场计费和支付状态更新不超过1秒", "检查订单金额、支付结果和车位释放"),
        ],
    )
    lines.extend([
        "### 4.2 系统开放性需求",
        "",
        "系统应通过RESTful接口、统一数据模型和可替换模拟服务保持开放性。AI视觉、GIS地图、支付通道、设备接入、诱导屏发布等外部能力在实训阶段采用模拟接口，后续可用真实服务替换而不改变18个一级模块的业务边界。",
        "",
        "### 4.3 界面友好性需求",
        "",
        "车主端页面应突出查找车位、预约、支付、寻车和优惠使用；管理端页面应突出表格筛选、状态标识、批量操作、图表统计和异常处理；监管端页面应突出全市态势、区域对比、违停审核和趋势分析。关键操作需有明确按钮、状态反馈和错误提示。",
        "",
        "### 4.4 系统可用性需求",
        "",
        "系统在本地演示环境下应保证核心车主闭环、运营后台闭环和监管查看闭环可连续演示。外部能力不可用时，页面应以模拟数据或降级展示方式继续表达需求，不应出现空白页面或无法说明的中断流程。",
        "",
        "### 4.5 可管理性需求",
        "",
        "管理员应能查看车场、车位、订单、设备、价格、营销、违停和统计数据的状态；运维人员应能查看设备告警和处理记录；监管人员应能查看汇总数据。关键配置、审核、发布和处理动作需要记录操作者、时间、对象和结果。",
        "",
        "## 5 接口需求",
        "",
        "### 5.1 用户接口",
        "",
        "车主端提供搜索、预约、支付、寻车、记录查询和发票申请界面；管理端提供车场、车位、订单、设备、违停、营销和运营统计界面；监管端提供停车态势、供需分析、收费统计、投诉热力和政策评估界面。",
        "",
        "### 5.2 软件接口",
        "",
        "前后端采用RESTful JSON接口，统一响应格式为 `{code, msg, data}`。后端通过SQLAlchemy访问SQLite。AI识别、地图、支付、设备接入和诱导屏发布在实训阶段以模拟接口呈现，后续可替换为真实服务。",
        "",
        "### 5.3 硬件接口",
        "",
        "系统不直接连接真实硬件。摄像头、道闸、诱导屏、充电桩等设备通过设备表、运行日志和模拟状态面板管理。设备接口需保留设备编号、设备类型、运行状态、最后心跳、故障等级和处理记录。",
        "",
        "### 5.4 通讯接口",
        "",
        "开发环境前端运行于localhost:5173，后端运行于localhost:8000，通信协议为HTTP/1.1。接口返回统一结构 `{code, msg, data}`，其中code=0表示成功，code=400表示参数错误，code=401表示未登录或登录失效，code=403表示无权限，code=500表示服务端异常。未来可扩展WebSocket推送实时车位状态。",
        "",
        "## 6 总体设计约束",
        "",
        "### 6.1 标准符合性",
        "",
        "系统需求、设计和实现必须符合项目要求.txt中的18个一级模块，不得将注册登录、基础CRUD或订单管理误列为独立一级模块。接口命名、数据库字段和页面模块应保持一致，便于后续测试和文档交叉核验。",
        "",
        "### 6.2 硬件约束",
        "",
        "实训环境不部署真实摄像头、道闸、诱导屏和充电桩。相关功能通过模拟数据、设备状态表、运行日志和预留接口呈现。硬件状态变化必须能在设备运维平台和运营后台中被查看。",
        "",
        "### 6.3 技术限制",
        "",
        "技术栈约束：前端使用Vue3、Element Plus、Vite、ECharts；后端使用Python 3.10+、FastAPI、SQLAlchemy；数据库使用SQLite。安全约束：密码不得明文存储，关键操作需记录日志，监管平台展示汇总数据时不得暴露用户敏感信息。业务约束：停车位状态变化必须可追踪，预约超时必须释放，价格调整必须可解释、可公示。",
        "",
        "## 7 软件质量特性",
        "",
        "### 7.1 可靠性",
        "",
        "核心业务状态必须可靠保存，包括车位状态、预约锁定、订单支付、违停审核、设备告警和营销核销。模拟接口失败时需要给出明确提示，并保留最近一次可用状态或异常记录。",
        "",
        "### 7.2 易用性",
        "",
        "车主端操作路径应简洁，管理端应支持筛选、排序、状态标识和批量处理，监管端应支持区域、时间、类型维度的统计查看。所有关键按钮和状态变化都需要可理解的中文提示。",
        "",
    ])
    add_md_table(
        lines,
        ["质量属性", "具体要求"],
        [
            ("可维护性", "前端组件按业务域划分，后端按路由、服务、模型分层；接口和数据字典同步维护。"),
            ("安全性", "密码哈希存储，登录鉴权，角色权限控制，关键操作日志记录，监管数据脱敏展示。"),
            ("可扩展性", "AI识别、地图、支付和硬件设备均保留扩展接口，可由模拟服务替换为真实服务。"),
            ("可测试性", "每个一级模块至少对应一个测试用例编号，覆盖正常、异常、权限和状态流转。"),
        ],
    )
    lines.extend([
        "## 8 需求分级",
        "",
        "P0 必须有：停车场数据联网接入、城市停车一张图、智能导航与车位预约、车牌识别与无感支付、车场运营管理后台、车主移动端APP/小程序。",
        "",
        "P1 重要：路内停车位视觉检测、停车诱导信息发布、停车行为大数据分析、动态定价策略引擎、共享停车管理平台、反向寻车导航、城市停车监管平台。",
        "",
        "P2 最好有：充电车位智能管理、违停自动抓拍与取证、长期停车月卡管理、设备运维管理平台、商圈停车联合营销。",
        "",
        "分级说明：P0模块构成车主停车和车场运营的最小闭环，必须优先完成；P1模块提升系统智能化、监管分析和运营效率，建议在P0稳定后并行推进；P2模块可采用模拟数据和重点页面完成演示，但仍需在需求、设计和测试中保留完整边界。",
        "",
        "## 9 附录",
        "",
        "### 9.1 需求追踪矩阵",
        "",
    ])
    add_md_table(
        lines,
        ["模块ID", "模块名", "需求来源", "优先级", "负责人", "参与角色", "测试用例"],
        [(m["id"], m["name"], m["source"], m["priority"], m["owner"], m["actor"], f"TC-{m['id']}") for m in MODULES],
    )
    lines.extend([
        "说明：P0为必须完成的核心闭环模块，P1为重要增强模块，P2为可模拟呈现的扩展模块。注册登录、停车场CRUD、订单管理、支付模拟、统计图表等实现项归入对应一级模块，不再误拆为额外一级模块。",
        "",
        "### 9.2 需求评审记录",
        "",
    ])
    add_md_table(
        lines,
        ["问题编号", "问题类型", "问题描述", "严重程度", "处理结果"],
        [
            ("IR-01", "遗漏", "原需求容易将登录、订单、CRUD误拆为一级模块", "高", "统一回归项目要求18个一级功能模块"),
            ("IR-02", "遗漏", "AI视觉、GIS地图、支付和硬件模块缺少实训实现策略", "高", "明确采用模拟数据、状态面板和接口预留方式"),
            ("IR-03", "不明确", "实时刷新未定义具体间隔", "低", "明确车位状态默认30秒刷新，可配置"),
            ("IR-04", "不明确", "车位预约锁定时间需明确", "中", "明确为15分钟，超时自动释放"),
            ("IR-05", "矛盾", "计费单位小时/分钟表述不一致", "低", "统一为按小时计费，不足1小时按1小时"),
            ("IR-06", "遗漏", "缺少模块用例图", "高", "为18个一级模块逐一补充用例图"),
            ("IR-07", "遗漏", "缺少接口统一返回格式", "中", "定义为{code, msg, data}"),
            ("IR-08", "遗漏", "缺少安全和权限约束", "中", "补充登录、角色权限、关键操作日志和监管数据脱敏要求"),
        ],
    )
    lines.extend([
        "评审结论：通过。18个一级功能模块已全部覆盖，并补充模块用例图、需求追踪矩阵、数据字典、非功能需求和实训阶段实现边界。",
        "",
        "### 9.3 RA细化记录",
        "",
    ])
    add_md_table(
        lines,
        ["步骤", "时间", "执行人", "内容"],
        [
            ("1", "2026.06.01上午", "程子浩（RA）", "逐项阅读项目要求中的18个功能模块，识别模糊点"),
            ("2", "2026.06.01下午", "程子浩（RA）", "将每个一级模块拆解为输入、处理、输出、业务规则和验收点"),
            ("3", "2026.06.01晚上", "丁梓钊（PM）", "从产品和用户体验视角补充用例与异常流程"),
            ("4", "2026.06.02上午", "全员", "需求评审，确认18模块口径和模拟实现边界"),
            ("5", "2026.06.02下午", "程晓洋（PM）", "汇总评审意见，生成SRS V1.0正式版"),
        ],
    )
    lines.append("关键决策：18个一级模块严格对应项目要求；实训无法完整落地的AI视觉、GIS地图、真实支付和硬件接入模块，以模拟数据、状态面板和接口预留方式实现；编码阶段优先完成车主停车闭环、运营管理闭环和监管分析闭环。")
    lines.append("")
    MARKDOWN_OUTPUT.write_text("\n".join(lines), encoding="utf-8")
    print(f"[保存] {MARKDOWN_OUTPUT}")


def main():
    doc = build_clean_srs()
    print(f"[保存] {OUTPUT}")
    doc.save(OUTPUT)
    write_markdown_srs()
    print(f"[完成] 文件大小: {OUTPUT.stat().st_size:,} bytes")


if __name__ == "__main__":
    main()
