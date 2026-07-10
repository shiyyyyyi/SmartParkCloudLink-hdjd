from pathlib import Path
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

import build_final_paper_dzz as base


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path(r"D:\cxdownload\指导老师-学生姓名-毕设论文 2026届软件工程 (1).docx")
OUT_DIR = ROOT / "文档（项目相关的作业所需提交文档生成在这里）" / "期末大作业"
OUT_FILE = OUT_DIR / "指导老师-丁梓钊-毕设论文 2026届软件工程-模板版.docx"


def style_exists(doc, name):
    try:
        doc.styles[name]
        return True
    except KeyError:
        return False


def para(doc, text="", style=None, align=None, font="宋体", size=None, bold=False):
    if style and not style_exists(doc, style):
        style = None
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    base.set_east_asia_font(run, font)
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    return p


def page_break(doc):
    doc.add_page_break()


def set_table_cell(cell, text):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    base.set_east_asia_font(r, "宋体")
    r.font.size = Pt(12)


def add_cover(doc):
    # The cover already exists in the template. Replace text in-place so the
    # original school-template geometry, spacing, and table formatting survive.
    for idx, text in {
        0: "20260310华东交通大学",
        1: "毕业设计（论文）",
        8: "题目：基于 Vue3 与 FastAPI 的智慧停车管理系统设计与实现",
    }.items():
        if idx < len(doc.paragraphs):
            set_paragraph_text(doc.paragraphs[idx], text)
            doc.paragraphs[idx].alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.tables[0]
    data = [
        ["学    院:", "信息与软件工程学院", "信息与软件工程学院", "信息与软件工程学院"],
        ["专    业:", "软件工程", "班    级:", ""],
        ["学生姓名:", "丁梓钊", "学    号:", ""],
        ["指导教师:", "________", "完成日期:", "2026年7月3日"],
    ]
    for row, values in zip(table.rows, data):
        for cell, text in zip(row.cells, values):
            set_table_cell(cell, text)


def add_integrity_statement(doc):
    table = doc.tables[1]
    cell = table.cell(0, 0)
    cell.text = ""
    lines = [
        "",
        "毕业设计（论文）诚信声明",
        "",
        "本人郑重声明：所呈交的毕业设计（论文）是我个人在指导教师指导下，结合工程实训项目资料、本地源码和课程文档独立整理完成的成果。",
        "文中引用的项目需求、系统设计、测试报告及相关技术资料均已在参考文献或正文说明中体现。除文中特别加以标注和致谢的内容外，本文不包含其他人已发表或撰写的研究成果。",
        "",
        "",
        "本人签名：丁梓钊            指导教师签名：",
        "",
        "2026 年 7 月 3 日",
    ]
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (1, 7, 9) else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(line)
        base.set_east_asia_font(r, "宋体")
        r.font.size = Pt(12)
        if i == 1:
            r.bold = True
            r.font.size = Pt(16)


def set_paragraph_text(paragraph, text, font="宋体", size=None):
    paragraph.text = ""
    run = paragraph.add_run(text)
    base.set_east_asia_font(run, font)
    if size:
        run.font.size = Pt(size)
    return run


def trim_after_second_template_table(doc):
    body = doc._element.body
    table_count = 0
    keep_through = None
    for i, child in enumerate(list(body)):
        if child.tag == qn("w:tbl"):
            table_count += 1
            if table_count == 2:
                keep_through = i
                break
    if keep_through is None:
        return
    children = list(body)
    for child in children[keep_through + 1:]:
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def add_team_page(doc):
    base.h(doc, 1, "小组成员及个人工作说明")
    base.p(doc, "本项目为工程实训课程小组项目，主题为“AI-driven 城市智慧停车管理与诱导系统”。按照期末大作业要求，本文在毕业设计论文模板基础上整理项目背景、技术路线、需求分析、系统设计、数据库设计、详细实现、测试结果与总结展望，并在文档开头说明小组成员和本人完成的主要工作。", style="Body Text")
    base.add_table(
        doc,
        ["成员", "角色", "负责模块", "主要工作"],
        [
            ["程晓洋", "组长/后端主程/数据库管理员", "M2、M7、M8、M13、M16、M17", "负责数据接入、统计分析、动态定价、运营后台、设备运维与监管平台。"],
            ["程子浩", "需求分析师/后端主程/QA", "M3、M4、M5、M6、M15、M18", "负责停车一张图、诱导发布、预约导航、无感支付、反向寻车和联合营销。"],
            ["丁梓钊", "产品经理/前端主程/测试工程师", "M1、M9、M10、M11、M12、M14", "参与需求细化与评审，负责车主端产品流、前端路由与组件设计，复核共享停车、充电车位、违停取证、月卡和车主端场景，并参与测试验证。"],
        ],
        [2.4, 3.0, 3.3, 7.0],
    )
    page_break(doc)


def add_abstracts_and_toc(doc):
    para(doc, "基于 Vue3 与 FastAPI 的智慧停车管理系统设计与实现", style="Title", align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "摘  要", style="摘要标题", align=WD_ALIGN_PARAGRAPH.CENTER)
    for text in [
        "随着城市机动车保有量持续增长，停车资源分布不均、空余车位信息不透明、人工收费效率低、停车场运营数据滞后等问题日益突出。为缓解车主找位难、停车场管理效率低和城市监管缺少统一视图等矛盾，本文围绕工程实训项目“AI-driven 城市智慧停车管理与诱导系统”，设计并实现了一个基于 Vue3 与 FastAPI 的智慧停车管理系统。",
        "系统采用前后端分离的 B/S 架构。前端使用 Vue3、Element Plus、Vite、Pinia、Axios 和 ECharts 实现车主端与管理端页面；后端使用 FastAPI、SQLAlchemy 和 SQLite 构建 RESTful API、数据模型、预约状态机、订单计费、统计分析和管理接口。系统重点实现停车场数据接入、城市停车查询、车位预约、车牌入出场模拟、无感支付模拟、运营后台、停车记录和反向寻车等核心闭环，并对共享停车、充电车位、违停取证、长期月卡等扩展模块进行了原型化设计和接口预留。",
        "测试结果表明，系统 V1.0 MVP 完成 38 条功能测试用例，认证、停车场查询、预约、订单支付、管理后台和数据接入等模块均能按预期运行。本文最后总结了项目实现过程中的经验，并从真实地图服务、AI视觉识别、硬件接入、并发控制和移动端体验等方面提出后续改进方向。",
    ]:
        para(doc, text, style="摘要正文")
    para(doc, "关键词：智慧停车；Vue3；FastAPI；前后端分离；车位预约；无感支付", style="关键词正文")
    page_break(doc)

    para(doc, "Design and Implementation of a Smart Parking Management System Based on Vue3 and FastAPI", style="Title", align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Abstract", style="英文摘要标题", align=WD_ALIGN_PARAGRAPH.CENTER)
    for text in [
        "With the continuous growth of urban vehicles, parking systems face problems such as uneven resource distribution, opaque vacancy information, inefficient manual charging, and delayed operational statistics. To address these issues, this paper presents the design and implementation of a smart parking management system based on Vue3 and FastAPI, derived from the engineering training project named AI-driven Urban Smart Parking Management and Guidance System.",
        "The system adopts a front-end and back-end separated B/S architecture. The front end is developed with Vue3, Element Plus, Vite, Pinia, Axios and ECharts, while the back end is built with FastAPI, SQLAlchemy and SQLite. The implemented MVP covers parking lot data access, parking search, parking space reservation, simulated license plate entry and exit, simulated non-stop payment, operation dashboard, parking records and reverse car searching. Extended modules such as shared parking, charging parking spaces, violation evidence collection and monthly card management are also designed for further development.",
        "Functional testing shows that all 38 test cases of the V1.0 MVP passed successfully. The system can support the main closed-loop process from searching a parking lot to reservation, entry, payment, exit and operation management. Finally, this paper summarizes the development experience and proposes future improvements in real map services, AI visual recognition, hardware integration, concurrency control and mobile user experience.",
    ]:
        para(doc, text, style="英文摘要正文", font="Times New Roman")
    para(doc, "Keywords: Smart Parking; Vue3; FastAPI; Front-end and Back-end Separation; Reservation; Non-stop Payment", style="英文关键词正文", font="Times New Roman")
    page_break(doc)

    para(doc, "目  录", style="目录标题", align=WD_ALIGN_PARAGRAPH.CENTER)
    toc = [
        ("toc 1", "第1章 绪论\t1"),
        ("toc 2", "1.1 研究的背景及意义\t1"),
        ("toc 3", "1.1.1 选题的背景\t1"),
        ("toc 3", "1.1.2 国内外研究现状\t2"),
        ("toc 3", "1.1.3 研究的意义\t3"),
        ("toc 2", "1.2 系统目标\t4"),
        ("toc 1", "第2章 相关技术\t5"),
        ("toc 2", "2.1 总体技术架构概述\t5"),
        ("toc 2", "2.2 关键技术分节详解\t6"),
        ("toc 2", "2.3 开发工具与环境\t8"),
        ("toc 1", "第3章 需求分析与概要设计\t9"),
        ("toc 1", "第4章 数据库设计\t15"),
        ("toc 1", "第5章 详细设计\t19"),
        ("toc 1", "第6章 系统测试\t27"),
        ("toc 1", "第7章 总结与展望\t31"),
        ("toc 1", "致 谢\t33"),
        ("toc 1", "参考文献\t34"),
        ("toc 1", "附录A 软件使用说明书\t35"),
        ("toc 1", "附录B 主要源代码\t36"),
    ]
    for style, text in toc:
        para(doc, text, style=style)
    page_break(doc)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(str(TEMPLATE))
    add_cover(doc)
    add_integrity_statement(doc)
    trim_after_second_template_table(doc)
    add_team_page(doc)
    add_abstracts_and_toc(doc)
    base.chapter_1(doc)
    base.chapter_2(doc)
    base.chapter_3(doc)
    base.chapter_4(doc)
    base.chapter_5(doc)
    base.chapter_6(doc)
    base.chapter_7(doc)
    base.appendices(doc)

    doc.core_properties.author = "丁梓钊"
    doc.core_properties.title = "基于 Vue3 与 FastAPI 的智慧停车管理系统设计与实现"
    doc.save(OUT_FILE)
    print(OUT_FILE)


if __name__ == "__main__":
    sys.exit(main())
