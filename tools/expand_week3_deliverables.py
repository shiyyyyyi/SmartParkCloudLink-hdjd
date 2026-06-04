# -*- coding: utf-8 -*-
"""Expand week-3 SRS and HLD deliverables in-place without changing templates.

The edits deliberately keep the existing template sections and table layouts.
They only add detail to existing paragraphs/tables in the generated documents.
"""
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOC_ROOT = ROOT / "文档（项目相关的作业所需提交文档生成在这里）"
SRS = DOC_ROOT / "03_需求分析" / "19组-程晓洋-需求分析.docx"
HLD = DOC_ROOT / "04_系统设计" / "19组-程晓洋-系统设计.docx"

MODULE_DETAIL = {
    "M1": "补充说明：该模块在实训阶段重点保证识别结果可录入、可复核、可追溯，页面应能区分自动识别结果与人工修正结果，避免低置信度数据直接进入监管或计费流程。",
    "M2": "补充说明：该模块需要保留接入配置、同步日志和异常告警，便于后续把模拟数据源替换为真实停车场接口时仍沿用相同字段和状态流转。",
    "M3": "补充说明：查询结果应同时服务车主端推荐、诱导屏发布和监管汇总，空位、价格、距离、更新时间等字段必须保持统一口径。",
    "M4": "补充说明：诱导内容发布前应进行空位阈值和设备状态校验，发布失败时保留上一条有效内容并记录失败原因，方便运营人员回溯。",
    "M5": "补充说明：预约状态至少包含待确认、已确认、已取消、已超时等状态，车位锁定和释放必须与订单、车位状态保持一致，避免重复预约。",
    "M6": "补充说明：无感支付在实训阶段采用模拟扣款，但仍需表现入场、出场、计费、支付成功/失败和车位释放的完整闭环。",
    "M7": "补充说明：分析指标应明确统计时间范围和数据口径，避免把异常订单、取消预约和人工处理记录混入正常经营指标。",
    "M8": "补充说明：动态定价规则需要保留启停状态、适用时段、价格系数和调价原因，使价格变化可解释、可回滚、可审计。",
    "M9": "补充说明：共享车位应处理时段冲突、车位主取消、预约用户爽约等边界，收益统计以完成订单为准。",
    "M10": "补充说明：充电车位既是车位资源又关联充电设备，预约、占用、故障和维护状态需要在车主端与管理端同步展示。",
    "M11": "补充说明：违停证据应包含车牌、位置、时间、图片/模拟证据和复核状态，未复核通过前不得进入处罚或监管确认状态。",
    "M12": "补充说明：月卡管理需要校验车牌、有效期、适用车场和续费状态，过期或不匹配时按临停规则处理。",
    "M13": "补充说明：运营后台是多个模块的数据汇聚入口，应提供筛选、导出、异常处理、二次确认和操作日志能力。",
    "M14": "补充说明：车主端需要保证移动端操作路径短、状态反馈明确，关键操作如预约、取消、支付应有明确确认和结果提示。",
    "M15": "补充说明：反向寻车应结合进行中订单、车位区域、楼层和入口信息生成可读路线，定位信息不足时提供人工问询提示。",
    "M16": "补充说明：设备运维模块需要覆盖设备台账、在线状态、故障等级、处理人和处理记录，避免设备异常影响核心停车流程而不可见。",
    "M17": "补充说明：监管平台展示汇总和脱敏数据，重点支持区域供需、违停治理、收费趋势和政策效果分析。",
    "M18": "补充说明：联合营销需要把优惠规则、适用商户、核销状态和订单关联起来，保证活动效果能够被统计和复盘。",
}


SRS_SECTION_EXPANSIONS = {
    "本文档用于明确AI-driven城市智慧停车管理与诱导系统在实训阶段需要实现和模拟呈现的需求范围": "扩展说明：本文档不仅描述功能边界，也用于约束后续概要设计、数据库设计、接口设计和测试用例编写。所有功能条目均需要能够映射到设计模块、页面入口、接口或数据表，避免出现需求已写但后续无法验证的断点。",
    "本文档覆盖18个一级功能模块的功能需求": "扩展说明：覆盖范围按时间线从需求采集、需求细化、需求评审到正式版归档展开。文档保留需求来源、优先级、责任人和验收标准，便于小组在编码阶段按模块拆分任务，也便于教师验收时快速核对交付物是否完整。",
    "城市停车资源分布不均": "扩展说明：系统分析时将车主、停车场管理员、交管部门、设备运维人员和商圈运营人员作为核心干系人，分别关注找位效率、运营效率、监管透明度、设备可用性和营销转化效果。需求细化时不把这些关注点拆散为零散功能，而是归入18个一级模块的业务闭环。",
    "系统采用前后端分离B/S架构": "扩展说明：该产品环境要求前端、后端和数据库都能在本地演示环境中运行。AI识别、GIS地图、支付和硬件设备均通过模拟接口呈现，但接口字段、状态码和数据流向需要按真实接入思路设计，保证后续扩展时不推翻需求定义。",
    "系统总体用例分为车主停车闭环": "扩展说明：总体用例强调跨模块协同，而不是单点页面功能。车主侧以查询、预约、入场、支付、寻车为主线；运营侧以车场接入、车位维护、订单处理、价格配置、设备处理为主线；监管侧以汇总态势、违停审核、供需分析和政策评估为主线。",
    "车主端提供搜索、预约、支付、寻车": "扩展说明：用户接口需要给出明确的状态提示，例如车位紧张、预约锁定中、预约超时、支付失败、设备离线和数据延迟。管理端和监管端的表格、筛选、统计图也应围绕这些状态设计，便于测试人员构造验收用例。",
    "系统需求、设计和实现必须符合项目要求.txt中的18个一级模块": "扩展说明：注册登录、基础CRUD、订单管理、支付模拟和统计图表均作为支撑能力存在，不作为额外一级模块编号。该约束用于保证需求文档、概要设计、编码任务和测试追踪矩阵之间保持一致。",
    "核心业务状态必须可靠保存": "扩展说明：可靠性重点体现在状态可恢复、操作可追踪和异常可解释。预约、订单、支付、违停、设备告警和营销核销等状态变化应记录变更时间和处理结果，避免演示过程中出现无法解释的数据跳变。",
}


HLD_SECTION_EXPANSIONS = {
    "本文档描述AI-driven城市智慧停车管理与诱导系统的概要设计": "扩展说明：本文档按照第三周任务从需求规格说明书向系统设计过渡，重点回答系统边界如何划分、模块如何分解、接口如何协同、数据如何落表、页面如何承载业务流程等问题。后续详细设计和编码实现应以本文档的模块、接口和数据结构为基础。",
    "系统覆盖SRS中的18个一级业务模块": "扩展说明：P0模块优先保证演示闭环稳定，P1模块补足运营、监管和分析能力，P2模块通过模拟数据和预留接口体现扩展方向。这样的分层可以让小组在有限实训周期内先完成主流程，再逐步增加增强能力。",
    "软件应用于城市停车资源查询": "扩展说明：应用场景不仅包括车主找车位，也包括停车场管理员维护资源、监管人员观察城市态势、设备运维人员处理异常和商圈运营人员核销优惠。界面和接口设计应让这些角色各自看到与职责匹配的数据。",
    "系统外部实体包括车主": "扩展说明：上下文边界内的系统负责业务状态保存、接口编排和页面展示；边界外的AI识别、GIS地图、支付、道闸、诱导屏和充电桩均以模拟适配器接入。这样既符合实训环境约束，也保留未来真实设备接入的扩展位置。",
    "系统采用前后端分离B/S结构": "扩展说明：表现层不直接访问数据库，所有业务操作均通过后端接口完成；服务层负责状态机、计费规则、统计聚合和异常处理；数据访问层统一管理事务，降低前端页面与数据库结构之间的耦合。",
    "核心业务流程以车主停车闭环为主线": "扩展说明：管理端和监管端流程都围绕该闭环产生的数据展开。预约会影响车位状态，入出场会生成订单，支付会释放车位，订单和违停记录会进入统计分析，设备异常会影响检测、诱导和出入场可用性。",
    "系统模块分解严格对应需求规格说明书中的18个一级业务模块": "扩展说明：每个一级模块都需要在概要设计中明确页面入口、服务职责、数据表和接口边界。实现支撑项可以跨模块复用，但不得破坏一级模块与需求追踪矩阵的对应关系。",
    "接口统一采用HTTP/1.1 + RESTful JSON": "扩展说明：接口设计应保持请求参数、返回结构和错误码一致，前端通过统一拦截器处理登录失效、权限不足、参数错误和服务异常。管理类接口需要校验角色，监管类接口需要执行数据脱敏。",
    "核心概念模型围绕用户、车辆、停车场": "扩展说明：数据库设计优先保证核心状态一致性，尤其是车位、预约、订单、支付和设备状态之间的联动。统计类数据可以通过快照表沉淀，减少页面每次打开都进行复杂聚合的成本。",
    "基础数据包括示例停车场": "扩展说明：基础数据应覆盖至少一个完整演示闭环：车主账号、车辆、停车场、空闲车位、预约记录、入场订单、支付记录、设备告警、违停证据和营销活动，确保答辩时不依赖临时手工录入。",
    "日志策略：关键操作记录操作者": "扩展说明：出错处理不仅面向用户提示，也服务于测试和复盘。每类异常都应能定位到接口、数据对象和处理结果；对于支付失败、设备离线、预约超时等场景，系统应保留可重试或人工处理入口。",
}


def set_run_style(run, template_run=None):
    if template_run is not None:
        run.font.name = template_run.font.name
        if template_run._element.rPr is not None and template_run._element.rPr.rFonts is not None:
            east_asia = template_run._element.rPr.rFonts.get(qn("w:eastAsia"))
            if east_asia:
                run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
        run.font.size = template_run.font.size or Pt(10.5)
        run.bold = template_run.bold
    else:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.5)


def append_to_paragraph(paragraph, addition):
    if addition in paragraph.text:
        return False
    template = paragraph.runs[-1] if paragraph.runs else None
    run = paragraph.add_run(addition)
    set_run_style(run, template)
    return True


def insert_after(doc, paragraph, text):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = paragraph._parent.paragraphs[-1]
    # python-docx does not expose the newly inserted paragraph object directly.
    # Resolve it by walking document paragraphs and matching XML identity.
    for candidate in doc.paragraphs:
        if candidate._p is new_p:
            inserted = candidate
            break
    inserted.style = paragraph.style
    inserted.paragraph_format.line_spacing = paragraph.paragraph_format.line_spacing
    inserted.paragraph_format.space_after = paragraph.paragraph_format.space_after
    run = inserted.add_run(text)
    set_run_style(run, paragraph.runs[-1] if paragraph.runs else None)
    return inserted


def add_detail_after_matching_paragraphs(doc, expansions):
    changed = 0
    for para in list(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        for needle, addition in expansions.items():
            if needle in text:
                # Keep section shape: add detail inside the same section, after the original paragraph.
                if addition not in "\n".join(p.text for p in doc.paragraphs):
                    insert_after(doc, para, addition)
                    changed += 1
                break
    return changed


def expand_srs_modules(doc):
    changed = 0
    paragraphs = list(doc.paragraphs)
    for idx, para in enumerate(paragraphs):
        text = para.text
        if "功能说明：" not in text:
            continue
        module_id = None
        # Look backwards for nearest heading containing Mx.
        for prev in reversed(paragraphs[max(0, idx - 6):idx]):
            for mid in MODULE_DETAIL:
                if mid in prev.text:
                    module_id = mid
                    break
            if module_id:
                break
        if module_id and MODULE_DETAIL[module_id] not in text:
            append_to_paragraph(para, MODULE_DETAIL[module_id])
            changed += 1
    return changed


def expand_hld_tables(doc):
    changed = 0
    for table in doc.tables:
        if not table.rows:
            continue
        headers = [cell.text for cell in table.rows[0].cells]
        if "功能摘要" in "".join(headers):
            for row in table.rows[1:]:
                cells = row.cells
                if len(cells) >= 4 and "设计补充：" not in cells[3].text:
                    mid = cells[0].text.strip()
                    cells[3].text = cells[3].text.strip() + f"\n设计补充：该模块需明确页面入口、服务方法、状态变化和异常提示，保证能从需求追踪到接口与数据表。"
                    changed += 1
        if "错误场景" in "".join(headers):
            for row in table.rows[1:]:
                cells = row.cells
                if len(cells) >= 2 and "记录处理结果" not in cells[1].text:
                    cells[1].text = cells[1].text.strip() + "；同时记录处理结果，便于测试复盘。"
                    changed += 1
    return changed


def count_chars(doc):
    return sum(len(p.text) for p in doc.paragraphs) + sum(len(c.text) for t in doc.tables for r in t.rows for c in r.cells)


def main():
    srs_doc = Document(str(SRS))
    hld_doc = Document(str(HLD))

    before_srs = count_chars(srs_doc)
    before_hld = count_chars(hld_doc)

    srs_changes = add_detail_after_matching_paragraphs(srs_doc, SRS_SECTION_EXPANSIONS)
    srs_changes += expand_srs_modules(srs_doc)

    hld_changes = add_detail_after_matching_paragraphs(hld_doc, HLD_SECTION_EXPANSIONS)
    hld_changes += expand_hld_tables(hld_doc)

    srs_doc.save(str(SRS))
    hld_doc.save(str(HLD))

    print(f"SRS changes={srs_changes}, chars {before_srs}->{count_chars(srs_doc)}")
    print(f"HLD changes={hld_changes}, chars {before_hld}->{count_chars(hld_doc)}")


if __name__ == "__main__":
    main()
