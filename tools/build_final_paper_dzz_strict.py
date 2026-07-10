from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path(r"D:\cxdownload\指导老师-学生姓名-毕设论文 2026届软件工程 (1).docx")
OUT_DIR = ROOT / "文档（项目相关的作业所需提交文档生成在这里）" / "期末大作业"
OUT_FILE = OUT_DIR / "指导老师-第19组-毕设论文 2026届软件工程-符合要求版.docx"

HLD_IMAGES = ROOT / "文档（项目相关的作业所需提交文档生成在这里）" / "04_系统设计" / "_generated_hld_diagrams"
USE_CASE_IMAGES = ROOT / "文档（项目相关的作业所需提交文档生成在这里）" / "03_需求分析" / "_generated_use_case_diagrams"
PROJECT = ROOT / "smart-park-v1.0"
THESIS_TITLE_CN = "AI-driven 城市智慧停车管理与诱导系统的设计与实现"
THESIS_TITLE_EN = "Design and Implementation of an AI-driven Urban Smart Parking Management and Guidance System"
RUNNING_HEADER_ODD = "华东交通大学毕业设计（论文）"
GROUP_NAME = "第19组"
GROUP_MEMBERS = "程晓洋、程子浩、丁梓钊"
RUNNING_HEADER_EVEN = f"{GROUP_NAME}：{THESIS_TITLE_CN}"


def style_name(doc: Document, preferred: str, fallback: str = "Normal") -> str:
    try:
        doc.styles[preferred]
        return preferred
    except KeyError:
        return fallback


def set_east_asia_font(run, east_asia: str = "宋体", west: str = "Times New Roman") -> None:
    run.font.name = west
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), west)
    rfonts.set(qn("w:hAnsi"), west)


def set_paragraph_text(paragraph, text: str, east_asia: str = "宋体", size: float | None = None, bold: bool = False):
    paragraph.text = ""
    run = paragraph.add_run(text)
    set_east_asia_font(run, east_asia)
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    return run


def normalize_styles(doc: Document) -> None:
    updates = {
        "Normal": ("宋体", 10.5, False, None),
        "Body Text": ("宋体", 10.5, False, None),
        "Heading 1": ("黑体", 15, True, "000000"),
        "Heading 2": ("黑体", 13, True, "000000"),
        "Heading 3": ("黑体", 12, True, "000000"),
        "Title": ("黑体", 16, True, "000000"),
        "摘要正文": ("宋体", 10.5, False, None),
        "英文摘要正文": ("Times New Roman", 10.5, False, None),
    }
    for name, (font, size, bold, color) in updates.items():
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        style.font.name = font
        style.font.size = Pt(size)
        style.font.bold = bold
        if color:
            style.font.color.rgb = RGBColor.from_string(color)
        rpr = style._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), font if font != "Times New Roman" else "宋体")
        rfonts.set(qn("w:ascii"), "Times New Roman")
        rfonts.set(qn("w:hAnsi"), "Times New Roman")
        pf = style.paragraph_format
        if name in ("Normal", "Body Text", "摘要正文"):
            pf.first_line_indent = Cm(0.74)
            pf.line_spacing = 1.5
            pf.space_after = Pt(0)
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def trim_after_second_template_table(doc: Document) -> None:
    body = doc._element.body
    table_count = 0
    keep_through = None
    for idx, child in enumerate(list(body)):
        if child.tag == qn("w:tbl"):
            table_count += 1
            if table_count == 2:
                keep_through = idx
                break
    if keep_through is None:
        raise RuntimeError("模板中未找到封面信息表和诚信声明表。")
    for child in list(body)[keep_through + 1 :]:
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def page_break(doc: Document) -> None:
    doc.add_page_break()


def p(
    doc: Document,
    text: str = "",
    style: str = "Body Text",
    align=None,
    bold: bool = False,
    size: float | None = None,
    east_asia: str = "宋体",
):
    para = doc.add_paragraph(style=style_name(doc, style, "Normal"))
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    set_east_asia_font(run, east_asia)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    return para


def h(doc: Document, level: int, text: str):
    para = doc.add_heading(text, level=level)
    for run in para.runs:
        set_east_asia_font(run, "黑体")
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.bold = True
    if level == 1:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return para


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, fill: str | None = None, center: bool = False) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(str(text))
    set_east_asia_font(run, "宋体")
    run.font.size = Pt(9)
    run.bold = bold
    if fill:
        shade_cell(cell, fill)


def apply_table_borders(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "666666")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_cm: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    apply_table_borders(table)
    for i, head in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, head, bold=True, fill="D9EAF7", center=True)
        cell.width = Cm(widths_cm[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, center=len(str(value)) <= 10)
            cells[i].width = Cm(widths_cm[i])
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        try:
            para = doc.add_paragraph(style="List Bullet")
        except KeyError:
            para = doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.74)
        para.paragraph_format.first_line_indent = None
        run = para.add_run(item)
        set_east_asia_font(run, "宋体")
        run.font.size = Pt(10.5)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        try:
            para = doc.add_paragraph(style="List Number")
        except KeyError:
            para = doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.74)
        para.paragraph_format.first_line_indent = None
        run = para.add_run(item)
        set_east_asia_font(run, "宋体")
        run.font.size = Pt(10.5)


def caption(doc: Document, text: str) -> None:
    para = p(doc, text, style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    para.paragraph_format.first_line_indent = None
    para.paragraph_format.space_after = Pt(6)


def add_image(doc: Document, path: Path, title: str, width: float = 5.45) -> None:
    if not path.exists():
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    pic = run.add_picture(str(path), width=Inches(width))
    pic._inline.docPr.set("title", title)
    pic._inline.docPr.set("descr", title)
    caption(doc, title)


def add_code_block(doc: Document, file_path: Path, title: str, start: int, end: int) -> None:
    h(doc, 2, title)
    lines = file_path.read_text(encoding="utf-8", errors="ignore").splitlines()
    snippet = "\n".join(f"{i + 1:03d}  {line}" for i, line in enumerate(lines[start - 1 : end], start - 1))
    table = doc.add_table(rows=1, cols=1)
    apply_table_borders(table)
    cell = table.cell(0, 0)
    shade_cell(cell, "F3F6FA")
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.first_line_indent = None
    para.paragraph_format.line_spacing = 1.0
    run = para.add_run(snippet)
    run.font.name = "Consolas"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Consolas")
    rfonts.set(qn("w:hAnsi"), "Consolas")
    run.font.size = Pt(7.5)
    doc.add_paragraph()


def update_cover(doc: Document) -> None:
    replacements = {
        0: "20260310华东交通大学",
        1: "毕业设计（论文）",
        8: f"题目：{THESIS_TITLE_CN}",
    }
    for idx, text in replacements.items():
        if idx < len(doc.paragraphs):
            set_paragraph_text(doc.paragraphs[idx], text, "宋体")
            doc.paragraphs[idx].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.tables[0]
    data = [
        ["学    院:", "信息与软件工程学院", "信息与软件工程学院", "信息与软件工程学院"],
        ["专    业:", "软件工程", "班    级:", ""],
        ["小组成员:", GROUP_MEMBERS, "小组编号:", "第19组"],
        ["指导教师:", "________", "完成日期:", "2026年7月3日"],
    ]
    for row, values in zip(table.rows, data):
        for cell, value in zip(row.cells, values):
            set_cell_text(cell, value, center=True)


def update_integrity_statement(doc: Document) -> None:
    cell = doc.tables[1].cell(0, 0)
    cell.text = ""
    lines = [
        "",
        "毕业设计（论文）诚信声明",
        "",
        "本组郑重声明：所呈交的毕业设计（论文）是第19组成员在指导教师指导下，结合工程实训项目资料、本地源码、需求分析报告、系统设计报告和测试材料共同整理完成的研究与实践成果。",
        "文中引用的项目文档、开源框架文档及相关技术资料均已在参考文献或正文中说明。除文中特别注明和致谢的内容外，本文不包含其他个人或集体已经发表或撰写的成果。",
        "",
        "小组成员签名：程晓洋、程子浩、丁梓钊            指导教师签名：",
        "",
        "2026 年 7 月 3 日",
    ]
    for i, line in enumerate(lines):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (1, 6, 8) else WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(line)
        set_east_asia_font(run, "宋体")
        run.font.size = Pt(12 if i != 1 else 16)
        run.bold = i == 1


def set_header_text(header, text: str) -> None:
    if header.paragraphs:
        para = header.paragraphs[0]
        para.text = ""
    else:
        para = header.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    set_east_asia_font(run, "宋体")
    run.font.size = Pt(9)


def update_headers(doc: Document) -> None:
    for idx, section in enumerate(doc.sections):
        if idx == 0:
            set_header_text(section.header, "")
            set_header_text(section.first_page_header, "")
            set_header_text(section.even_page_header, "")
            continue
        set_header_text(section.header, RUNNING_HEADER_ODD)
        set_header_text(section.first_page_header, "")
        set_header_text(section.even_page_header, RUNNING_HEADER_EVEN)


def add_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    if footer.paragraphs:
        para = footer.paragraphs[0]
        para.text = ""
    else:
        para = footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("第19组  2026届软件工程毕业设计（论文）")
    set_east_asia_font(run, "宋体")
    run.font.size = Pt(9)


def front_matter(doc: Document) -> None:
    h(doc, 1, "小组成员及分工说明")
    p(doc, "本项目为工程实训课程第19组项目，主题为“AI-driven 城市智慧停车管理与诱导系统”。本文根据课程期末大作业要求，按照附件毕业设计（论文）格式，对项目背景、需求分析、总体设计、数据库设计、详细实现、测试结果和总结展望进行系统整理。")
    p(doc, "项目采用小组协作完成。本文由第19组成员基于本地工程文件、项目计划、需求分析、系统设计、测试报告和成员周报共同整理撰写，重点说明系统整体方案、各成员承担的模块、核心实现过程和测试验证结果。项目后端未采用 SpringBoot，而是使用 FastAPI + SQLAlchemy + SQLite；前端使用 Vue3 + Element Plus + Vite，文档中已按实际完成工作进行说明。")
    add_table(
        doc,
        ["成员", "项目开发工作", "文档撰写/整理工作", "完成说明"],
        [
            ["程晓洋", "负责 M2、M7、M8、M13、M16、M17，完成项目统筹、数据库设计、数据接入、运营后台、设备运维和监管分析相关工作。", "参与项目立项、开发计划、系统设计和最终文档资料汇总，提供数据库、后台管理和统计分析部分材料。", "组长/PM/数据库管理员/后端主程"],
            ["程子浩", "负责 M3、M4、M5、M6、M15、M18，完成停车一张图、诱导发布、预约导航、无感支付、反向寻车、联合营销和接口联调。", "参与需求分析、概要设计、接口说明和测试用例整理，提供车主停车路径和后端核心流程部分材料。", "需求分析师/后端主程/QA"],
            ["丁梓钊", "负责 M1、M9、M10、M11、M12、M14，完成车主端页面流、共享停车、充电车位、违停取证、月卡管理、前端交互和测试验证。", "参与需求评审、页面原型、测试报告和本文最终格式整理，重点整理车主端、扩展模块和测试验证内容。", "产品经理/前端主程/测试工程师"],
        ],
        [1.8, 5.5, 5.5, 2.7],
    )
    page_break(doc)

    p(doc, THESIS_TITLE_CN, style="Title", align=WD_ALIGN_PARAGRAPH.CENTER, east_asia="黑体")
    p(doc, "摘  要", style="摘要标题", align=WD_ALIGN_PARAGRAPH.CENTER, east_asia="黑体")
    for text in [
        "随着城市机动车保有量持续增长，停车资源分布不均、空余车位信息不透明、人工收费效率低、停车场运营数据滞后等问题日益突出。智慧停车系统通过数据接入、车位预约、车牌识别、支付模拟、运营统计和监管分析，将车主、停车场运营方与城市监管方纳入统一平台，是智慧交通和智慧城市建设中的重要应用场景。",
        "本文围绕工程实训项目“AI-driven 城市智慧停车管理与诱导系统”，设计并实现了一个基于 Vue3 与 FastAPI 的智慧停车管理系统。系统采用前后端分离 B/S 架构，前端使用 Vue3、Element Plus、Vite、Pinia、Axios 和 ECharts 实现车主端与管理端页面；后端使用 FastAPI、SQLAlchemy 和 SQLite 构建 RESTful API、数据模型、预约状态机、订单计费、统计分析和管理接口。",
        "系统 V1.0 MVP 聚焦停车场数据联网接入、城市停车一张图、智能导航与车位预约、车牌识别与无感支付、车场运营管理后台和车主移动端六个核心模块，同时对路内视觉检测、共享停车、充电车位、违停取证、长期月卡、反向寻车和监管分析等扩展能力进行需求设计与原型化说明。测试结果表明，系统完成 38 条功能测试用例，核心流程能够稳定支撑搜索、预约、入场、出场、支付、记录查询和后台管理。最后，本文从真实地图接入、AI视觉识别、硬件设备接入、并发控制和移动端体验等方面提出后续改进方向。",
    ]:
        p(doc, text, style="摘要正文")
    p(doc, "关键词：智慧停车；Vue3；FastAPI；前后端分离；车位预约；无感支付", style="关键词正文", bold=True)
    page_break(doc)

    p(doc, THESIS_TITLE_EN, style="Title", align=WD_ALIGN_PARAGRAPH.CENTER, east_asia="黑体")
    p(doc, "Abstract", style="英文摘要标题", align=WD_ALIGN_PARAGRAPH.CENTER, east_asia="Times New Roman")
    for text in [
        "With the continuous growth of urban vehicles, parking systems face practical problems such as uneven resource distribution, opaque vacancy information, inefficient manual charging and delayed operational statistics. Smart parking systems integrate data access, reservation, license plate recognition, payment simulation, operation statistics and management services, and therefore become an important application scenario of smart transportation.",
        "Based on the engineering training project named AI-driven Urban Smart Parking Management and Guidance System, this paper designs and implements a smart parking management system using Vue3 and FastAPI. The front end is developed with Vue3, Element Plus, Vite, Pinia, Axios and ECharts, while the back end is implemented with FastAPI, SQLAlchemy and SQLite. The system provides RESTful APIs, database models, reservation state transitions, order billing, analytics and administration functions.",
        "The V1.0 MVP focuses on six core modules: parking lot data access, parking search, intelligent reservation, simulated license plate entry and exit, simulated non-stop payment, operation management and the vehicle owner portal. Functional testing shows that all 38 test cases pass successfully, supporting the end-to-end process from search and reservation to entry, payment, records and administration. Future improvements include real map services, AI visual recognition, hardware integration, concurrency control and mobile optimization.",
    ]:
        p(doc, text, style="英文摘要正文", east_asia="Times New Roman")
    p(doc, "Keywords: Smart Parking; Vue3; FastAPI; Front-end and Back-end Separation; Reservation; Non-stop Payment", style="英文关键词正文", east_asia="Times New Roman", bold=True)
    page_break(doc)

    p(doc, "目  录", style="目录标题", align=WD_ALIGN_PARAGRAPH.CENTER, east_asia="黑体")
    toc_lines = [
        ("toc 1", "第一章 绪  论\t9"),
        ("toc 2", "1.1 研究的背景及意义\t9"),
        ("toc 2", "1.2 国内外研究与应用现状\t9"),
        ("toc 2", "1.3 论文主要工作\t10"),
        ("toc 1", "第二章 相关技术\t10"),
        ("toc 2", "2.1 前端开发技术\t11"),
        ("toc 2", "2.2 后端开发技术\t11"),
        ("toc 2", "2.3 数据库与接口技术\t12"),
        ("toc 2", "2.4 开发工具与运行环境\t12"),
        ("toc 1", "第三章 需求分析与概要设计\t13"),
        ("toc 2", "3.1 系统角色与业务边界\t13"),
        ("toc 2", "3.2 功能性需求\t13"),
        ("toc 2", "3.3 非功能性需求\t15"),
        ("toc 2", "3.4 系统总体架构与业务流程\t15"),
        ("toc 2", "3.5 小组分工与需求落点\t17"),
        ("toc 1", "第四章 数据库设计\t18"),
        ("toc 2", "4.1 概念结构设计\t18"),
        ("toc 2", "4.2 逻辑结构设计\t19"),
        ("toc 2", "4.3 物理结构与字段约束\t19"),
        ("toc 2", "4.4 数据一致性设计\t20"),
        ("toc 1", "第五章 系统详细设计与实现\t20"),
        ("toc 2", "5.1 前端总体设计\t20"),
        ("toc 2", "5.2 后端接口设计\t25"),
        ("toc 2", "5.3 关键业务实现\t26"),
        ("toc 3", "5.3.1 停车场查询与推荐\t26"),
        ("toc 3", "5.3.2 车位预约状态机\t27"),
        ("toc 3", "5.3.3 车牌入出场与计费支付\t27"),
        ("toc 3", "5.3.4 管理后台与统计分析\t27"),
        ("toc 2", "5.4 小组负责模块的详细设计\t31"),
        ("toc 2", "5.5 部署与运行设计\t33"),
        ("toc 1", "第六章 系统测试\t34"),
        ("toc 2", "6.1 测试环境与方法\t34"),
        ("toc 2", "6.2 测试范围与用例统计\t34"),
        ("toc 2", "6.3 典型测试用例\t35"),
        ("toc 2", "6.4 测试结果与分析\t35"),
        ("toc 1", "第七章 总结与展望\t36"),
        ("toc 2", "7.1 设计工作总结\t36"),
        ("toc 2", "7.2 不足与未来展望\t36"),
        ("toc 1", "致  谢\t37"),
        ("toc 1", "参考文献\t37"),
        ("toc 1", "附录A 软件使用说明书\t37"),
        ("toc 1", "附录B 主要源代码\t40"),
    ]
    for style, text in toc_lines:
        para = p(doc, text, style=style)
        para.paragraph_format.first_line_indent = None
    page_break(doc)


def chapter_1(doc: Document) -> None:
    h(doc, 1, "第一章 绪  论")
    h(doc, 2, "1.1 研究的背景及意义")
    paragraphs = [
        "城市停车问题是城市交通治理中长期存在的基础性问题。随着机动车保有量持续增长，城市核心区、商业综合体、医院、学校、交通枢纽等区域经常出现停车供需不平衡的情况。车主在出行前无法准确掌握目的地周边停车场的空余车位、价格、距离和出入口位置，往往需要绕行寻找车位，既增加了出行时间，也会形成无效交通流，进一步加重道路拥堵。",
        "传统停车场系统多以单个停车场为单位进行收费、入出场和车位管理，信息只在场内流转，难以对外提供实时余位和统一查询服务。对于车主而言，停车服务链条被切分为查找停车场、到场排队、人工或扫码缴费、离场、查找车辆等多个相互割裂的环节。对于停车场运营方而言，车位周转率、订单收入、设备运行状态、异常订单和用户投诉等数据分散在不同台账中，人工统计效率低，难以及时支撑运营决策。",
        "智慧停车管理系统的建设目标，是通过统一数据接入、停车场查询、车位预约、车牌识别、无感支付、运营分析和监管统计，把停车资源、用户行为和管理流程连接起来。该类系统不仅可以帮助车主减少找位时间，也可以帮助停车场提高车位利用率、提升收费效率、降低人工管理成本，并为城市停车政策评估和精细化治理提供数据基础。",
        "本项目来源于软件工程实训课程。课程要求以小组为单位完成项目立项、计划、需求分析、系统设计、编码实现、测试验收和文档整理。本组选题为“AI-driven 城市智慧停车管理与诱导系统”，在课程周期内以本地可运行的 Web 应用形式实现 MVP。由于真实地图、AI视觉模型、支付通道、摄像头、道闸和充电桩等硬件条件有限，项目采用模拟数据、状态面板和接口预留的方式实现教学场景下的完整业务闭环。",
    ]
    for text in paragraphs:
        p(doc, text)
    h(doc, 2, "1.2 国内外研究与应用现状")
    for text in [
        "从实际应用来看，国内外已有较多停车平台支持线上缴费、余位展示、车牌识别和会员优惠。大型商业区、机场、高铁站和城市道路停车系统通常会部署诱导屏、车牌识别设备和收费系统，并通过小程序或 App 提供停车缴费入口。这些系统证明了停车服务线上化和自动化的必要性，也说明停车管理正在从单点收费系统转向数据驱动的综合服务平台。",
        "但现有系统仍然存在若干局限。第一，部分系统只服务单个停车场或单一运营方，缺少跨停车场的统一检索和推荐能力；第二，车位预约、反向寻车、共享停车、充电车位、违停取证和城市监管等功能之间经常独立建设，难以形成完整闭环；第三，很多系统对设备状态、订单异常和运营指标的管理能力不足，不能很好地支撑运营人员进行快速决策。",
        "从技术角度看，前端单页应用、RESTful API、关系型数据库、数据可视化和状态机模型已经能够支撑中小规模智慧停车平台。真实工业系统通常还需要接入地图服务、物联网设备、支付网关、目标检测模型和消息推送服务。本项目在课程约束下，将这些外部能力抽象为可替换接口或模拟数据，既保留智慧停车业务的核心逻辑，又保证系统可以在本地独立运行和演示。",
        "与传统 CRUD 管理系统相比，智慧停车系统更强调状态一致性和流程闭环。预约会改变车位状态，入场会生成订单并占用车位，出场会触发计费，支付完成后需要释放车位并写入记录。任何一个环节处理不当，都可能导致余位不准确、重复预约、订单金额错误或后台统计失真。因此，本文在系统设计中重点讨论预约状态机、订单状态机、车位状态同步和后台统计口径。",
    ]:
        p(doc, text)
    h(doc, 2, "1.3 论文主要工作")
    add_numbered(
        doc,
        [
            "梳理智慧停车业务需求，明确车主、管理员和监管视角下的核心场景，并将项目要求映射为18个一级功能模块。",
            "设计前后端分离的总体架构，使用 Vue3 构建车主端和管理端页面，使用 FastAPI 提供认证、停车场、预约、订单、后台管理和数据分析接口。",
            "完成核心数据库建模，围绕用户、停车场、车位、预约、订单和同步日志构建可运行的数据闭环，并对扩展模块保留设计空间。",
            "实现 V1.0 MVP 的停车查询、车位预约、模拟入出场、计费支付、停车记录、反向寻车、运营后台和数据同步功能。",
            "结合小组分工，说明数据运营、车主停车路径、车主端体验、扩展模块原型和功能测试等工作内容。",
        ],
    )


def chapter_2(doc: Document) -> None:
    h(doc, 1, "第二章 相关技术")
    h(doc, 2, "2.1 前端开发技术")
    for text in [
        "系统前端采用 Vue3 作为核心框架。Vue3 的组件化机制适合把停车场卡片、搜索框、预约弹窗、订单列表、统计卡片和管理表格拆分为独立组件，有利于降低页面复杂度。Composition API 使状态、请求和交互逻辑更容易按照业务功能组织，便于在车主端和管理端之间复用。",
        "Vite 作为前端构建工具，具有启动快、热更新快和配置简洁的特点，适合课程实训中高频调试的开发方式。Element Plus 提供表单、按钮、表格、菜单、弹窗、消息提示、标签和分页等通用组件，帮助项目在较短时间内形成较完整的用户界面。对后台管理系统而言，Element Plus 的表格和表单组件能够显著提升开发效率。",
        "Pinia 用于保存登录用户、角色、token、车牌等全局状态。Vue Router 用于组织登录、注册、车主首页、停车场详情、我的预约、我的订单、停车记录、反向寻车、个人中心和管理后台等页面。Axios 统一封装 HTTP 请求，在请求拦截器中自动携带认证信息，在响应拦截器中统一处理异常提示，避免不同页面重复编写错误处理代码。",
        "ECharts 用于管理后台的数据可视化。停车场运营管理需要展示订单量、收入趋势、车位占用率、饱和度和停车场分布等指标，图表比纯文本更适合呈现趋势和对比关系。本项目在后台仪表盘中使用图表展示运营数据，为后续监管分析和运营优化提供基础。",
    ]:
        p(doc, text)
    h(doc, 2, "2.2 后端开发技术")
    for text in [
        "后端采用 FastAPI 框架。FastAPI 基于 Python 类型标注和 Pydantic 数据模型，可以自动生成 Swagger 接口文档，便于前端开发人员和测试人员查看请求参数、返回结构和接口路径。在课程项目中，自动化接口文档能减少沟通成本，也方便演示系统能力。",
        "系统通过 APIRouter 将接口拆分为认证、停车场、预约、订单、后台管理和数据分析等模块。模块化路由结构使各业务边界更加清晰，便于多人协作和后续扩展。认证模块负责注册、登录、个人信息和角色识别；停车场模块负责列表查询、详情、车位和数据同步；预约模块负责创建、确认、取消和超时释放；订单模块负责入出场、计费、支付、记录和反向寻车。",
        "SQLAlchemy 用于对象关系映射。通过模型类定义 users、parking_lots、parking_spots、reservations、orders 和 sync_logs 等表，后端业务代码可以通过对象操作数据库记录，而不需要在每个接口中手写复杂 SQL。Pydantic schema 用于区分请求模型、响应模型和数据库模型，降低数据校验和接口返回格式出错的概率。",
        "SQLite 作为本地数据库，部署简单、无需单独安装数据库服务，适合课程演示和小规模数据测试。虽然真实生产环境通常会选择 MySQL、PostgreSQL 等数据库，但 SQLite 能够满足本项目本地运行、状态持久化和功能测试的需求。",
    ]:
        p(doc, text)
    h(doc, 2, "2.3 数据库与接口技术")
    add_table(
        doc,
        ["技术", "项目用途", "采用原因"],
        [
            ["Vue3", "构建车主端和管理端单页应用", "组件化能力强，生态成熟，适合快速构建交互界面"],
            ["Element Plus", "表单、表格、弹窗、菜单和消息提示", "减少基础 UI 开发工作量，界面风格统一"],
            ["Pinia", "用户状态、角色、token 和车牌缓存", "API 简洁，适合 Vue3 项目状态管理"],
            ["FastAPI", "RESTful API 和 Swagger 文档", "类型清晰，接口文档自动生成，利于联调"],
            ["SQLAlchemy", "数据库模型和持久化操作", "ORM 方式降低 SQL 编写复杂度"],
            ["SQLite", "本地数据库存储", "无需服务端安装，适合实训演示"],
            ["ECharts", "后台统计图表", "支持收入趋势、占用率和饱和度可视化"],
        ],
        [2.4, 5.5, 6.2],
    )
    h(doc, 2, "2.4 开发工具与运行环境")
    add_table(
        doc,
        ["类别", "工具/环境", "说明"],
        [
            ["操作系统", "Windows", "本地开发、调试和演示环境"],
            ["后端运行", "Python 3.10+", "运行 FastAPI、SQLAlchemy 和数据初始化脚本"],
            ["前端运行", "Node 18+ / npm", "安装依赖并启动 Vite 开发服务器"],
            ["浏览器", "Chrome / Edge", "访问前端页面和 Swagger 文档"],
            ["代码管理", "Git / GitHub", "保存项目源码和阶段性成果"],
            ["编辑器", "VS Code / PyCharm", "前后端代码编辑和调试"],
        ],
        [2.8, 4.3, 7.0],
    )


def chapter_3(doc: Document) -> None:
    h(doc, 1, "第三章 需求分析与概要设计")
    h(doc, 2, "3.1 系统角色与业务边界")
    for text in [
        "系统主要面向三类角色：车主用户、停车场管理员和城市监管人员。车主用户关注附近停车场、空余车位、价格、预约、入出场、支付、停车记录和反向寻车；管理员关注停车场、车位、订单、收入、设备和同步日志；监管人员关注区域停车供需、收费统计、违停处理和政策评估。",
        "课程阶段的 V1.0 MVP 重点实现车主和管理员两类角色。监管人员视角在需求和设计中体现，部分统计能力通过后台数据分析接口预留。外部 AI 视觉识别、地图导航、支付通道、摄像头、道闸、诱导屏和充电桩等能力均采用模拟数据或预留接口实现，确保系统能够在本地独立运行。",
    ]:
        p(doc, text)
    add_table(
        doc,
        ["角色", "核心目标", "主要功能"],
        [
            ["车主用户", "快速找到合适车位并完成停车闭环", "搜索停车场、查看详情、预约车位、查看订单、支付、停车记录、反向寻车"],
            ["管理员", "维护停车资源并掌握运营状态", "车场管理、车位管理、订单管理、统计看板、数据同步、异常释放"],
            ["监管人员", "掌握城市停车运行情况", "供需分析、收费统计、违停处理、政策评估和脱敏汇总"],
        ],
        [2.2, 5.0, 7.7],
    )
    h(doc, 2, "3.2 功能性需求")
    modules = [
        ["M1", "路内停车位视觉检测", "模拟检测点位、车位状态和人工复核记录，由丁梓钊负责需求与页面原型思路"],
        ["M2", "停车场数据联网接入", "停车场余位和同步日志，支持手动触发模拟同步"],
        ["M3", "城市停车一张图", "停车场列表、搜索、筛选、排序、详情和附近推荐"],
        ["M4", "停车诱导信息发布", "诱导屏台账、内容生成和发布记录，课程阶段预留设计"],
        ["M5", "智能导航与车位预约", "预约创建、确认、取消、15分钟锁定和重复预约校验"],
        ["M6", "车牌识别与无感支付", "模拟入场、出场、计费、支付和订单状态变化"],
        ["M7", "停车行为大数据分析", "收入趋势、订单量、周转率和饱和度统计"],
        ["M8", "动态定价策略引擎", "价格规则和峰谷时段策略，课程阶段以规则说明体现"],
        ["M9", "共享停车管理平台", "共享时段发布、预约冲突校验和收益展示，由丁梓钊负责原型设计"],
        ["M10", "充电车位智能管理", "充电桩状态、预约和异常占用告警，由丁梓钊负责原型设计"],
        ["M11", "违停自动抓拍与取证", "违停证据、审核状态和处理记录，由丁梓钊负责原型设计"],
        ["M12", "长期停车月卡管理", "月卡申请、审批、续费和到期提醒，由丁梓钊负责原型设计"],
        ["M13", "车场运营管理后台", "仪表盘、车场管理、订单管理和收入统计"],
        ["M14", "车主移动端APP/小程序", "车主端搜索、预约、支付、记录和个人中心，由丁梓钊重点负责"],
        ["M15", "反向寻车导航", "根据车牌查询进行中停车位置和步行提示"],
        ["M16", "设备运维管理平台", "摄像头、道闸、诱导屏、充电桩等设备台账"],
        ["M17", "城市停车监管平台", "区域供需、收费、违停和政策评估汇总"],
        ["M18", "商圈停车联合营销", "优惠活动、停车券核销和营销效果统计"],
    ]
    add_table(doc, ["编号", "一级功能模块", "需求说明"], modules, [1.5, 4.4, 9.0])
    h(doc, 2, "3.3 非功能性需求")
    add_bullets(
        doc,
        [
            "可用性：车主端操作路径应围绕搜索、预约、支付和记录组织，减少无关入口；管理端应突出列表、筛选、统计和异常处理。",
            "性能：本地演示环境下常规查询应快速响应，统计接口应避免阻塞核心预约和订单接口。",
            "安全性：登录、预约、订单和管理后台接口需要校验用户身份，管理接口应限制管理员角色访问。",
            "一致性：预约、入场、支付、取消和过期释放都必须同步更新车位状态和停车场余位。",
            "可维护性：前端按 views、router、stores、utils 分层，后端按 routers、models、schemas、database 分层。",
            "可扩展性：真实地图、AI识别、支付网关和硬件设备通过预留字段、模拟接口和状态面板逐步替换。",
        ],
    )
    h(doc, 2, "3.4 系统总体架构与业务流程")
    p(doc, "系统采用浏览器端 SPA 与 FastAPI 服务端分离的 B/S 架构。前端通过 Axios 访问 RESTful API，后端路由完成业务校验和状态转换，SQLAlchemy 负责数据库访问，SQLite 保存本地数据。管理后台通过统计接口获得停车场、订单、收入和占用率等汇总指标。")
    add_image(doc, HLD_IMAGES / "02_architecture.png", "图 3-1 系统总体架构图", 5.7)
    p(doc, "核心业务流程从车主搜索停车场开始，经过查看详情、预约车位、到场确认、模拟入场、模拟出场、计费支付，再进入停车记录和后台运营统计。预约状态和车位状态是流程控制的核心，二者必须保持一致。")
    add_image(doc, HLD_IMAGES / "03_activity.png", "图 3-2 核心业务活动图", 5.6)
    add_image(doc, HLD_IMAGES / "04_sequence.png", "图 3-3 预约与支付核心时序图", 5.6)
    h(doc, 2, "3.5 小组分工与需求落点")
    for text in [
        "小组按照数据运营、车主停车路径和车主端扩展场景三个责任域划分工作。程晓洋负责数据接入、统计分析、动态定价、运营后台、设备运维和监管平台；程子浩负责城市停车一张图、诱导发布、预约导航、无感支付、反向寻车和联合营销；丁梓钊负责路内检测、共享停车、充电车位、违停取证、月卡管理和车主移动端。",
        "设计阶段，小组共同确认前后端接口、页面路径、状态提示、异常反馈和测试验收口径。测试阶段，围绕核心停车闭环和各成员负责模块设计测试用例，重点关注预约冲突、支付状态、充电桩不可预约、共享车位时段冲突、违停证据字段完整、月卡有效期等场景。",
    ]:
        p(doc, text)
    add_image(doc, USE_CASE_IMAGES / "M14_use_case.png", "图 3-4 M14车主移动端用例图", 5.3)
    add_image(doc, USE_CASE_IMAGES / "M9_use_case.png", "图 3-5 M9共享停车管理用例图", 5.3)


def chapter_4(doc: Document) -> None:
    h(doc, 1, "第四章 数据库设计")
    h(doc, 2, "4.1 概念结构设计")
    for text in [
        "系统核心数据围绕用户、停车场、车位、预约、订单和同步日志展开。用户可以拥有车辆和车牌信息，车主发起预约时需要绑定停车场、车位和车牌；预约确认后可以进入停车订单流程；订单记录入场时间、出场时间、停车时长、费用和支付状态；停车场和车位保存总量、余位、位置、价格和状态；同步日志记录停车场数据接入行为。",
        "数据库设计的关键在于状态一致性。车位状态会受到预约、入场、取消、过期、出场和支付等操作影响；停车场余位需要与车位状态保持一致；订单状态需要与支付状态、停车记录和后台统计保持一致。因此，数据库表之间不仅存在外键关系，也存在业务状态转换关系。",
    ]:
        p(doc, text)
    add_image(doc, HLD_IMAGES / "05_er.png", "图 4-1 系统核心ER图", 5.8)
    h(doc, 2, "4.2 逻辑结构设计")
    add_table(
        doc,
        ["数据表", "主要字段", "说明"],
        [
            ["users", "id, username, password_hash, phone, plate_numbers, role, created_at", "保存用户账号、联系方式、车牌集合、角色和密码哈希"],
            ["parking_lots", "id, name, address, lat, lng, total_spots, available_spots, price_per_hour, status", "保存停车场基础信息、坐标、总车位、余位、价格和接入状态"],
            ["parking_spots", "id, lot_id, spot_number, status", "保存停车场下具体车位编号以及 free、reserved、occupied 等状态"],
            ["reservations", "id, user_id, lot_id, spot_id, plate_number, status, expires_at, confirmed_at", "保存预约记录、车牌、状态、过期时间和确认时间"],
            ["orders", "id, user_id, lot_id, spot_id, plate_number, entry_time, exit_time, amount, status", "保存停车订单、入出场时间、费用和支付状态"],
            ["sync_logs", "id, lot_id, status, message, created_at", "保存停车场数据同步记录和结果说明"],
        ],
        [2.6, 6.6, 5.7],
    )
    h(doc, 2, "4.3 物理结构与字段约束")
    add_table(
        doc,
        ["对象", "状态/约束", "转换说明"],
        [
            ["车位", "free / reserved / occupied", "预约时由 free 变为 reserved；入场或确认后变为 occupied；取消、过期或支付后释放"],
            ["预约", "created / confirmed / cancelled / expired", "创建后 15 分钟内有效；确认、取消或过期后不可重复使用"],
            ["订单", "parking / pending_pay / paid / exception", "入场生成 parking；出场后进入 pending_pay；支付后变为 paid"],
            ["用户", "user / admin", "普通用户访问车主端，管理员访问运营后台和统计接口"],
            ["同步日志", "success / failed", "记录停车场数据接入结果，便于后台追踪"],
        ],
        [2.4, 4.8, 7.6],
    )
    h(doc, 2, "4.4 数据一致性设计")
    for text in [
        "预约创建接口需要检查用户是否已有未完成预约，并从目标停车场中选择可用车位。如果没有可用车位，应直接返回失败提示；如果预约成功，应同时写入预约记录、修改车位状态、减少停车场余位。取消预约和预约过期时，需要恢复车位状态并增加停车场余位。",
        "模拟入场接口需要根据车牌和停车场创建订单，绑定车位并修改状态。模拟出场接口根据入场时间和停车场费率计算费用，将订单置为待支付。支付接口在校验订单归属和状态后，将订单置为已支付，并释放对应车位。后台统计接口只读取已完成或进行中的订单，不直接修改业务数据。",
        "为了方便课程演示，当前项目使用 SQLite 保存数据。后续如迁移到 MySQL 或 PostgreSQL，应进一步补充事务隔离、唯一索引、乐观锁或行级锁，避免并发预约同一车位时发生状态冲突。",
    ]:
        p(doc, text)


def chapter_5(doc: Document) -> None:
    h(doc, 1, "第五章 系统详细设计与实现")
    h(doc, 2, "5.1 前端总体设计")
    for text in [
        "前端页面按照车主端和管理端两条主线组织。车主端强调低门槛和短路径，首页提供停车场搜索、筛选、推荐和快捷入口；停车场详情页展示车场位置、余位、价格和车位信息；我的预约、我的订单、停车记录和反向寻车页面组成车主停车闭环。管理端采用侧边栏布局，突出仪表盘、车场管理、车场概览和订单管理等高频后台操作。",
        "小组在前端设计中重点关注 M14 车主端体验。车主进入系统后应能迅速完成“搜索停车场—查看详情—预约车位—到场确认—入场停车—出场支付—查看记录”的流程；当预约失败、车位已满、支付失败、无权限或车牌为空时，页面应给出明确提示，避免用户误以为系统无响应。",
    ]:
        p(doc, text)
    add_table(
        doc,
        ["路由", "页面文件", "主要功能"],
        [
            ["/login", "Login.vue", "账号登录、角色识别和 token 保存"],
            ["/register", "Register.vue", "普通用户注册和基础信息填写"],
            ["/home", "Home.vue", "停车场搜索、筛选、推荐和预约入口"],
            ["/lots/:id", "LotDetail.vue", "停车场详情、车位列表和预约确认"],
            ["/reservations", "MyReservations.vue", "我的预约、确认和取消"],
            ["/orders", "MyOrders.vue", "进行中订单、待支付订单和支付入口"],
            ["/records", "ParkingRecords.vue", "历史停车记录查询"],
            ["/find-car", "FindCar.vue", "根据车牌查询当前停车位置"],
            ["/admin", "Dashboard.vue", "运营概览、收入和车位占用统计"],
            ["/admin/lots", "LotManagement.vue", "停车场管理、同步和编辑"],
        ],
        [3.0, 4.1, 6.8],
    )
    add_image(doc, HLD_IMAGES / "ui_owner_01_search.png", "图 5-1 车主端搜索与推荐页面原型", 4.9)
    add_image(doc, HLD_IMAGES / "ui_owner_02_detail.png", "图 5-2 停车场详情页面原型", 4.9)
    add_image(doc, HLD_IMAGES / "ui_owner_03_reserve.png", "图 5-3 预约确认页面原型", 4.9)
    add_image(doc, HLD_IMAGES / "ui_owner_05_findcar.png", "图 5-4 反向寻车页面原型", 4.9)

    h(doc, 2, "5.2 后端接口设计")
    for text in [
        "后端接口按照业务领域拆分。认证模块负责注册、登录、获取个人信息和权限识别；停车场模块负责停车场列表、详情、车位状态和数据同步；预约模块负责预约创建、确认、取消和过期释放；订单模块负责模拟车牌入出场、费用计算、支付、停车记录和反向寻车；管理后台模块负责运营概览、停车场管理、订单管理和异常处理。",
        "统一返回格式有利于前端处理。系统接口通常返回 code、msg 和 data 字段，前端可以根据 code 判断请求是否成功，根据 msg 展示提示，根据 data 渲染页面。对异常情况，如未登录、无权限、车位不足、重复预约、订单不存在等，后端应返回明确错误信息，前端再以 Element Plus 消息组件提示用户。",
    ]:
        p(doc, text)
    add_table(
        doc,
        ["模块", "接口示例", "说明"],
        [
            ["认证", "POST /api/auth/login", "校验用户名和密码，返回用户角色和 token"],
            ["停车场", "GET /api/lots", "按关键字、价格、距离和空位查询停车场"],
            ["车位", "GET /api/lots/{id}/spots", "查询目标停车场的车位状态"],
            ["预约", "POST /api/reservations", "创建预约并锁定一个可用车位"],
            ["预约", "POST /api/reservations/{id}/cancel", "取消预约并释放车位"],
            ["订单", "POST /api/license-plate/events", "模拟车牌入场或出场事件"],
            ["支付", "POST /api/orders/{id}/pay", "模拟支付并释放车位"],
            ["后台", "GET /api/admin/dashboard", "返回车场、车位、订单和收入汇总指标"],
            ["分析", "GET /api/analytics/revenue-trend", "返回收入趋势等图表数据"],
        ],
        [2.2, 5.0, 7.2],
    )
    h(doc, 2, "5.3 关键业务实现")
    for subtitle, texts in [
        (
            "5.3.1 停车场查询与推荐",
            [
                "停车场查询接口支持关键字、价格、区域和距离条件。若前端传入用户经纬度，后端可根据停车场坐标计算距离，并按距离或空位进行排序。停车场卡片需要同时展示名称、地址、余位、价格和状态，帮助车主快速判断是否适合前往。",
                "课程阶段未接入真实地图服务，因此前端用列表和模拟坐标展示停车场信息。该设计虽然简化了地图交互，但保留了后续接入地图 SDK 的数据结构：停车场表中包含 lat、lng 字段，前端也保留了定位和附近推荐的接口参数。",
            ],
        ),
        (
            "5.3.2 车位预约状态机",
            [
                "预约创建是系统最重要的状态转换之一。后端需要检查用户登录态、检查是否已有活跃预约、检查停车场是否存在可用车位，然后创建预约记录并把车位置为 reserved。预约默认锁定 15 分钟，超时后应释放车位，避免用户长期占用资源。",
                "预约确认表示车主已到达或准备入场，系统将预约状态更新为 confirmed。取消预约则把状态更新为 cancelled，并释放原锁定车位。该状态机保证车位不会被重复预约，也让前端能够清晰展示“待确认、已确认、已取消、已过期”等不同状态。",
            ],
        ),
        (
            "5.3.3 车牌入出场与计费支付",
            [
                "车牌识别在课程阶段采用模拟事件实现。入场事件创建 parking 状态订单，记录车牌、停车场、车位和入场时间；出场事件根据入场时间和停车场小时费率计算费用，将订单置为 pending_pay；支付接口完成后把订单置为 paid，并释放对应车位。",
                "计费逻辑需要考虑免费时长、按小时向上取整和异常订单。本项目的测试报告将入场、出场、支付、订单查询、停车记录和反向寻车作为 M6 模块重点测试内容，确保核心停车闭环能够正常演示。",
            ],
        ),
        (
            "5.3.4 管理后台与统计分析",
            [
                "管理后台面向停车场运营人员。仪表盘展示停车场数量、总车位、占用车位、今日收入、今日订单和活跃用户；车场管理支持查看停车场基础信息、余位和同步状态；订单管理支持查看进行中、待支付和已完成订单。",
                "数据分析接口基于订单和车位数据生成收入趋势、饱和度、周转率和转化率等指标。虽然课程阶段数据量较小，但统计口径的设计为后续接入真实业务数据提供了基础。",
            ],
        ),
    ]:
        h(doc, 3, subtitle)
        for text in texts:
            p(doc, text)
    add_image(doc, HLD_IMAGES / "ui_admin_01_overview.png", "图 5-5 管理后台运营概览页面原型", 5.5)
    add_image(doc, HLD_IMAGES / "ui_admin_02_access.png", "图 5-6 停车场接入管理页面原型", 5.5)

    h(doc, 2, "5.4 小组负责模块的详细设计")
    for text in [
        "小组在详细设计阶段将工作拆分到各责任域推进。M14 车主端作为系统面向普通用户的入口，需要把停车场查询、预约、支付、记录和寻车整合在一个清晰路径中。页面设计上应避免让用户理解复杂后台术语，而是使用“附近停车场”“我的预约”“待支付订单”“停车记录”等直接表达。",
        "M1 路内停车位视觉检测在课程阶段采用模拟识别数据展示。页面重点展示检测点位、车位状态、最后识别时间和人工复核入口。M9 共享停车关注时段冲突，要求同一车位同一时段不可重复预约。M10 充电车位关注充电桩状态，故障或占用状态不可预约。M11 违停取证关注证据字段完整性，至少应包含车牌、时间、地点、图片或视频、违停类型和审核状态。M12 月卡管理关注有效期、续费和到期提醒。",
    ]:
        p(doc, text)
    add_table(
        doc,
        ["模块", "设计重点", "关键验收点"],
        [
            ["M14 车主端", "首页搜索、详情、预约、支付、记录和寻车入口", "用户能完整走通停车闭环"],
            ["M1 路内检测", "模拟检测状态和人工复核", "状态包含空闲、占用、违停和异常"],
            ["M9 共享停车", "共享时段、预约冲突和收益展示", "同一时段不可重复预约"],
            ["M10 充电车位", "充电桩状态、预约和告警", "故障充电桩不可预约"],
            ["M11 违停取证", "证据列表、详情和审核处理", "证据信息完整且状态可追踪"],
            ["M12 月卡管理", "申请、审批、续费和到期提醒", "同车同场仅保留有效月卡"],
        ],
        [2.2, 6.0, 6.0],
    )
    add_image(doc, USE_CASE_IMAGES / "M1_use_case.png", "图 5-7 M1路内停车位视觉检测用例图", 5.25)
    add_image(doc, USE_CASE_IMAGES / "M10_use_case.png", "图 5-8 M10充电车位智能管理用例图", 5.25)
    add_image(doc, USE_CASE_IMAGES / "M11_use_case.png", "图 5-9 M11违停自动抓拍与取证用例图", 5.25)
    add_image(doc, USE_CASE_IMAGES / "M12_use_case.png", "图 5-10 M12长期停车月卡管理用例图", 5.25)

    h(doc, 2, "5.5 部署与运行设计")
    p(doc, "系统采用前后端分别启动的方式运行。后端进入 backend 目录后安装 requirements.txt 中的依赖，并通过 python main.py 启动；前端进入 frontend 目录后执行 npm install 和 npm run dev 启动 Vite 服务。演示时先启动后端，再启动前端，浏览器访问 http://localhost:5173 即可进入系统。")
    p(doc, "为了降低答辩演示风险，系统使用本地 SQLite 数据库和种子数据，避免依赖外部数据库、地图、支付或硬件服务。这样既能保证项目在普通开发环境下可运行，也便于测试人员复现功能流程。")


def chapter_6(doc: Document) -> None:
    h(doc, 1, "第六章 系统测试")
    h(doc, 2, "6.1 测试环境与方法")
    p(doc, "系统测试在 Windows 本地环境完成，后端使用 Python 3.10+ 启动 FastAPI 服务，前端使用 Node 18+ 启动 Vite 开发服务器，浏览器使用 Chrome 或 Edge。测试方法以手工功能测试和接口联调测试为主，重点验证认证、停车场、预约、订单、管理后台和数据接入模块。")
    add_table(
        doc,
        ["项目", "内容"],
        [
            ["测试版本", "SmartPark V1.0 MVP"],
            ["测试时间", "2026年6月22日 - 2026年6月24日"],
            ["后端环境", "Python 3.10+ / FastAPI / SQLAlchemy / SQLite"],
            ["前端环境", "Node 18+ / Vue3 / Vite / Chrome"],
            ["测试方式", "手工功能测试、接口调用验证、前后端联调验证"],
        ],
        [3.2, 10.5],
    )
    h(doc, 2, "6.2 测试范围与用例统计")
    add_table(
        doc,
        ["模块", "用例数", "测试重点", "结果"],
        [
            ["认证模块", "8", "注册、登录、重复注册、错误密码、个人信息和修改密码", "通过"],
            ["停车场模块", "9", "列表、搜索、排序、筛选、详情、车位和附近推荐", "通过"],
            ["预约模块", "7", "创建预约、车位已满、重复预约、确认、取消、过期释放", "通过"],
            ["支付/订单模块", "8", "入场、出场、计费、支付、订单、记录、反向寻车", "通过"],
            ["管理后台", "8", "仪表盘、车场概览、流量、收入、订单、释放车位和权限", "通过"],
            ["数据接入", "2", "数据同步和同步日志查询", "通过"],
        ],
        [2.9, 2.0, 7.2, 2.0],
    )
    h(doc, 2, "6.3 典型测试用例")
    add_table(
        doc,
        ["用例ID", "测试场景", "操作步骤", "预期结果"],
        [
            ["TC-AUTH-01", "正常登录", "输入 admin/admin123 并提交", "登录成功，返回管理员信息并进入后台"],
            ["TC-M3-02", "停车场搜索", "输入停车场名称或地址关键字", "列表仅展示匹配停车场"],
            ["TC-M5-01", "正常预约", "选择有空位停车场并提交车牌", "预约成功，车位锁定15分钟"],
            ["TC-M5-03", "重复预约", "已有活跃预约时再次预约", "系统提示已有活跃预约"],
            ["TC-M6-01", "模拟入场", "提交车牌和停车场入场事件", "创建 parking 状态订单并占用车位"],
            ["TC-M6-03", "模拟出场", "对进行中订单提交出场事件", "计算费用，订单进入待支付"],
            ["TC-M6-04", "模拟支付", "对待支付订单提交支付", "支付成功，车位释放"],
            ["TC-M13-08", "普通用户拒绝", "普通用户访问管理接口", "返回无权限提示"],
            ["TC-M2-01", "数据同步", "管理员触发停车场同步", "更新时间戳并写入同步日志"],
        ],
        [2.3, 3.0, 5.8, 4.4],
    )
    h(doc, 2, "6.4 测试结果与分析")
    for text in [
        "根据项目测试报告，V1.0 MVP 共执行 38 条功能测试用例，全部通过。认证、停车场、预约、订单支付、管理后台和数据接入模块均能按预期运行。测试结果说明，系统能够支撑从搜索停车场、预约车位、模拟入场、模拟出场、支付离场到后台统计的端到端流程。",
        "测试中重点关注状态一致性。例如创建预约后车位应被锁定，取消预约后车位应释放；模拟入场后订单应处于停车中状态，模拟出场后订单应进入待支付状态，支付完成后车位应释放。上述测试场景均通过，说明核心状态机设计可满足课程演示要求。",
        "当前测试仍存在不足。项目尚未进行高并发预约测试、压力测试、真实移动端兼容性测试和异常网络测试。后续如果将系统扩展为生产级应用，需要补充数据库事务测试、接口自动化测试、浏览器兼容性测试和安全性测试。",
    ]:
        p(doc, text)


def chapter_7(doc: Document) -> None:
    h(doc, 1, "第七章 总结与展望")
    h(doc, 2, "7.1 设计工作总结")
    for text in [
        "本文基于工程实训项目“AI-driven 城市智慧停车管理与诱导系统”，按照毕业设计（论文）格式完成了系统背景、相关技术、需求分析、概要设计、数据库设计、详细实现和测试结果的整理。系统采用 Vue3 + FastAPI + SQLAlchemy + SQLite 技术栈，实现了停车查询、车位预约、模拟入出场、模拟支付、停车记录、反向寻车、运营后台和数据分析等核心功能。",
        "从软件工程角度看，本项目覆盖了课程实训要求的多个关键环节：项目立项、任务分工、需求分析、系统架构设计、数据库建模、接口设计、前端实现、后端实现、测试验证和文档交付。通过将工程项目整理成毕业设计格式文档，可以更清楚地呈现从需求到实现再到测试的完整过程。",
        "通过本项目，小组成员进一步理解了用户路径设计、前后端接口协作、状态一致性、异常提示和测试用例设计的重要性。一个可演示系统不仅需要代码能够运行，还需要业务流程清晰、数据状态准确、用户操作可理解、异常情况可恢复。",
    ]:
        p(doc, text)
    h(doc, 2, "7.2 不足与未来展望")
    add_numbered(
        doc,
        [
            "真实地图与导航能力仍不足。当前主要使用列表、坐标和静态示意方式展示停车资源，后续可接入高德、百度或开源地图服务，实现真实位置展示、路径规划和距离计算。",
            "AI视觉识别尚未真实落地。课程阶段使用模拟检测数据，后续可接入摄像头、视频流和目标检测模型，实现路内车位状态识别和违停抓拍。",
            "支付与硬件均为模拟。后续可接入真实支付网关、道闸、诱导屏、充电桩和设备状态采集协议，让系统从教学演示走向真实业务。",
            "并发控制需要增强。多个用户同时预约同一车位时，应引入数据库事务、唯一约束、乐观锁或消息队列，避免状态冲突。",
            "测试体系需要完善。后续应补充接口自动化测试、端到端测试、压力测试、安全测试和移动端兼容性测试，提高系统质量。",
        ],
    )
    h(doc, 1, "致  谢")
    p(doc, "感谢指导老师在工程实训课程中对项目选题、阶段任务和文档格式的指导。感谢小组成员程晓洋、程子浩在需求分析、系统设计、后端接口、数据建模、测试验证和文档整理过程中的协作。通过本次实训，我对软件工程项目从需求到实现再到测试验收的全过程有了更具体的认识，也提升了将工程项目整理为正式技术文档的能力。")
    h(doc, 1, "参考文献")
    refs = [
        "[1] Vue.js 官方文档. Vue 3 Guide.",
        "[2] Element Plus 官方文档. Component Documentation.",
        "[3] FastAPI 官方文档. FastAPI Framework Documentation.",
        "[4] SQLAlchemy 官方文档. SQLAlchemy ORM Documentation.",
        "[5] SQLite 官方文档. SQLite Documentation.",
        "[6] Apache ECharts 官方文档. ECharts Handbook.",
        "[7] Axios 官方文档. Promise based HTTP client for the browser and node.js.",
        "[8] 项目组. AI-driven城市智慧停车管理与诱导系统项目立项报告, 2026.",
        "[9] 项目组. AI-driven城市智慧停车管理与诱导系统软件开发计划, 2026.",
        "[10] 项目组. AI-driven城市智慧停车管理与诱导系统需求分析报告, 2026.",
        "[11] 项目组. AI-driven城市智慧停车管理与诱导系统概要设计说明书, 2026.",
        "[12] 项目组. SmartPark V1.0 功能测试报告, 2026.",
    ]
    for ref in refs:
        para = p(doc, ref, style="参考文献")
        para.paragraph_format.first_line_indent = None


def appendices(doc: Document) -> None:
    h(doc, 1, "附录A 软件使用说明书")
    h(doc, 2, "A.1 后端启动")
    add_table(
        doc,
        ["步骤", "命令/操作", "说明"],
        [
            ["1", "cd smart-park-v1.0/backend", "进入后端目录"],
            ["2", "pip install -r requirements.txt", "安装 FastAPI、SQLAlchemy 等依赖"],
            ["3", "python main.py", "启动后端服务"],
            ["4", "访问 http://localhost:8000/docs", "查看 Swagger API 文档"],
        ],
        [1.5, 6.5, 5.8],
    )
    h(doc, 2, "A.2 前端启动")
    add_table(
        doc,
        ["步骤", "命令/操作", "说明"],
        [
            ["1", "cd smart-park-v1.0/frontend", "进入前端目录"],
            ["2", "npm install", "安装前端依赖"],
            ["3", "npm run dev", "启动 Vite 开发服务"],
            ["4", "访问 http://localhost:5173", "进入系统前端页面"],
        ],
        [1.5, 6.5, 5.8],
    )
    h(doc, 2, "A.3 演示账号")
    add_table(
        doc,
        ["角色", "用户名", "密码", "说明"],
        [
            ["管理员", "admin", "admin123", "进入管理后台，查看仪表盘、车场管理和订单管理"],
            ["普通用户", "user", "user123", "进入车主端，完成查询、预约、支付和记录查看"],
        ],
        [2.8, 3.0, 3.0, 5.5],
    )
    h(doc, 2, "A.4 演示流程")
    add_numbered(
        doc,
        [
            "使用普通用户账号登录，进入车主端首页。",
            "搜索停车场，查看停车场卡片中的余位、价格和地址。",
            "进入停车场详情页，选择车位并提交车牌号完成预约。",
            "在我的预约页面确认预约，模拟到场。",
            "通过后端或页面入口模拟车辆入场、出场和支付。",
            "查看我的订单、停车记录和反向寻车页面。",
            "使用管理员账号登录后台，查看停车场、订单和统计数据。",
        ],
    )
    page_break(doc)

    h(doc, 1, "附录B 主要源代码")
    add_code_block(doc, PROJECT / "backend" / "main.py", "B.1 FastAPI 应用入口 main.py", 1, 80)
    add_code_block(doc, PROJECT / "backend" / "models.py", "B.2 核心数据模型 models.py", 1, 130)
    add_code_block(doc, PROJECT / "backend" / "routers" / "lots.py", "B.3 停车场路由 lots.py", 1, 180)
    add_code_block(doc, PROJECT / "backend" / "routers" / "reservations.py", "B.4 预约路由 reservations.py", 1, 190)
    add_code_block(doc, PROJECT / "backend" / "routers" / "orders.py", "B.5 订单与支付路由 orders.py", 1, 210)
    add_code_block(doc, PROJECT / "frontend" / "src" / "router" / "index.js", "B.6 前端路由 router/index.js", 1, 120)
    add_code_block(doc, PROJECT / "frontend" / "src" / "utils" / "api.js", "B.7 Axios 封装 utils/api.js", 1, 110)


def build() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(str(TEMPLATE))
    normalize_styles(doc)
    update_cover(doc)
    update_integrity_statement(doc)
    trim_after_second_template_table(doc)
    update_headers(doc)
    add_footer(doc)

    front_matter(doc)
    chapter_1(doc)
    chapter_2(doc)
    chapter_3(doc)
    chapter_4(doc)
    chapter_5(doc)
    chapter_6(doc)
    chapter_7(doc)
    appendices(doc)

    doc.core_properties.author = "第19组"
    doc.core_properties.title = THESIS_TITLE_CN
    doc.core_properties.subject = "工程实训期末大作业"
    doc.save(OUT_FILE)
    return OUT_FILE


if __name__ == "__main__":
    print(build())
