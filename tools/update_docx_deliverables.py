# -*- coding: utf-8 -*-
"""Update generated Word deliverables to the 18-module SmartPark scope.

This only edits files under 文档（项目相关的作业所需提交文档生成在这里）.
It never touches the template directory under 交付物/巩固式解密模板.
"""
from pathlib import Path

from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn


ROOT = Path(r"D:\实训\SmartParkCloudLink-hdjd")
DOC_ROOT = ROOT / "文档（项目相关的作业所需提交文档生成在这里）"

MODULES = [
    ("M1", "路内停车位视觉检测", "模拟识别路内车位空闲、占用、违停和车牌信息，支持人工复核。", "P1", "丁梓钊"),
    ("M2", "停车场数据联网接入", "通过标准化接口汇聚路外停车场空位、价格和运营状态。", "P0", "程晓洋"),
    ("M3", "城市停车一张图", "展示停车场位置、空位、价格、距离和区域热度，支持筛选推荐。", "P0", "程子浩"),
    ("M4", "停车诱导信息发布", "维护诱导屏并按空位状态生成、模拟发布诱导内容。", "P1", "程子浩"),
    ("M5", "智能导航与车位预约", "推荐停车方案，支持预约车位并锁定15分钟。", "P0", "程子浩"),
    ("M6", "车牌识别与无感支付", "模拟车牌入出场、自动计费和无感支付，完成订单闭环。", "P0", "程子浩"),
    ("M7", "停车行为大数据分析", "统计周转率、平均停车时长、饱和度和区域热度。", "P1", "程晓洋"),
    ("M8", "动态定价策略引擎", "按时段、利用率、区域热度等规则模拟动态调价并记录原因。", "P1", "程晓洋"),
    ("M9", "共享停车管理平台", "支持车位主发布共享时段，用户预约共享车位并查看收益。", "P1", "丁梓钊"),
    ("M10", "充电车位智能管理", "管理充电车位和充电桩状态，处理预约、故障和异常占用。", "P2", "丁梓钊"),
    ("M11", "违停自动抓拍与取证", "模拟违停抓拍证据，支持审核、处理和监管推送。", "P2", "丁梓钊"),
    ("M12", "长期停车月卡管理", "支持月卡申请、审批、续费、绑定车辆和到期提醒。", "P2", "丁梓钊"),
    ("M13", "车场运营管理后台", "提供车位、订单、收入、设备和远程操作管理能力。", "P0", "程晓洋"),
    ("M14", "车主移动端APP/小程序", "提供搜索、预约、支付、停车记录、发票和寻车入口。", "P0", "丁梓钊"),
    ("M15", "反向寻车导航", "按车牌或车位号定位车辆，生成楼层区域和步行指引。", "P1", "程子浩"),
    ("M16", "设备运维管理平台", "集中管理摄像头、道闸、诱导屏、充电桩等设备台账和告警。", "P2", "程晓洋"),
    ("M17", "城市停车监管平台", "展示全市停车态势、供需分析、收费统计和政策评估结果。", "P1", "程晓洋"),
    ("M18", "商圈停车联合营销", "配置商户优惠、停车券核销和营销效果统计。", "P2", "程子浩"),
]


def set_run_style(run, bold=None, size=9.5):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.bold = bold


def set_cell(cell, text, bold=False, size=9.5):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    set_run_style(run, bold=bold, size=size)


def resize_table(table, row_count):
    while len(table.rows) < row_count:
        table.add_row()
    # python-docx cannot remove rows through public API.
    while len(table.rows) > row_count:
        tr = table.rows[-1]._tr
        tr.getparent().remove(tr)


def set_table(table, headers, rows):
    # Rebuild data rows instead of reusing old rows. Some source templates have
    # merged cells in existing rows, and writing into those rows can duplicate
    # one value across every column.
    while len(table.rows) > 1:
        tr = table.rows[-1]._tr
        tr.getparent().remove(tr)
    while len(table.rows) < len(rows) + 1:
        table.add_row()
    for col, header in enumerate(headers):
        set_cell(table.cell(0, col), header, bold=True)
    for row_index, row in enumerate(rows, start=1):
        for col, value in enumerate(row):
            set_cell(table.cell(row_index, col), value)


def replace_paragraph_contains(doc, needles, text):
    for paragraph in doc.paragraphs:
        if all(needle in paragraph.text for needle in needles):
            for run in paragraph.runs:
                run.text = ""
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.text = text
            set_run_style(run, size=10.5)
            return True
    return False


def remove_duplicate_paragraphs(doc, text):
    seen = False
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip() == text:
            if seen:
                p = paragraph._p
                p.getparent().remove(p)
            else:
                seen = True


def replace_first_paragraph_contains(doc, needles, text):
    return replace_paragraph_contains(doc, needles, text)


def find_docx(name_part):
    matches = [p for p in DOC_ROOT.rglob("*.docx") if p.name.startswith("19组") and name_part in p.name]
    if not matches:
        raise FileNotFoundError(name_part)
    return matches[0]


def update_start_report():
    path = find_docx("项目立项")
    doc = Document(path)
    if len(doc.tables) >= 4:
        set_table(
            doc.tables[3],
            ["一级模块", "功能说明", "优先级"],
            [(f"{mid} {name}", desc, priority) for mid, name, desc, priority, _ in MODULES],
        )
    doc.save(path)
    print(f"updated {path}")


def update_project_plan():
    path = find_docx("项目计划")
    doc = Document(path)
    replace_paragraph_contains(
        doc,
        ["项目目标"],
        "项目目标：构建覆盖18个一级功能模块的城市智慧停车管理与诱导系统，实现停车场数据联网接入、城市停车一张图、预约导航、车牌识别与无感支付、运营后台、车主移动端、数据分析、动态定价、共享停车、充电车位、违停取证、月卡、设备运维、监管平台和商圈联合营销等能力。实训阶段对AI视觉、GIS地图、真实支付和硬件设备采用模拟数据、状态面板和接口预留方式实现。",
    )
    if len(doc.tables) >= 5:
        wbs_rows = [
            ("1", "项目启动与需求调研", "1", "—", "易", "全体"),
            ("2", "项目计划、立项报告与需求分析", "3", "1", "中", "全体"),
            ("3", "系统设计与数据库设计", "3", "2", "中", "程晓洋"),
        ]
        for index, (mid, name, _desc, priority, owner) in enumerate(MODULES, start=4):
            difficulty = "中" if priority == "P0" else "中-易"
            effort = "3" if priority == "P0" else "2"
            wbs_rows.append((str(index), f"{mid} {name}", effort, "3", difficulty, owner))
        wbs_rows.extend(
            [
                (str(len(wbs_rows) + 1), "前后端联调与测试", "4", "4-21", "中", "全体"),
                (str(len(wbs_rows) + 2), "验收文档、演示脚本与答辩准备", "3", "22", "易", "全体"),
            ]
        )
        set_table(doc.tables[4], ["序号", "工作包", "工作量 / （人天）", "前置任务", "任务 / 易难度", "负责人"], wbs_rows)
    if len(doc.tables) >= 6:
        set_table(
            doc.tables[5],
            ["编号/任务", "W1 启动", "W2 需求", "W3 设计", "W4 核心开发", "W5 扩展开发", "W6 联调测试", "W7 文档验收", "W8 答辩"],
            [
                ("1 启动+立项", "█", "", "", "", "", "", "", ""),
                ("2 需求分析+计划", "█", "█", "", "", "", "", "", ""),
                ("3 系统设计+数据库", "", "█", "█", "", "", "", "", ""),
                ("4 M2/M3/M5/M6/M13/M14核心闭环", "", "", "█", "█", "", "", "", ""),
                ("5 M1/M4/M7/M8/M9/M15/M17重要模块", "", "", "", "█", "█", "", "", ""),
                ("6 M10/M11/M12/M16/M18扩展模块", "", "", "", "", "█", "█", "", ""),
                ("7 集成测试+问题修复", "", "", "", "", "", "█", "█", ""),
                ("8 验收材料+演示答辩", "", "", "", "", "", "", "█", "█"),
            ],
        )
    remove_duplicate_paragraphs(doc, "4  项目甘特图")
    remove_duplicate_paragraphs(doc, "4  项目甘特图")
    remove_duplicate_paragraphs(doc, "4 项目甘特图")
    remove_duplicate_paragraphs(doc, "说明： █ = 该周在施工，第1-3周为需求/设计阶段，第4-8周为编码/测试/验收阶段，W8为验收交付周。")
    doc.save(path)
    print(f"updated {path}")


def update_hld():
    path = find_docx("系统设计")
    doc = Document(path)
    replace_paragraph_contains(
        doc,
        ["系统共"],
        "系统共18个一级功能模块，按三级优先级分层：P0核心6个（M2、M3、M5、M6、M13、M14完整实现核心闭环）、P1重要7个（M1、M4、M7、M8、M9、M15、M17增强管理与监管价值）、P2扩展5个（M10、M11、M12、M16、M18采用模拟数据、状态面板和接口预留方式实现）。",
    )
    replace_first_paragraph_contains(
        doc,
        ["数据库使用SQLite"],
        "数据库使用SQLite，围绕用户、车辆、停车场、车位、预约、订单、共享车位、充电桩、违停证据、月卡、设备、诱导屏、动态定价、营销活动和统计快照等核心数据表设计，覆盖18个一级模块的数据需求。",
    )
    replace_first_paragraph_contains(
        doc,
        ["API端点总计"],
        "API端点按18个一级模块和支撑能力划分，覆盖停车场数据接入、一张图查询、诱导发布、预约、车牌识别与无感支付、统计分析、动态定价、共享停车、充电车位、违停取证、月卡、运营后台、车主端、反向寻车、设备运维、监管平台和联合营销等接口。",
    )
    replace_first_paragraph_contains(
        doc,
        ["前端路由：车主端"],
        "前端路由按18个一级模块组织：车主端覆盖城市停车一张图、预约导航、无感支付、共享停车、充电车位、月卡、反向寻车和营销活动；管理端覆盖数据接入、路内检测、诱导发布、运营后台、大数据分析、动态定价、违停取证、设备运维和城市监管。",
    )
    if len(doc.tables) >= 3:
        api_rows = []
        for mid, name, _desc, _priority, _owner in MODULES:
            route = {
                "M1": "/admin/roadside-detection",
                "M2": "/admin/data-access",
                "M3": "/map",
                "M4": "/admin/guide-screens",
                "M5": "/reserve",
                "M6": "/orders/checkout",
                "M7": "/admin/analytics",
                "M8": "/admin/pricing",
                "M9": "/shared-spots",
                "M10": "/charging-spots",
                "M11": "/admin/violations",
                "M12": "/monthly-cards",
                "M13": "/admin/operations",
                "M14": "/app",
                "M15": "/find-car",
                "M16": "/admin/devices",
                "M17": "/regulation",
                "M18": "/marketing",
            }[mid]
            api = {
                "M1": "GET/POST /api/roadside-detections",
                "M2": "GET/POST /api/data-access",
                "M3": "GET /api/lots/search",
                "M4": "POST /api/guide-screens/publish",
                "M5": "POST /api/reservations",
                "M6": "POST /api/license-plate/events, POST /api/payments/mock",
                "M7": "GET /api/analytics/*",
                "M8": "GET/POST /api/pricing-rules",
                "M9": "GET/POST /api/shared-spots",
                "M10": "GET/POST /api/charging-spots",
                "M11": "GET/POST /api/violations",
                "M12": "GET/POST /api/monthly-cards",
                "M13": "GET /api/admin/operations",
                "M14": "GET /api/app/home",
                "M15": "GET /api/find-car",
                "M16": "GET/POST /api/devices",
                "M17": "GET /api/regulation/dashboard",
                "M18": "GET/POST /api/marketing-campaigns",
            }[mid]
            tables = {
                "M1": "parking_spots, detections, violations",
                "M2": "parking_lots, sync_logs",
                "M3": "parking_lots, parking_spots",
                "M4": "guide_screens, parking_lots",
                "M5": "reservations, parking_spots",
                "M6": "vehicles, parking_orders",
                "M7": "parking_orders, analytics_snapshots",
                "M8": "pricing_rules, parking_orders",
                "M9": "shared_spots, reservations",
                "M10": "charging_piles, parking_spots",
                "M11": "violations, devices",
                "M12": "monthly_cards, vehicles",
                "M13": "parking_lots, orders, devices",
                "M14": "users, vehicles, reservations, orders",
                "M15": "parking_orders, parking_spots",
                "M16": "devices, maintenance_logs",
                "M17": "analytics_snapshots, violations",
                "M18": "marketing_campaigns, coupons",
            }[mid]
            api_rows.append((f"{mid} {name}", route, api, tables))
        set_table(doc.tables[2], ["模块", "前端路由示例", "后端API端点", "数据表"], api_rows)
    if len(doc.tables) >= 4:
        set_table(
            doc.tables[3],
            ["表名", "用途", "核心字段"],
            [
                ("users", "用户与角色权限", "id, username, password_hash, phone, role"),
                ("vehicles", "车辆与车牌信息", "id, user_id, plate_number, vehicle_type"),
                ("parking_lots", "停车场基础信息和接入状态", "id, name, address, region, total_spots, rate_per_hour"),
                ("parking_spots", "车位编号、类型和状态", "id, lot_id, spot_number, status, spot_type"),
                ("reservations", "预约锁定、确认、取消和过期", "id, user_id, spot_id, status, expire_at"),
                ("parking_orders", "入出场、计费和支付订单", "id, user_id, spot_id, plate_number, entry_time, amount, status"),
                ("shared_spots", "共享车位和共享时段", "id, owner_id, spot_id, available_start, available_end, price"),
                ("charging_piles", "充电桩状态和故障", "id, spot_id, status, fault_type, last_update"),
                ("violations", "违停证据和审核处理", "id, plate_number, location, violation_type, review_status"),
                ("monthly_cards", "月卡申请、审批和续费", "id, user_id, plate_number, card_type, valid_to, status"),
                ("devices", "设备台账、在线状态和运维日志", "id, lot_id, device_type, status, firmware_version"),
                ("guide_screens", "诱导屏位置和发布内容", "id, location, related_lot_id, content, publish_status"),
                ("pricing_rules", "动态定价规则和调价记录", "id, lot_id, period, factor, reason, enabled"),
                ("marketing_campaigns", "商圈营销活动和优惠规则", "id, name, merchant, rule_desc, valid_to, status"),
                ("analytics_snapshots", "统计指标快照", "id, lot_id, metric_type, metric_value, snapshot_at"),
            ],
        )
    doc.save(path)
    print(f"updated {path}")


def main():
    update_start_report()
    update_project_plan()
    update_hld()


if __name__ == "__main__":
    main()
