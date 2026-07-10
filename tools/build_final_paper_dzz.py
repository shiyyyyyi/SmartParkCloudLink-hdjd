from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path(r"D:\cxdownload\指导老师-学生姓名-毕设论文 2026届软件工程 (1).docx")
OUT_DIR = ROOT / "文档（项目相关的作业所需提交文档生成在这里）" / "期末大作业"
OUT_FILE = OUT_DIR / "指导老师-丁梓钊-毕设论文 2026届软件工程.docx"

HLD_IMAGES = ROOT / "文档（项目相关的作业所需提交文档生成在这里）" / "04_系统设计" / "_generated_hld_diagrams"
USE_CASE_IMAGES = ROOT / "文档（项目相关的作业所需提交文档生成在这里）" / "03_需求分析" / "_generated_use_case_diagrams"


def clear_document(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_east_asia_font(run, font_name="宋体") -> None:
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:ascii"), "Times New Roman")


def set_style(doc, style_name, font="宋体", size=11, bold=False, color=None, before=0, after=6, line=1.35):
    style = doc.styles[style_name]
    style.font.name = font
    style.font.size = Pt(size)
    style.font.bold = bold
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:ascii"), "Times New Roman")
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)

    set_style(doc, "Normal", "宋体", 10.5, False, None, 0, 6, 1.5)
    set_style(doc, "Heading 1", "黑体", 16, True, "000000", 18, 10, 1.35)
    set_style(doc, "Heading 2", "黑体", 14, True, "000000", 12, 6, 1.3)
    set_style(doc, "Heading 3", "黑体", 12, True, "000000", 8, 4, 1.25)
    set_style(doc, "Body Text", "宋体", 10.5, False, None, 0, 6, 1.5)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("丁梓钊  2026届软件工程毕业设计格式文档")
    set_east_asia_font(run, "宋体")
    run.font.size = Pt(9)


def p(doc, text="", style=None, align=None, bold=False, size=None, font="宋体", color=None):
    para = doc.add_paragraph(style=style)
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    set_east_asia_font(run, font)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return para


def h(doc, level, text):
    para = doc.add_heading(text, level=level)
    for run in para.runs:
        set_east_asia_font(run, "黑体")
        run.font.color.rgb = RGBColor(0, 0, 0)
    return para


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def apply_table_grid(table):
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "9EADBA")


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_cell_text(cell, text, bold=False, fill=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(str(text))
    set_east_asia_font(run, "宋体")
    run.font.size = Pt(9.5)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        shade_cell(cell, fill)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    apply_table_grid(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, head in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], head, bold=True, fill="D9EAF7", align=WD_ALIGN_PARAGRAPH.CENTER)
        if widths:
            table.rows[0].cells[i].width = Cm(widths[i])
    mark_header_row(table.rows[0])
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.CENTER if len(str(value)) <= 8 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], value, align=align)
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        try:
            para = doc.add_paragraph(style="List Bullet")
        except KeyError:
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(0.8)
            para.add_run("· ")
        run = para.add_run(item)
        set_east_asia_font(run, "宋体")
        run.font.size = Pt(10.5)


def add_numbered(doc, items):
    for idx, item in enumerate(items, 1):
        try:
            para = doc.add_paragraph(style="List Number")
        except KeyError:
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(0.8)
            para.add_run(f"{idx}. ")
        run = para.add_run(item)
        set_east_asia_font(run, "宋体")
        run.font.size = Pt(10.5)


def add_caption(doc, caption):
    para = p(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER, font="宋体", size=9.5)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(8)


def add_image(doc, path: Path, caption: str, width=5.7):
    if path.exists():
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        inline = run.add_picture(str(path), width=Inches(width))
        inline._inline.docPr.set("descr", caption)
        inline._inline.docPr.set("title", caption)
        add_caption(doc, caption)


def add_code_block(doc, file_path: Path, title: str, start=1, end=80):
    h(doc, 2, title)
    if not file_path.exists():
        p(doc, f"未找到文件：{file_path}")
        return
    lines = file_path.read_text(encoding="utf-8", errors="ignore").splitlines()
    snippet = "\n".join(f"{i + 1:03d}  {line}" for i, line in enumerate(lines[start - 1:end], start - 1))
    table = doc.add_table(rows=1, cols=1)
    apply_table_grid(table)
    mark_header_row(table.rows[0])
    cell = table.cell(0, 0)
    shade_cell(cell, "F3F6FA")
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(snippet)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    run.font.size = Pt(8)
    para.paragraph_format.line_spacing = 1.0
    doc.add_paragraph()


def add_cover(doc):
    for _ in range(2):
        p(doc)
    p(doc, "华东交通大学", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=22, font="黑体")
    p(doc, "毕业设计（论文）格式文档", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=20, font="黑体")
    for _ in range(3):
        p(doc)
    p(doc, "题    目：基于 Vue3 与 FastAPI 的智慧停车管理系统设计与实现", align=WD_ALIGN_PARAGRAPH.CENTER, size=15, font="宋体")
    p(doc, "项目名称：AI-driven 城市智慧停车管理与诱导系统", align=WD_ALIGN_PARAGRAPH.CENTER, size=15, font="宋体")
    p(doc, "学生姓名：丁梓钊", align=WD_ALIGN_PARAGRAPH.CENTER, size=15, font="宋体")
    p(doc, "专    业：软件工程", align=WD_ALIGN_PARAGRAPH.CENTER, size=15, font="宋体")
    p(doc, "届    别：2026届", align=WD_ALIGN_PARAGRAPH.CENTER, size=15, font="宋体")
    p(doc, "指导教师：________", align=WD_ALIGN_PARAGRAPH.CENTER, size=15, font="宋体")
    p(doc, "完成日期：2026年7月", align=WD_ALIGN_PARAGRAPH.CENTER, size=15, font="宋体")
    doc.add_page_break()


def add_front_matter(doc):
    h(doc, 1, "小组成员及个人工作说明")
    p(doc, "本项目来源于工程实训课程的“AI-driven 城市智慧停车管理与诱导系统”小组项目。根据课程期末大作业要求，本文按照毕业设计（论文）格式整理项目背景、需求分析、系统设计、数据库设计、详细实现、测试结果与总结展望。")
    add_table(
        doc,
        ["成员", "角色", "主要负责模块", "主要工作"],
        [
            ["程晓洋", "组长 / 后端主程 / 数据库管理员", "M2、M7、M8、M13、M16、M17", "负责数据接入、统计分析、动态定价、运营后台、设备运维与监管平台等设计和实现。"],
            ["程子浩", "需求分析师 / 后端主程 / QA", "M3、M4、M5、M6、M15、M18", "负责城市停车一张图、诱导发布、预约导航、无感支付、反向寻车和联合营销等业务闭环。"],
            ["丁梓钊", "产品经理 / 前端主程 / 测试工程师", "M1、M9、M10、M11、M12、M14", "参与需求细化与评审，负责车主端产品流、前端路由与组件设计，复核共享停车、充电车位、违停取证、月卡和车主端场景，并参与测试用例设计与验收。"],
        ],
        [2.2, 3.2, 3.2, 7.2],
    )
    p(doc, "说明：本文由丁梓钊基于小组实训项目资料、本地源码、需求分析文档、概要设计文档和测试报告独立整理撰写，内容重点突出本人参与的产品设计、前端设计和测试验证工作。")
    doc.add_page_break()

    h(doc, 1, "摘  要")
    p(doc, "随着城市机动车保有量持续增长，停车资源分布不均、空余车位信息不透明、人工收费效率低、停车场运营数据滞后等问题日益突出。为缓解车主找位难、停车场管理效率低和城市监管缺少统一视图等矛盾，本文围绕工程实训项目“AI-driven 城市智慧停车管理与诱导系统”，设计并实现了一个基于 Vue3 与 FastAPI 的智慧停车管理系统。")
    p(doc, "系统采用前后端分离的 B/S 架构。前端使用 Vue3、Element Plus、Vite、Pinia、Axios 和 ECharts 实现车主端与管理端页面；后端使用 FastAPI、SQLAlchemy 和 SQLite 构建 RESTful API、数据模型、预约状态机、订单计费、统计分析和管理接口。系统重点实现停车场数据接入、城市停车查询、车位预约、车牌入出场模拟、无感支付模拟、运营后台、停车记录和反向寻车等核心闭环，并对共享停车、充电车位、违停取证、长期月卡等扩展模块进行了原型化设计和接口预留。")
    p(doc, "测试结果表明，系统 V1.0 MVP 完成 38 条功能测试用例，认证、停车场查询、预约、订单支付、管理后台和数据接入等模块均能按预期运行。本文最后总结了项目实现过程中的经验，并从真实地图服务、AI视觉识别、硬件接入、并发控制和移动端体验等方面提出后续改进方向。")
    p(doc, "关键词：智慧停车；Vue3；FastAPI；前后端分离；车位预约；无感支付", bold=True)
    doc.add_page_break()

    h(doc, 1, "Abstract")
    p(doc, "With the continuous growth of urban vehicles, parking systems face problems such as uneven resource distribution, opaque vacancy information, inefficient manual charging, and delayed operational statistics. To address these issues, this paper presents the design and implementation of a smart parking management system based on Vue3 and FastAPI, derived from the engineering training project named “AI-driven Urban Smart Parking Management and Guidance System”.")
    p(doc, "The system adopts a front-end and back-end separated B/S architecture. The front end is developed with Vue3, Element Plus, Vite, Pinia, Axios and ECharts, while the back end is built with FastAPI, SQLAlchemy and SQLite. The implemented MVP covers parking lot data access, parking search, parking space reservation, simulated license plate entry and exit, simulated non-stop payment, operation dashboard, parking records and reverse car searching. Extended modules such as shared parking, charging parking spaces, violation evidence collection and monthly card management are also designed for further development.")
    p(doc, "Functional testing shows that all 38 test cases of the V1.0 MVP passed successfully. The system can support the main closed-loop process from searching a parking lot to reservation, entry, payment, exit and operation management. Finally, this paper summarizes the development experience and proposes future improvements in real map services, AI visual recognition, hardware integration, concurrency control and mobile user experience.")
    p(doc, "Keywords: Smart Parking; Vue3; FastAPI; Front-end and Back-end Separation; Reservation; Non-stop Payment", bold=True)
    doc.add_page_break()

    h(doc, 1, "目  录")
    toc = [
        "第1章 绪论",
        "1.1 研究背景与意义",
        "1.2 国内外研究与应用现状",
        "1.3 论文主要工作",
        "第2章 相关技术",
        "第3章 需求分析与系统概要设计",
        "第4章 数据库设计",
        "第5章 系统详细设计与实现",
        "第6章 系统测试",
        "第7章 总结与展望",
        "致谢",
        "参考文献",
        "附录A 系统使用说明书",
        "附录B 主要源代码",
    ]
    for item in toc:
        para = doc.add_paragraph()
        if re.match(r"^\d+\.", item) or item.startswith("附录"):
            para.paragraph_format.left_indent = Cm(0.8)
        run = para.add_run(item)
        set_east_asia_font(run, "宋体")
        run.font.size = Pt(11)
    doc.add_page_break()


def chapter_1(doc):
    h(doc, 1, "第1章 绪论")
    h(doc, 2, "1.1 研究背景与意义")
    p(doc, "城市停车问题是交通治理和智慧城市建设中的典型场景。传统停车系统往往以单个停车场为单位进行收费和管理，停车场之间缺少数据共享，车主在出行前难以准确掌握目的地周边空余车位、停车价格和入出场规则，容易出现绕行找位、排队缴费和取车困难等问题。对停车场运营方而言，车位状态、订单收入、设备状态、异常订单等数据分散在不同系统中，人工统计效率低，难以及时支撑运营决策。")
    p(doc, "智慧停车管理系统通过数据接入、车位预约、车牌识别、支付模拟、运营统计和监管分析，将车主、停车场和监管方纳入统一的信息平台。它不仅能够提升车主停车体验，也能够帮助停车场提高周转率和管理效率，并为城市停车政策评估提供数据基础。本项目在课程实训周期内采用模拟数据和可替换接口完成核心闭环，重点训练需求分析、系统设计、前后端开发、数据库建模和测试验收能力。")
    h(doc, 2, "1.2 国内外研究与应用现状")
    p(doc, "目前商业化停车产品已经普遍支持线上缴费、车牌识别和停车记录查询，一些大型商圈和交通枢纽也提供停车诱导屏、余位公示和会员优惠功能。但是多数系统仍然偏向单场景应用，城市级统一接入、共享停车、充电车位管理、违停取证、设备运维和监管分析等能力之间缺少完整闭环。")
    p(doc, "从技术路线看，前端单页应用、RESTful API、关系型数据库、数据可视化和状态机模型已经能够支撑中小规模智慧停车系统的教学演示。真实工业系统通常还需要接入地图服务、摄像头、道闸、支付网关、物联网设备和AI视觉模型。考虑到课程环境的时间和资源限制，本项目将这些外部能力抽象为模拟服务或预留接口，保证系统可本地运行、可演示、可测试。")
    h(doc, 2, "1.3 论文主要工作")
    add_numbered(doc, [
        "完成智慧停车项目的需求梳理，明确车主、管理员、监管人员等角色的核心业务场景，并形成18个一级功能模块的需求跟踪关系。",
        "设计前后端分离的总体架构，使用 Vue3 负责车主端和管理端交互，使用 FastAPI 提供认证、停车场、预约、订单、管理后台和数据分析接口。",
        "设计并实现 SQLite 数据库模型，围绕用户、停车场、车位、预约、订单和同步日志构建核心数据闭环。",
        "实现停车场查询、车位预约、模拟入场、模拟出场、计费支付、停车记录、反向寻车和管理后台统计等 MVP 功能。",
        "结合本人在小组中的职责，重点说明车主端产品流程、前端路由组织、组件化设计、异常提示口径和功能测试工作。",
    ])


def chapter_2(doc):
    h(doc, 1, "第2章 相关技术")
    h(doc, 2, "2.1 前端技术")
    p(doc, "Vue3 是本系统前端开发的核心框架。其组件化思想适合将停车场卡片、车位网格、订单列表、统计卡片等页面元素拆分为可维护的独立单元。Vite 提供快速启动和热更新能力，适合课程实训中的高频调试。Element Plus 提供表单、表格、菜单、弹窗和消息提示等通用组件，降低了管理后台和车主端页面的实现成本。")
    p(doc, "Pinia 用于保存登录用户、角色、令牌等状态；Vue Router 用于区分登录页、车主端首页、车场详情、我的预约、我的订单、停车记录、反向寻车和管理后台等页面；Axios 通过统一实例封装后端请求，并在请求拦截器中携带 token，在响应拦截器中统一处理错误提示。ECharts 用于展示运营后台收入趋势、饱和度、订单量等统计图表。")
    h(doc, 2, "2.2 后端技术")
    p(doc, "FastAPI 是基于 Python 类型标注的高性能 Web 框架，能够自动生成 Swagger API 文档，便于前后端联调和测试。系统使用 FastAPI 的 APIRouter 将认证、停车场、预约、订单、管理后台和分析模块拆分为独立路由文件，提升了代码可读性。")
    p(doc, "SQLAlchemy 用于对象关系映射，将 users、parking_lots、parking_spots、reservations、orders、sync_logs 等数据表映射为 Python 类。Pydantic 模型负责请求参数和响应数据的结构约束，例如注册字段长度、手机号格式、预约创建参数和订单响应结构。SQLite 作为本地文件数据库，部署成本低，适合课程演示。")
    h(doc, 2, "2.3 安全与接口技术")
    p(doc, "系统采用 token 方式保存用户登录状态，并通过后端工具函数解析当前用户身份。管理端接口在进入业务逻辑前检查用户角色，普通用户访问管理接口时返回无权限错误。密码使用哈希方式存储，避免明文密码直接进入数据库。接口设计遵循 RESTful 风格，以 JSON 作为前后端通信格式。")
    add_table(
        doc,
        ["层级", "技术选型", "在系统中的作用"],
        [
            ["前端框架", "Vue3 + Vite", "实现单页应用、组件化页面和快速开发调试"],
            ["UI与图表", "Element Plus + ECharts", "构建表单、表格、仪表盘和统计图表"],
            ["状态与请求", "Pinia + Axios", "管理登录状态、封装API请求和错误处理"],
            ["后端框架", "FastAPI", "提供RESTful API和Swagger接口文档"],
            ["数据访问", "SQLAlchemy + SQLite", "实现ORM持久化和本地演示数据库"],
        ],
        [2.8, 4.0, 7.0],
    )


def chapter_3(doc):
    h(doc, 1, "第3章 需求分析与系统概要设计")
    h(doc, 2, "3.1 用户角色与业务场景")
    add_table(
        doc,
        ["用户角色", "核心需求", "系统支持"],
        [
            ["车主", "快速找车位、预约、入出场、支付、查看记录、反向寻车", "车主端首页、停车场列表、预约、订单、停车记录、寻车页面"],
            ["停车场管理员", "查看车位和订单状态，统计收入，处理异常车位", "管理后台仪表盘、车场管理、订单管理、车位释放"],
            ["监管/运营人员", "掌握停车态势、分析收入和饱和度", "数据分析接口、ECharts统计图、报表数据"],
            ["车位主/扩展用户", "共享闲置车位、管理月卡和充电车位", "扩展模块设计与接口预留"],
        ],
        [2.5, 5.2, 5.7],
    )
    h(doc, 2, "3.2 功能需求")
    p(doc, "系统按项目要求划分为18个一级模块，其中 V1.0 MVP 聚焦 P0 核心闭环：停车场数据联网接入、城市停车一张图、智能导航与车位预约、车牌识别与无感支付、车场运营管理后台和车主移动端。扩展模块通过原型页面、模拟数据或接口预留体现。")
    add_table(
        doc,
        ["优先级", "模块", "实现说明"],
        [
            ["P0", "M2/M3/M5/M6/M13/M14", "构成查询、预约、入场、计费、支付、记录和运营后台的完整闭环"],
            ["P1", "M1/M4/M7/M8/M9/M15/M17", "增强感知、诱导、分析、共享、寻车和监管能力"],
            ["P2", "M10/M11/M12/M16/M18", "课程阶段以模拟数据、状态面板和接口预留呈现"],
        ],
        [2.0, 5.0, 6.0],
    )
    h(doc, 2, "3.3 非功能性需求")
    add_bullets(doc, [
        "性能要求：常规列表查询和详情查询应在本地环境中快速响应，统计类接口应避免阻塞核心操作。",
        "可用性要求：车主端页面需要适配常见浏览器和移动端宽度，按钮和提示应明确表达下一步操作。",
        "安全要求：认证接口、管理接口和用户数据接口需要校验登录态和角色权限。",
        "可维护性要求：前端按照 views、router、stores、utils 分层，后端按照 routers、models、schemas、database 分层。",
        "可扩展性要求：AI视觉、GIS地图、真实支付和硬件接入均通过模拟服务或预留字段支持后续替换。",
    ])
    h(doc, 2, "3.4 系统总体架构")
    p(doc, "系统采用浏览器端 SPA 与 FastAPI 后端服务分离的架构。浏览器通过 Axios 调用 RESTful API，后端路由处理业务请求并通过 SQLAlchemy 访问 SQLite 数据库。")
    add_image(doc, HLD_IMAGES / "02_architecture.png", "图 3-1 系统总体架构图", 5.8)
    h(doc, 2, "3.5 核心业务流程")
    p(doc, "核心业务流程从车主搜索停车场开始，经过车场详情查看、车位预约、到场确认、模拟入场、模拟出场、计费支付，再进入停车记录和运营统计。预约状态和车位状态是流程控制的关键。")
    add_image(doc, HLD_IMAGES / "03_activity.png", "图 3-2 核心业务活动图", 5.7)
    add_image(doc, HLD_IMAGES / "04_sequence.png", "图 3-3 预约与支付核心时序图", 5.7)
    h(doc, 2, "3.6 本人承担工作的需求落点")
    p(doc, "丁梓钊在小组中主要承担 PM、前端主程和测试工程师职责。在需求阶段，重点从可设计性、可测试性、可追溯性和用户体验角度审查需求；在设计阶段，复核车主端页面流、前端路由、组件划分和异常提示；在编码准备与测试阶段，围绕 M14 车主端、M1 路内检测、M9 共享停车、M10 充电车位、M11 违停取证和 M12 月卡管理形成页面与测试思路。")
    add_image(doc, USE_CASE_IMAGES / "M14_use_case.png", "图 3-4 M14车主移动端用例图", 5.5)


def chapter_4(doc):
    h(doc, 1, "第4章 数据库设计")
    h(doc, 2, "4.1 概念结构设计")
    p(doc, "系统核心数据围绕用户、停车场、车位、预约、订单和同步日志展开。用户发起预约，预约绑定停车场和车位；订单记录车辆入场、出场、费用和支付状态；停车场和车位保存总量、余位、价格和状态；同步日志记录停车场数据接入行为。")
    add_image(doc, HLD_IMAGES / "05_er.png", "图 4-1 系统核心ER图", 5.8)
    h(doc, 2, "4.2 逻辑结构设计")
    add_table(
        doc,
        ["数据表", "主要字段", "说明"],
        [
            ["users", "id, username, password_hash, phone, role", "保存用户账号、联系方式、角色和密码哈希"],
            ["parking_lots", "id, name, address, lat, lng, total_spots, available_spots, price_per_hour", "保存停车场基础信息、坐标、总车位、余位和价格"],
            ["parking_spots", "id, lot_id, spot_number, status", "保存车位编号和空闲、占用、预约等状态"],
            ["reservations", "id, user_id, lot_id, spot_id, plate_number, status, expires_at", "保存预约记录、车牌、预约状态和过期时间"],
            ["orders", "id, user_id, lot_id, spot_id, plate_number, entry_time, exit_time, amount, status", "保存停车订单、入出场时间、费用和支付状态"],
            ["sync_logs", "id, lot_id, status, message, created_at", "保存停车场数据同步记录和结果说明"],
        ],
        [2.8, 6.2, 5.0],
    )
    h(doc, 2, "4.3 关键状态约束")
    add_table(
        doc,
        ["对象", "状态", "转换说明"],
        [
            ["车位", "free / occupied / reserved", "预约时由 free 变为 reserved，到场或入场后变为 occupied，支付或取消后释放"],
            ["预约", "created / confirmed / cancelled / expired", "创建后15分钟内确认有效，取消或超时后释放车位"],
            ["订单", "parking / pending_pay / paid / exception", "入场生成 parking，出场后待支付，支付完成后 paid，异常进入 exception"],
            ["用户", "user / admin", "普通用户访问车主端，管理员可进入运营后台"],
        ],
        [2.4, 4.0, 7.6],
    )
    h(doc, 2, "4.4 数据一致性设计")
    p(doc, "预约、入场和支付等操作都会影响停车场余位和车位状态。后端在创建预约时检查是否存在活跃预约和可用车位，避免同一用户重复预约；取消预约、预约过期和支付完成后释放车位，保证车位状态与停车场余位一致。")


def chapter_5(doc):
    h(doc, 1, "第5章 系统详细设计与实现")
    h(doc, 2, "5.1 前端页面与路由设计")
    p(doc, "前端路由分为车主端和管理端。车主端包含登录、注册、首页、停车场详情、我的预约、我的订单、停车记录、反向寻车和个人中心；管理端采用嵌套路由，包括运营仪表盘、车场管理、车场概览和订单管理。")
    add_table(
        doc,
        ["路由", "页面", "功能"],
        [
            ["/home", "Home.vue", "搜索停车场、展示推荐停车场、进入详情和预约流程"],
            ["/lots/:id", "LotDetail.vue", "展示车场详情、车位列表和预约入口"],
            ["/reservations", "MyReservations.vue", "展示我的预约并支持确认、取消"],
            ["/orders", "MyOrders.vue", "展示订单并进入支付"],
            ["/records", "ParkingRecords.vue", "展示停车历史记录"],
            ["/find-car", "FindCar.vue", "根据车牌查找进行中停车位置"],
            ["/admin", "Dashboard.vue", "展示运营概览和统计图表"],
            ["/admin/lots", "LotManagement.vue", "管理停车场和同步状态"],
        ],
        [3.1, 4.0, 6.9],
    )
    add_image(doc, HLD_IMAGES / "ui_owner_01_search.png", "图 5-1 车主端搜索与推荐页面原型", 4.8)
    add_image(doc, HLD_IMAGES / "ui_owner_03_reserve.png", "图 5-2 车位预约页面原型", 4.8)
    add_image(doc, HLD_IMAGES / "ui_admin_01_overview.png", "图 5-3 管理后台运营概览页面原型", 5.5)
    h(doc, 2, "5.2 后端路由与接口设计")
    p(doc, "后端使用 APIRouter 组织接口。认证模块负责注册、登录、个人信息和密码修改；停车场模块提供列表、搜索、详情、附近推荐、车位查询和数据同步；预约模块负责创建、查询、确认、取消和过期释放；订单模块负责模拟入出场、支付、停车记录和反向寻车；管理后台模块提供统计概览、车场管理、订单管理和手动释放车位。")
    add_table(
        doc,
        ["模块", "接口示例", "说明"],
        [
            ["认证", "POST /api/auth/login", "校验用户名和密码，返回登录token和用户角色"],
            ["停车场", "GET /api/lots", "按关键字、价格、距离、空位等条件查询停车场"],
            ["预约", "POST /api/reservations", "创建预约并锁定可用车位15分钟"],
            ["订单", "POST /api/license-plate/events", "模拟车牌入场或出场事件"],
            ["支付", "POST /api/orders/{id}/pay", "模拟订单支付并释放车位"],
            ["管理后台", "GET /api/admin/dashboard", "返回停车场、车位、订单、收入等概览指标"],
            ["分析", "GET /api/analytics/revenue-trend", "返回收入趋势数据用于图表展示"],
        ],
        [2.4, 5.3, 6.6],
    )
    h(doc, 2, "5.3 关键业务实现")
    h(doc, 3, "5.3.1 停车场查询与距离排序")
    p(doc, "停车场列表接口支持关键字、价格区间、经纬度和排序方式。若前端传入用户位置，后端根据停车场坐标计算距离，并支持按距离从近到远排序；若传入价格或空位排序，则按业务字段返回结果。")
    h(doc, 3, "5.3.2 车位预约与超时释放")
    p(doc, "预约创建时，后端首先解析当前登录用户，再检查该用户是否已有 active/created 状态的预约；随后选择目标停车场下可用车位，将车位状态改为 reserved，将预约过期时间设为创建时间后15分钟。取消预约和过期释放都会恢复车位可用状态。")
    h(doc, 3, "5.3.3 车牌入出场与模拟支付")
    p(doc, "入场事件会创建 parking 状态订单并占用车位；出场事件根据入场时间和停车场单价计算停车时长与费用，并将订单置为 pending_pay；支付接口检查订单归属和状态，支付成功后标记 paid，同时释放对应车位。")
    h(doc, 3, "5.3.4 管理后台与统计分析")
    p(doc, "管理后台接口只允许管理员访问。仪表盘统计停车场数量、总车位数、占用车位数、今日收入、今日订单和活跃用户；分析模块以订单和车位数据为基础，返回收入趋势、饱和度、周转率和转化率等指标。")
    h(doc, 2, "5.4 本人前端与产品设计实现")
    p(doc, "本人在项目中重点关注车主端产品体验。车主端页面按照“搜索停车场 → 查看详情 → 预约车位 → 到场确认 → 入场停车 → 出场支付 → 查看记录/反向寻车”的路径组织，避免用户在多个入口之间迷失。前端通过路由守卫控制登录态，通过 Axios 拦截器统一携带 token，通过 Element Plus 消息组件展示预约失败、支付失败、权限不足等异常。")
    add_bullets(doc, [
        "M14车主端：明确首页搜索、推荐停车场、快捷入口、预约提醒和订单记录的页面组织。",
        "M1路内检测：以模拟检测数据和人工复核入口作为教学演示方案。",
        "M9共享停车：设计发布、列表、预约和收益查看流程，并强调时段冲突校验。",
        "M10充电车位：设计充电状态、预约和异常占用告警页面。",
        "M11违停取证：设计证据列表、详情查看和审核处理流程。",
        "M12月卡管理：设计申请、审批、续费和到期提醒流程。",
    ])


def chapter_6(doc):
    h(doc, 1, "第6章 系统测试")
    h(doc, 2, "6.1 测试环境与方法")
    p(doc, "系统测试在 Windows 本地环境完成，后端使用 Python 3.10+ 和 FastAPI/Uvicorn 启动，前端使用 Node 18+ 和 Vite 启动，浏览器使用 Chrome。测试以功能测试为主，覆盖认证、停车场查询、预约、支付、订单、管理后台和数据接入模块。")
    add_table(
        doc,
        ["项目", "内容"],
        [
            ["测试版本", "V1.0 MVP"],
            ["测试时间", "2026年6月22日 - 2026年6月24日"],
            ["后端环境", "Python 3.10+ / FastAPI / SQLAlchemy / SQLite"],
            ["前端环境", "Node 18+ / Vue3 / Vite / Chrome"],
            ["测试方式", "手工功能测试、接口调用验证、前后端联调验证"],
        ],
        [3.2, 10.4],
    )
    h(doc, 2, "6.2 测试范围")
    add_table(
        doc,
        ["模块", "用例数", "测试重点", "结果"],
        [
            ["认证模块", "8", "注册、登录、重复注册、错误密码、个人信息和修改密码", "通过"],
            ["停车场模块", "9", "列表、搜索、排序、筛选、详情、车位和附近推荐", "通过"],
            ["预约模块", "7", "创建预约、车位已满、重复预约、确认、取消、过期释放", "通过"],
            ["支付/订单模块", "8", "入场、出场、计费、支付、订单、记录、反向寻车", "通过"],
            ["管理后台", "8", "仪表盘、车场概览、流量、收入、订单、释放车位和权限", "通过"],
            ["数据接入", "2", "数据同步和同步日志查询", "通过"],
        ],
        [3.0, 2.2, 7.5, 2.2],
    )
    h(doc, 2, "6.3 典型测试用例")
    add_table(
        doc,
        ["用例ID", "场景", "操作", "预期结果"],
        [
            ["TC-M5-01", "正常预约", "选择有空位停车场并提交车牌", "预约成功，车位锁定15分钟"],
            ["TC-M5-03", "重复预约", "已有活跃预约时再次预约", "系统提示已有活跃预约"],
            ["TC-M6-03", "模拟出场", "对进行中订单提交出场事件", "计算停车时长和费用，订单进入待支付"],
            ["TC-M6-04", "模拟支付", "对待支付订单提交支付", "支付成功，车位释放"],
            ["TC-M13-08", "普通用户拒绝", "普通用户访问管理接口", "返回403无权限"],
            ["TC-M2-01", "数据同步", "管理员触发停车场同步", "更新时间戳并写入同步日志"],
        ],
        [2.2, 3.0, 5.8, 4.7],
    )
    h(doc, 2, "6.4 测试结论")
    p(doc, "V1.0 MVP 共设计并执行 38 条功能测试用例，全部通过，通过率为 100%。测试结果说明，系统能够稳定完成从搜索停车场、预约车位、模拟入出场、支付离场到管理后台统计的端到端流程。后续仍需补充并发预约、异常网络、真实移动端适配和压力测试。")


def chapter_7(doc):
    h(doc, 1, "第7章 总结与展望")
    h(doc, 2, "7.1 工作总结")
    p(doc, "本文基于工程实训项目，按照毕业设计（论文）格式完成了智慧停车管理系统的整理与论述。系统采用 Vue3 + FastAPI + SQLAlchemy + SQLite 技术栈，实现了停车查询、车位预约、模拟入出场、模拟支付、停车记录、反向寻车、运营后台和数据分析等核心功能，并结合课程要求对18个一级模块进行了需求跟踪和扩展设计。")
    p(doc, "在个人工作方面，本人主要承担 PM、前端主程和测试工程师职责。通过本项目，我更加熟悉了从需求审查、页面流程设计、前端路由组织、接口联调到测试验收的完整工程过程，也认识到一个可演示系统不仅需要代码能运行，还需要业务流程清晰、异常提示明确、数据状态一致。")
    h(doc, 2, "7.2 不足与改进方向")
    add_numbered(doc, [
        "真实地图与导航能力不足：当前主要使用列表、坐标和静态示意方式展示，后续可接入高德、百度或开源地图服务。",
        "AI视觉识别尚未真实落地：课程阶段使用模拟检测数据，后续可接入摄像头与目标检测模型。",
        "支付与硬件均为模拟：后续可扩展真实支付网关、道闸、诱导屏和充电桩设备协议。",
        "并发控制仍需增强：多个用户同时预约同一车位时，应引入数据库事务、锁机制或乐观并发控制。",
        "移动端体验可继续优化：车主端可进一步适配小程序或移动端 PWA，提高触摸操作体验。",
    ])
    h(doc, 1, "致  谢")
    p(doc, "感谢指导老师在工程实训课程中对项目选题、文档规范和阶段任务的指导。感谢小组成员程晓洋、程子浩在需求、设计、后端、测试和文档整理过程中的配合。通过本次实训，我对软件工程项目从需求到实现再到测试验收的全过程有了更具体的认识，也提升了将工程项目整理成正式技术文档的能力。")
    h(doc, 1, "参考文献")
    refs = [
        "[1] Vue.js 官方文档. Vue 3 Guide.",
        "[2] Element Plus 官方文档. Component Documentation.",
        "[3] FastAPI 官方文档. FastAPI Framework Documentation.",
        "[4] SQLAlchemy 官方文档. SQLAlchemy ORM Documentation.",
        "[5] SQLite 官方文档. SQLite Documentation.",
        "[6] ECharts 官方文档. Apache ECharts Handbook.",
        "[7] Axios 官方文档. Promise based HTTP client for the browser and node.js.",
        "[8] 项目组. AI-driven城市智慧停车管理与诱导系统需求分析报告, 2026.",
        "[9] 项目组. AI-driven城市智慧停车管理与诱导系统概要设计说明书, 2026.",
        "[10] 项目组. SmartPark V1.0 功能测试报告, 2026.",
    ]
    for ref in refs:
        p(doc, ref)


def appendices(doc):
    h(doc, 1, "附录A 系统使用说明书")
    h(doc, 2, "A.1 后端启动")
    add_code_instruction(doc, [
        "cd smart-park-v1.0/backend",
        "pip install -r requirements.txt",
        "python main.py",
        "访问 http://localhost:8000/docs 查看 Swagger API 文档",
    ])
    h(doc, 2, "A.2 前端启动")
    add_code_instruction(doc, [
        "cd smart-park-v1.0/frontend",
        "npm install",
        "npm run dev",
        "访问 http://localhost:5173 进入系统前端",
    ])
    h(doc, 2, "A.3 演示账号")
    add_table(
        doc,
        ["角色", "用户名", "密码", "说明"],
        [
            ["管理员", "admin", "admin123", "进入管理后台，查看仪表盘、车场管理和订单管理"],
            ["普通用户", "user", "user123", "进入车主端，完成查询、预约、支付和记录查看"],
        ],
        [3.0, 3.0, 3.0, 5.0],
    )
    h(doc, 1, "附录B 主要源代码")
    add_code_block(doc, ROOT / "smart-park-v1.0" / "backend" / "main.py", "B.1 FastAPI应用入口 main.py", 1, 45)
    add_code_block(doc, ROOT / "smart-park-v1.0" / "backend" / "models.py", "B.2 核心数据模型 models.py", 1, 115)
    add_code_block(doc, ROOT / "smart-park-v1.0" / "backend" / "routers" / "reservations.py", "B.3 预约路由 reservations.py", 1, 155)
    add_code_block(doc, ROOT / "smart-park-v1.0" / "frontend" / "src" / "router" / "index.js", "B.4 前端路由 router/index.js", 1, 60)
    add_code_block(doc, ROOT / "smart-park-v1.0" / "frontend" / "src" / "utils" / "api.js", "B.5 Axios封装 utils/api.js", 1, 80)


def add_code_instruction(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    apply_table_grid(table)
    mark_header_row(table.rows[0])
    cell = table.cell(0, 0)
    shade_cell(cell, "F3F6FA")
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    run.font.size = Pt(9)
    doc.add_paragraph()


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    clear_document(doc)
    setup_document(doc)

    add_cover(doc)
    add_front_matter(doc)
    chapter_1(doc)
    chapter_2(doc)
    chapter_3(doc)
    chapter_4(doc)
    chapter_5(doc)
    chapter_6(doc)
    chapter_7(doc)
    appendices(doc)

    doc.core_properties.author = "丁梓钊"
    doc.core_properties.title = "基于 Vue3 与 FastAPI 的智慧停车管理系统设计与实现"
    doc.save(OUT_FILE)
    print(OUT_FILE)


if __name__ == "__main__":
    main()
