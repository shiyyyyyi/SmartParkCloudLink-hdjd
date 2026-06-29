# -*- coding: utf-8 -*-
"""Build the system design deliverable from the converted HLD template.

The source template is copied/converted into the document output directory as
``04_系统设计/_template_base.docx``. This script preserves the template's cover
tables, sections, page setup, headers, footers, and styles, then replaces the
template guidance body with SmartPark HLD content.
"""
from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOC_ROOT = ROOT / "文档（项目相关的作业所需提交文档生成在这里）"
HLD_DIR = DOC_ROOT / "04_系统设计"
TEMPLATE = HLD_DIR / "_template_base.docx"
OUTPUT = HLD_DIR / "19组-程晓洋-系统设计.docx"
DIAGRAM_DIR = HLD_DIR / "_generated_hld_diagrams"

PROJECT_NAME = "AI-driven城市智慧停车管理与诱导系统"
GROUP = "第19组"
DOC_ID = "SmartPark-SD-HLD-001"
DATE = "2026.06.05"

MODULES = [
    ("M1", "路内停车位视觉检测", "P1", "模拟识别车位占用、违停和车牌，提供人工复核。", "roadside_detections, violations"),
    ("M2", "停车场数据联网接入", "P0", "接入路外停车场空位、价格和运营状态。", "parking_lots, sync_logs"),
    ("M3", "城市停车一张图", "P0", "汇总路内外停车资源，支持地图/列表查询和筛选。", "parking_lots, parking_spots"),
    ("M4", "停车诱导信息发布", "P1", "按空位与拥堵状态生成并发布诱导屏内容。", "guide_screens, publish_logs"),
    ("M5", "智能导航与车位预约", "P0", "推荐停车方案，锁定车位并管理预约状态。", "reservations, parking_spots"),
    ("M6", "车牌识别与无感支付", "P0", "模拟车牌入出场、自动计费和无感支付。", "vehicles, parking_orders"),
    ("M7", "停车行为大数据分析", "P1", "统计周转率、饱和度、停留时长和收入趋势。", "parking_orders, analytics_snapshots"),
    ("M8", "动态定价策略引擎", "P1", "按时段、热度和利用率计算价格系数。", "pricing_rules, pricing_logs"),
    ("M9", "共享停车管理平台", "P1", "支持共享车位发布、预约和收益统计。", "shared_spots, reservations"),
    ("M10", "充电车位智能管理", "P2", "管理充电车位、充电桩状态和异常占用。", "charging_piles, parking_spots"),
    ("M11", "违停自动抓拍与取证", "P2", "模拟违停抓拍、证据归档和审核处理。", "violations, review_logs"),
    ("M12", "长期停车月卡管理", "P2", "支持月卡申请、审批、续费和到期提醒。", "monthly_cards, vehicles"),
    ("M13", "车场运营管理后台", "P0", "承载车场、车位、订单、设备、价格和统计管理。", "parking_lots, parking_orders, devices"),
    ("M14", "车主移动端APP/小程序", "P0", "承载搜索、预约、支付、记录、寻车和优惠入口。", "users, vehicles, reservations"),
    ("M15", "反向寻车导航", "P1", "按车牌/订单定位车辆，生成楼层和路线指引。", "parking_orders, parking_spots"),
    ("M16", "设备运维管理平台", "P2", "管理摄像头、道闸、诱导屏、充电桩和告警。", "devices, maintenance_logs"),
    ("M17", "城市停车监管平台", "P1", "展示城市停车态势、违停、收费和政策评估。", "analytics_snapshots, violations"),
    ("M18", "商圈停车联合营销", "P2", "配置优惠活动、停车券核销和效果统计。", "marketing_campaigns, coupons"),
]

MODULE_DETAIL = {
    "M1": {
        "actor": "停车监管人员、停车场管理员",
        "page": "路内泊位检测面板、检测结果列表、人工复核弹窗、违停证据详情页",
        "service": "RoadsideDetectionService负责接收模拟识别结果，ViolationService负责违停记录归档和复核状态流转。",
        "api": "POST /api/roadside/detections, GET /api/roadside/detections, PUT /api/violations/{id}/review",
        "data": "roadside_detections保存泊位编号、识别状态、置信度、识别时间；violations保存车牌、位置、证据、复核状态。",
        "flow": "系统按摄像头或模拟数据源生成泊位识别结果，低置信度记录进入人工复核；复核通过后同步更新车位状态或生成违停记录。",
        "state": "detected、low_confidence、review_pending、confirmed、rejected。",
        "error": "识别结果缺少车牌或置信度过低时不直接改变车位状态，而是进入review_pending并提示管理员复核。"
    },
    "M2": {
        "actor": "停车场管理员、系统管理员",
        "page": "停车场接入配置页、同步状态面板、离线告警列表、接入日志详情",
        "service": "LotSyncService负责标准化接入、字段校验、空位数同步和离线告警；LotService负责车场基础信息维护。",
        "api": "POST /api/lots, PUT /api/lots/{id}, POST /api/lots/{id}/sync, GET /api/lots/{id}/sync-logs",
        "data": "parking_lots保存车场基础数据，parking_spots保存车位状态，sync_logs保存同步时间、结果和异常原因。",
        "flow": "管理员配置车场和模拟接入参数后，系统周期性同步空位、价格和营业状态；同步失败时保留最近一次有效数据。",
        "state": "active、inactive、syncing、sync_failed、offline。",
        "error": "接口超时、字段缺失或空位数异常跳变时生成接入告警，页面标注最后更新时间。"
    },
    "M3": {
        "actor": "车主、监管人员",
        "page": "城市停车一张图、车场列表、筛选面板、车场详情、空位热度面板",
        "service": "LotQueryService负责区域筛选、距离排序、价格排序和空位热度计算；MapAdapter在实训阶段使用模拟坐标。",
        "api": "GET /api/lots/search, GET /api/lots/nearby, GET /api/lots/{id}/spots, GET /api/lots/hotspots",
        "data": "parking_lots提供车场位置、价格和营业状态；parking_spots提供空闲、锁定、占用和不可用数量。",
        "flow": "车主输入目的地后，系统按区域、距离、价格、空位数和更新时间聚合车场结果，并支持进入详情查看车位状态。",
        "state": "normal、few_spots、full、closed、data_stale。",
        "error": "查询无结果时展示附近区域推荐；模拟地图不可用时切换列表模式并保留筛选能力。"
    },
    "M4": {
        "actor": "停车场管理员、诱导屏运维人员",
        "page": "诱导屏管理页、发布内容编辑页、发布记录页、设备状态列表",
        "service": "GuidePublishService负责根据车场空位、拥堵和设备在线状态生成诱导内容，并记录发布日志。",
        "api": "GET /api/guide-screens, POST /api/guide-screens/publish, GET /api/guide-screens/publish-logs",
        "data": "guide_screens保存屏幕位置和关联车场，publish_logs保存内容、发布时间、结果和失败原因。",
        "flow": "系统根据空位阈值生成诱导文案，管理员确认后模拟发布到诱导屏；发布成功后更新屏幕当前内容。",
        "state": "draft、publishing、published、failed、offline。",
        "error": "诱导屏离线时不覆盖上一条有效内容，发布日志记录失败原因并进入设备告警。"
    },
    "M5": {
        "actor": "车主、停车场管理员",
        "page": "推荐车场列表、预约确认页、预约倒计时、到场确认页、预约记录页",
        "service": "ReservationService负责车位锁定、预约确认、取消和超时释放；SpotService负责车位状态一致性。",
        "api": "POST /api/reservations, PUT /api/reservations/{id}/confirm, PUT /api/reservations/{id}/cancel, POST /api/reservations/release-expired",
        "data": "reservations保存预约用户、车位、状态和过期时间；parking_spots保存free/locked/occupied状态。",
        "flow": "车主选择车位后创建预约，系统锁定车位15分钟；到场确认后进入停车订单流程，超时或取消则释放车位。",
        "state": "created、locked、confirmed、cancelled、expired。",
        "error": "并发预约导致车位状态变化时返回409，前端刷新车位网格并提示重新选择。"
    },
    "M6": {
        "actor": "车主、停车场管理员",
        "page": "入出场模拟页、订单详情页、支付确认页、异常订单处理页",
        "service": "PlateRecognitionAdapter模拟车牌识别，OrderService生成订单并计费，PaymentAdapter模拟无感支付。",
        "api": "POST /api/license-plate/events, GET /api/orders/{id}, POST /api/orders/{id}/checkout, POST /api/payments/mock",
        "data": "vehicles保存车牌与用户关系，parking_orders保存入场、出场、金额和支付状态。",
        "flow": "车辆入场时识别车牌并生成订单，出场时按停车时长、免费时长和价格规则计算费用，支付成功后释放车位。",
        "state": "parking、pending_pay、paid、failed、manual_release。",
        "error": "车牌无法识别时生成待人工确认订单；支付失败时订单保持pending_pay并允许重试。"
    },
    "M7": {
        "actor": "停车场管理员、监管人员",
        "page": "行为分析图表页、收入趋势页、周转率面板、导出报表页",
        "service": "AnalyticsService负责按时间、车场、区域聚合停车时长、周转率、饱和度和收入趋势。",
        "api": "GET /api/analytics/turnover, GET /api/analytics/revenue, GET /api/analytics/saturation, GET /api/analytics/export",
        "data": "parking_orders提供订单样本，analytics_snapshots保存统计快照和统计口径。",
        "flow": "系统从订单、车位和车场数据聚合指标，前端使用ECharts展示趋势、对比和异常峰值。",
        "state": "calculated、sample_insufficient、stale、exporting。",
        "error": "样本不足或异常订单过多时图表标注数据口径，不把取消预约和异常订单混入正常统计。"
    },
    "M8": {
        "actor": "停车场管理员、运营人员",
        "page": "动态定价规则页、价格预览页、调价日志页、规则启停页",
        "service": "PricingService负责价格系数计算、规则冲突检查、价格预览和调价记录。",
        "api": "POST /api/pricing-rules, PUT /api/pricing-rules/{id}, POST /api/pricing-rules/evaluate, GET /api/pricing-logs",
        "data": "pricing_rules保存时段、车场、系数和启停状态，pricing_logs保存计算结果和调价原因。",
        "flow": "管理员配置时段、热度和利用率条件，系统计算价格系数并在订单计费时读取有效规则。",
        "state": "draft、enabled、disabled、conflict、expired。",
        "error": "规则冲突或超过监管限价时禁止发布，提示冲突规则并保留原价格。"
    },
    "M9": {
        "actor": "车位主、车主、平台管理员",
        "page": "共享车位发布页、共享车位列表、共享预约页、收益统计页",
        "service": "SharedSpotService负责共享时段发布、冲突校验、预约生成和收益汇总。",
        "api": "POST /api/shared-spots, GET /api/shared-spots, POST /api/shared-spots/{id}/reservations, GET /api/shared-spots/income",
        "data": "shared_spots保存共享车位、可用时段和价格，reservations保存共享预约关系。",
        "flow": "车位主发布可共享时段，车主查询并预约，系统校验时段冲突并在完成后生成收益记录。",
        "state": "published、reserved、in_use、closed、cancelled。",
        "error": "共享时段冲突、车位主取消或用户爽约时更新预约状态，收益只统计完成订单。"
    },
    "M10": {
        "actor": "新能源车主、停车场管理员",
        "page": "充电车位列表、充电桩状态页、充电预约页、异常占用告警页",
        "service": "ChargingService负责充电车位状态维护、故障标记、异常占用识别和充电预约。",
        "api": "GET /api/charging-spots, POST /api/charging-spots/{id}/reserve, PUT /api/charging-piles/{id}/status",
        "data": "charging_piles保存功率、故障和状态，parking_spots保存车位类型和占用状态。",
        "flow": "车主查看充电车位并预约，管理员维护充电桩状态；故障或普通车辆占用时生成告警。",
        "state": "available、reserved、charging、fault、occupied_abnormal。",
        "error": "故障充电桩不可预约，燃油车占用充电车位时进入异常占用告警。"
    },
    "M11": {
        "actor": "监管人员、停车场管理员",
        "page": "违停证据列表、证据详情、审核处理页、统计页",
        "service": "ViolationService负责模拟抓拍、证据归档、审核状态流转和监管汇总。",
        "api": "POST /api/violations, GET /api/violations, GET /api/violations/{id}, PUT /api/violations/{id}/review",
        "data": "violations保存车牌、位置、证据URL、违停时间和审核状态，review_logs保存复核意见。",
        "flow": "系统生成模拟违停证据，管理员或监管人员复核，通过后进入监管统计，驳回则保留记录。",
        "state": "pending_review、approved、rejected、archived。",
        "error": "证据缺失、车牌冲突或位置不明确时不得通过审核，只能进入人工补充。"
    },
    "M12": {
        "actor": "车主、停车场管理员",
        "page": "月卡申请页、审批页、续费页、到期提醒页",
        "service": "MonthlyCardService负责申请校验、审批、续费、到期提醒和临停规则切换。",
        "api": "POST /api/monthly-cards, PUT /api/monthly-cards/{id}/approve, POST /api/monthly-cards/{id}/renew, GET /api/monthly-cards/alerts",
        "data": "monthly_cards保存车牌、适用车场、有效期和审批状态，vehicles保存车辆绑定信息。",
        "flow": "车主提交月卡申请，管理员审批后生效；系统在到期前提醒续费，过期后按临停规则计费。",
        "state": "applied、approved、rejected、active、expired、renewed。",
        "error": "车牌不一致、适用车场不匹配或月卡过期时不按月卡放行，并提示按临停规则处理。"
    },
    "M13": {
        "actor": "停车场管理员、运营人员",
        "page": "运营概览、车场管理、车位管理、订单管理、收入报表、异常处理页",
        "service": "AdminService聚合车场、车位、订单、设备、收入和异常告警，提供管理端统一入口。",
        "api": "GET /api/admin/dashboard, GET /api/admin/lots, GET /api/admin/orders, GET /api/admin/reports, PUT /api/admin/exceptions/{id}",
        "data": "parking_lots、parking_spots、parking_orders、devices和analytics_snapshots共同支撑运营视图。",
        "flow": "管理员进入运营概览后可查看关键指标、处理异常、导出报表，并进入各类资源管理页面。",
        "state": "normal、warning、processing、resolved、exported。",
        "error": "删除、调价、关闭车场等敏感操作必须二次确认并写入操作日志。"
    },
    "M14": {
        "actor": "车主",
        "page": "车主首页、搜索页、车场详情、预约页、支付页、停车记录、个人中心",
        "service": "UserPortalService聚合车主端入口，调用车场、预约、订单、寻车和营销服务。",
        "api": "GET /api/user/home, GET /api/user/profile, GET /api/user/records, GET /api/user/coupons",
        "data": "users、vehicles、reservations、parking_orders和coupons共同组成车主端数据视图。",
        "flow": "车主从首页进入搜索、预约、支付、寻车和优惠使用，所有关键状态通过卡片和标签明确展示。",
        "state": "guest、logged_in、has_reservation、parking、pending_pay。",
        "error": "登录失效、预约超时、支付失败和数据延迟均通过页面顶部提示和按钮状态反馈。"
    },
    "M15": {
        "actor": "车主、停车场管理员",
        "page": "反向寻车入口、车辆位置详情、楼层/区域指引、寻车记录页",
        "service": "FindCarService根据进行中订单、车位区域和楼层信息生成可读路线。",
        "api": "GET /api/find-car, GET /api/records/{id}/location, POST /api/find-car/manual-help",
        "data": "parking_orders提供进行中停车记录，parking_spots提供楼层、区域和车位编号。",
        "flow": "车主输入车牌或选择进行中订单，系统返回车辆所在楼层、区域、车位号和文字路线。",
        "state": "located、not_found、manual_help_required。",
        "error": "定位信息不足时提示联系管理员或进入人工问询，不返回虚假路线。"
    },
    "M16": {
        "actor": "设备运维人员、停车场管理员",
        "page": "设备台账、故障告警、维护记录、设备详情页",
        "service": "DeviceService负责设备在线状态、故障等级、维护记录和心跳时间更新。",
        "api": "GET /api/devices, PUT /api/devices/{id}/status, POST /api/devices/{id}/maintenance, GET /api/devices/alerts",
        "data": "devices保存设备类型、位置、状态和心跳时间，maintenance_logs保存处理过程。",
        "flow": "系统展示摄像头、道闸、诱导屏、充电桩等设备状态，运维人员处理告警并记录维护结果。",
        "state": "online、offline、fault、maintenance、resolved。",
        "error": "设备离线影响核心业务时，系统生成告警并触发降级策略，例如保留最近识别结果或禁止诱导发布。"
    },
    "M17": {
        "actor": "交管监管人员、城市管理人员",
        "page": "监管大屏、区域供需分析、违停治理、收费趋势、政策评估页",
        "service": "RegulationService聚合多车场统计、区域饱和度、违停趋势和政策效果。",
        "api": "GET /api/regulation/dashboard, GET /api/regulation/regions, GET /api/regulation/violations, GET /api/regulation/policy-effect",
        "data": "analytics_snapshots、parking_lots、parking_orders和violations提供监管汇总数据。",
        "flow": "监管端读取脱敏后的汇总数据，按区域、时间和指标展示停车供需、违停和收费趋势。",
        "state": "normal、high_saturation、violation_risk、data_delay。",
        "error": "监管请求涉及个人明细时执行脱敏，数据延迟时标注更新时间和统计口径。"
    },
    "M18": {
        "actor": "商圈运营人员、车主、停车场管理员",
        "page": "营销活动配置页、停车券列表、核销页、活动效果统计页",
        "service": "MarketingService负责活动规则、停车券发放、核销和效果统计。",
        "api": "POST /api/campaigns, GET /api/campaigns, POST /api/coupons/redeem, GET /api/campaigns/{id}/stats",
        "data": "marketing_campaigns保存活动规则，coupons保存停车券状态，parking_orders关联核销订单。",
        "flow": "运营人员配置活动，车主领取或使用停车券，出场支付时按规则抵扣并生成核销记录。",
        "state": "draft、active、paused、expired、redeemed。",
        "error": "优惠券过期、重复核销或不满足活动条件时拒绝抵扣并提示原因。"
    },
}

CORE_TABLE_FIELDS = {
    "users": [
        ("id", "INTEGER", "主键", "用户唯一编号"),
        ("username", "VARCHAR(50)", "唯一、非空", "登录用户名"),
        ("password_hash", "VARCHAR(255)", "非空", "bcrypt密码哈希"),
        ("phone", "VARCHAR(20)", "可空", "手机号"),
        ("role", "VARCHAR(20)", "非空", "user/admin/owner/gov"),
        ("created_at", "DATETIME", "默认当前时间", "注册时间"),
    ],
    "parking_lots": [
        ("id", "INTEGER", "主键", "车场编号"),
        ("name", "VARCHAR(100)", "唯一、非空", "车场名称"),
        ("address", "VARCHAR(255)", "可空", "详细地址"),
        ("region", "VARCHAR(50)", "可空", "行政区域"),
        ("total_spots", "INTEGER", "非空", "总车位数"),
        ("rate_per_hour", "DECIMAL(10,2)", "非空", "基础费率"),
        ("status", "VARCHAR(20)", "默认active", "营业/离线/维护状态"),
    ],
    "parking_spots": [
        ("id", "INTEGER", "主键", "车位编号"),
        ("lot_id", "INTEGER", "外键", "所属车场"),
        ("spot_number", "VARCHAR(20)", "非空", "车位号"),
        ("floor", "VARCHAR(20)", "可空", "楼层"),
        ("zone", "VARCHAR(20)", "可空", "区域"),
        ("status", "VARCHAR(20)", "非空", "free/locked/occupied/disabled"),
        ("spot_type", "VARCHAR(20)", "默认normal", "普通/充电/月卡/无障碍"),
    ],
    "reservations": [
        ("id", "INTEGER", "主键", "预约编号"),
        ("user_id", "INTEGER", "外键", "预约用户"),
        ("spot_id", "INTEGER", "外键", "预约车位"),
        ("status", "VARCHAR(20)", "非空", "created/confirmed/cancelled/expired"),
        ("expire_at", "DATETIME", "非空", "锁定过期时间"),
        ("confirm_at", "DATETIME", "可空", "到场确认时间"),
    ],
    "parking_orders": [
        ("id", "INTEGER", "主键", "订单编号"),
        ("user_id", "INTEGER", "外键", "用户编号"),
        ("spot_id", "INTEGER", "外键", "停车车位"),
        ("plate_number", "VARCHAR(20)", "可空", "车牌号"),
        ("entry_time", "DATETIME", "非空", "入场时间"),
        ("exit_time", "DATETIME", "可空", "出场时间"),
        ("amount", "DECIMAL(10,2)", "可空", "应付金额"),
        ("status", "VARCHAR(20)", "非空", "parking/pending_pay/paid/exception"),
    ],
    "devices": [
        ("id", "INTEGER", "主键", "设备编号"),
        ("lot_id", "INTEGER", "外键", "所属车场"),
        ("device_type", "VARCHAR(30)", "非空", "camera/gate/screen/charger"),
        ("status", "VARCHAR(20)", "非空", "online/offline/fault/maintenance"),
        ("last_heartbeat", "DATETIME", "可空", "最近心跳时间"),
        ("remark", "TEXT", "可空", "故障或维护备注"),
    ],
}


def set_run(run, size=10.5, bold=False, font="宋体", color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def set_cell(cell, text, bold=False, size=9.5):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(str(text)) <= 14 else WD_ALIGN_PARAGRAPH.LEFT
    for part_index, part in enumerate(str(text).split("\n")):
        if part_index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(part)
        set_run(run, size=size, bold=bold, font="黑体" if bold else "宋体")


def clear_paragraph(paragraph):
    p_pr = paragraph._p.pPr
    for child in list(paragraph._p):
        if child is not p_pr:
            paragraph._p.remove(child)


def set_paragraph(paragraph, text, bold=False, size=10.5, font="宋体"):
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run(run, size=size, bold=bold, font=font)


def rebuild_table(table, headers, rows):
    while len(table.rows) > 1:
        tr = table.rows[-1]._tr
        tr.getparent().remove(tr)
    while len(table.rows) < len(rows) + 1:
        table.add_row()
    for col, header in enumerate(headers):
        set_cell(table.cell(0, col), header, bold=True)
    for row_index, row in enumerate(rows, start=1):
        for col, value in enumerate(row):
            set_cell(table.cell(row_index, col), value, size=9)


def paragraph_text(element):
    return "".join(node.text or "" for node in element.iter(qn("w:t")))


def strip_template_body(doc):
    body = doc.element.body
    children = list(body)
    start = None
    for index, child in enumerate(children):
        if child.tag == qn("w:p") and "Keywords 关键词" in paragraph_text(child):
            start = index
            break
    if start is None:
        raise RuntimeError("Cannot locate HLD template body start")
    for child in children[start:]:
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {min(level, 4)}")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run(run, size={1: 16, 2: 14, 3: 12, 4: 11}.get(level, 10.5), bold=True, font="黑体")
    return paragraph


def add_para(doc, text="", bold=False):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.2
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    set_run(run, bold=bold)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.24)
    paragraph.paragraph_format.first_line_indent = Inches(-0.12)
    run = paragraph.add_run(f"• {text}")
    set_run(run)


def add_code(doc, lines):
    for line in lines:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line)
        set_run(run, size=8.5, font="Consolas")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for col, header in enumerate(headers):
        set_cell(table.cell(0, col), header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for col, value in enumerate(row):
            set_cell(cells[col], value, size=8.8)
    doc.add_paragraph()
    return table


def add_module_expansion(doc):
    add_heading(doc, "2.2.2.3 18个一级模块详细展开", 4)
    add_para(doc, "本节按照教师任务图要求，对18个一级业务模块逐一展开设计。每个模块均说明使用角色、页面入口、后端服务、接口、数据对象、主流程、状态流转、异常处理和验收关注点，保证后续编码时能直接从模块设计追踪到接口、数据表和测试用例。")
    for index, (mid, name, priority, summary, objects) in enumerate(MODULES, start=1):
        detail = MODULE_DETAIL[mid]
        add_heading(doc, f"2.2.2.3.{index} {mid} {name}", 4)
        add_para(doc, f"{mid}属于{priority}优先级模块。模块目标是：{summary}该模块与数据对象 {objects} 直接相关，并通过统一接口返回格式与车主端、管理端或监管端页面联动。")
        add_para(doc, f"设计边界：{detail['actor']}是该模块的主要使用者；前端通过{detail['page']}承载业务操作；后端由{detail['service']}模块内部只处理本模块业务状态，不把登录、通用表格、图表渲染等支撑能力误列为额外一级模块。")
        add_para(doc, f"业务流程：{detail['flow']}状态口径为{detail['state']}。接口层负责身份校验、参数校验和错误码返回，服务层负责业务规则与状态流转，数据访问层负责事务提交和一致性保护。")
        add_table(
            doc,
            ["设计项", "详细说明"],
            [
                ("使用角色", detail["actor"]),
                ("页面/交互", detail["page"]),
                ("后端服务", detail["service"]),
                ("主要接口", detail["api"]),
                ("数据对象", detail["data"]),
                ("主流程", detail["flow"]),
                ("状态流转", detail["state"]),
                ("异常处理", detail["error"]),
                ("验收关注", f"测试用例需覆盖正常流程、异常输入、权限拦截、状态变化、数据库记录和页面提示；该模块优先级为{priority}，应在第4周编码计划中按优先级拆分。"),
            ],
        )
        add_para(doc, f"实现补充：{mid}在演示阶段应准备至少3组正常数据、1组边界数据和1组异常数据。页面需要展示当前状态和最近更新时间；接口需要返回可读msg；数据库记录需要保留创建时间、更新时间和处理人，便于日报、周报和测试记录中追踪。")
        add_para(doc, f"页面细化：{name}的页面不只展示静态信息，还需要体现查询条件、当前状态、操作入口和异常反馈。列表型页面应包含关键字搜索、状态筛选、分页、详情入口和刷新时间；表单型页面应包含必填校验、提交确认、取消返回和提交结果提示；审核型页面应保留审核意见、处理人和处理时间。")
        add_para(doc, f"接口细化：{mid}相关接口入参应包含业务对象ID、当前用户身份、分页/筛选参数和必要的状态字段；出参应包含业务数据、状态标签、更新时间和错误提示。对修改类接口，后端需在服务层再次读取数据库当前状态，避免前端缓存状态导致误操作。")
        add_para(doc, f"数据一致性：{mid}涉及的数据对象为{objects}。设计时要求页面状态、接口返回和数据库记录三者一致；当业务状态发生变化时，服务层必须先完成状态校验，再写入数据表，最后返回统一响应。对于影响预约、订单、支付、设备和监管统计的操作，需要写入操作日志。")
        add_table(
            doc,
            ["扩展设计点", "设计要求", "落地说明"],
            [
                ("页面控件", "搜索、筛选、详情、提交、取消、刷新、状态标签", f"{mid}页面根据角色显示不同按钮，禁止无权限操作直接出现在可提交状态。"),
                ("接口入参", "业务ID、用户身份、筛选条件、目标状态、备注", "修改类接口必须携带目标对象ID，批量操作需传入ID列表并逐项返回结果。"),
                ("接口出参", "code、msg、data、状态、更新时间", "错误时data可返回错误字段、冲突对象或建议下一步操作。"),
                ("数据库写入", "状态字段、时间字段、操作人、备注", "关键状态变化必须留痕，避免演示和测试时无法解释数据来源。"),
                ("权限控制", "车主、管理员、监管人员按角色访问", "车主只能访问本人数据，管理员访问所属车场，监管端访问脱敏汇总数据。"),
                ("异常反馈", "前端可读提示 + 后端错误码", detail["error"]),
            ],
        )
        add_table(
            doc,
            ["测试编号", "测试场景", "预期结果"],
            [
                (f"TC-{mid}-01", "正常查询/打开页面", "页面加载成功，列表或详情数据与数据库模拟数据一致"),
                (f"TC-{mid}-02", "正常新增/提交/发布/确认", "接口返回code=0，数据库新增记录或状态正确变化"),
                (f"TC-{mid}-03", "异常输入或缺少必填字段", "接口返回400，前端展示字段级提示，不写入错误数据"),
                (f"TC-{mid}-04", "权限不足访问", "接口返回401或403，前端提示登录或无权操作"),
                (f"TC-{mid}-05", f"{detail['error']}", "系统保留可追踪记录，页面提示清晰，核心状态不被错误覆盖"),
            ],
        )


def add_database_field_expansion(doc):
    add_heading(doc, "3.2.1 核心数据表字段展开", 3)
    add_para(doc, "以下字段设计用于支撑18个一级模块的编码落地。正式编码时可在不改变业务含义的前提下补充索引、默认值和审计字段，但不得删除影响需求追踪的关键字段。")
    for table_name, rows in CORE_TABLE_FIELDS.items():
        add_heading(doc, f"3.2.1 {table_name} 表字段说明", 4)
        add_table(doc, ["字段名", "类型", "约束", "说明"], rows)
        add_para(doc, f"{table_name}表的设计重点是保证业务状态可追踪。新增或更新记录时应写入更新时间，关键状态变化由服务层统一处理，避免页面直接拼接SQL造成状态不一致。")


def load_font(size=26, bold=False):
    candidates = [
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\simsun.ttc"),
        Path(r"C:\Windows\Fonts\msyh.ttc"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def make_diagram(name, title, boxes, arrows=None, size=(1300, 760)):
    DIAGRAM_DIR.mkdir(exist_ok=True)
    img = Image.new("RGB", size, "white")
    draw = ImageDraw.Draw(img)
    title_font = load_font(34, bold=True)
    body_font = load_font(22)
    small_font = load_font(18)
    draw.text((40, 28), title, fill=(30, 30, 30), font=title_font)
    for box in boxes:
        x, y, w, h, label, fill = box
        draw.rounded_rectangle((x, y, x + w, y + h), radius=8, fill=fill, outline=(60, 60, 60), width=2)
        lines = label.split("\n")
        for idx, line in enumerate(lines):
            draw.text((x + 18, y + 16 + idx * 30), line, fill=(20, 20, 20), font=body_font if idx == 0 else small_font)
    for arrow in arrows or []:
        x1, y1, x2, y2, label = arrow
        draw.line((x1, y1, x2, y2), fill=(40, 70, 120), width=3)
        # simple arrow head
        if x2 >= x1:
            draw.polygon([(x2, y2), (x2 - 12, y2 - 7), (x2 - 12, y2 + 7)], fill=(40, 70, 120))
        else:
            draw.polygon([(x2, y2), (x2 + 12, y2 - 7), (x2 + 12, y2 + 7)], fill=(40, 70, 120))
        if label:
            draw.text(((x1 + x2) // 2 - 45, (y1 + y2) // 2 - 28), label, fill=(40, 70, 120), font=small_font)
    out = DIAGRAM_DIR / f"{name}.png"
    img.save(out)
    return out


# ====== UI Mockup helpers ======

class _UI:
    pass

def UI_BAR(kind, text):
    return ("bar", kind, text)

def UI_SECTION(text):
    return ("section", text)

def UI_CARD(title, subtitle, color="blue"):
    return ("card", title, subtitle, color)

def UI_INFO(text):
    return ("info", text)

def UI_BUTTON(text, color="blue"):
    return ("button", text, color)

def UI_GRID(cols, rows, states):
    return ("grid", cols, rows, states)

def UI_STATS_ROW(items):
    return ("stats", items)


def make_ui_mockup(name, title, page_title, elements, size=(420, 760)):
    """Draw a mobile-style UI mockup."""
    DIAGRAM_DIR.mkdir(exist_ok=True)
    img = Image.new("RGB", size, (245, 245, 248))
    draw = ImageDraw.Draw(img)
    title_font = load_font(24, bold=True)
    body_font = load_font(16)
    small_font = load_font(13)
    tiny_font = load_font(11)
    
    # Status bar
    draw.rectangle((0, 0, size[0], 28), fill=(50, 50, 55))
    draw.text((15, 5), "9:41", fill=(255, 255, 255), font=small_font)
    draw.text((size[0] - 80, 5), "100% ▮▮▮▮", fill=(255, 255, 255), font=tiny_font)
    
    # Header bar
    draw.rectangle((0, 28, size[0], 72), fill=(24, 144, 255))
    draw.text((20, 38), page_title, fill=(255, 255, 255), font=title_font)
    
    y = 84
    for elem in elements:
        kind = elem[0]
        if kind == "bar":
            _, bar_type, text = elem
            if bar_type == "search":
                draw.rectangle((12, y, size[0] - 12, y + 38), fill=(255, 255, 255), outline=(200, 200, 210))
                draw.text((22, y + 10), text, fill=(160, 160, 170), font=small_font)
                y += 52
            elif bar_type == "back":
                draw.text((20, y + 2), "← " + text, fill=(255, 255, 255), font=body_font)
                y += 6
            elif bar_type == "input":
                draw.rectangle((12, y, size[0] - 12, y + 38), fill=(255, 255, 255), outline=(200, 200, 210))
                draw.text((22, y + 10), text, fill=(100, 100, 110), font=small_font)
                y += 52
            elif bar_type == "dashboard":
                draw.text((15, y + 2), text, fill=(255, 255, 255), font=body_font)
                y += 6
        elif kind == "section":
            _, text = elem
            draw.rectangle((0, y, size[0], y + 32), fill=(235, 237, 240))
            draw.text((16, y + 8), text, fill=(100, 100, 110), font=small_font)
            y += 40
        elif kind == "card":
            _, title, subtitle, color = elem
            colors = {"green": (0, 150, 80), "orange": (230, 140, 20), "red": (220, 50, 50), "blue": (24, 144, 255), "gray": (120, 120, 130)}
            c = colors.get(color, (24, 144, 255))
            draw.rectangle((12, y, size[0] - 12, y + 64), fill=(255, 255, 255), outline=(220, 220, 225))
            draw.rectangle((12, y, 18, y + 64), fill=c)
            draw.text((28, y + 8), title, fill=(30, 30, 30), font=body_font)
            draw.text((28, y + 32), subtitle, fill=(120, 120, 130), font=small_font)
            y += 74
        elif kind == "info":
            _, text = elem
            draw.text((16, y + 4), text, fill=(80, 80, 90), font=small_font)
            y += 26
        elif kind == "button":
            _, text, color = elem
            colors = {"green": (52, 199, 89), "orange": (255, 149, 0), "red": (255, 59, 48), "blue": (0, 122, 255), "gray": (142, 142, 147)}
            c = colors.get(color, (0, 122, 255))
            draw.rounded_rectangle((20, y, size[0] - 20, y + 42), radius=10, fill=c)
            tw = draw.textlength(text, font=body_font) if hasattr(draw, 'textlength') else len(text) * 16
            draw.text(((size[0] - tw) // 2, y + 10), text, fill=(255, 255, 255), font=body_font)
            y += 54
        elif kind == "grid":
            _, cols, rows, states = elem
            grid_y = y
            cell_w = (size[0] - 40) // cols
            cell_h = 50
            for r in range(rows):
                for c in range(cols):
                    cx = 20 + c * cell_w
                    cy = grid_y + r * cell_h
                    idx = r * cols + c
                    if idx < len(states):
                        st = states[idx]
                        if st == 0:
                            fill_c = (200, 240, 210)  # free green
                            label = "空闲"
                        elif st == 1:
                            fill_c = (255, 230, 180)  # locked orange
                            label = "锁定"
                        elif st == 2:
                            fill_c = (255, 200, 200)  # occupied red
                            label = "占用"
                        else:
                            fill_c = (220, 220, 220)
                            label = ""
                        draw.rectangle((cx + 2, cy + 2, cx + cell_w - 4, cy + cell_h - 4), fill=fill_c, outline=(180, 180, 185))
                        draw.text((cx + cell_w // 2 - 12, cy + cell_h // 2 - 10), label, fill=(80, 80, 80), font=tiny_font)
            y += rows * cell_h + 12
        elif kind == "stats":
            _, items = elem
            n = len(items)
            box_w = (size[0] - 30) // n
            for i, item in enumerate(items):
                sx = 12 + i * box_w
                draw.rectangle((sx, y, sx + box_w - 6, y + 56), fill=(255, 255, 255), outline=(225, 225, 230))
                draw.text((sx + box_w // 2 - 30, y + 8), item, fill=(30, 30, 30), font=tiny_font)
            y += 68
    
    # Bottom tab bar for owner pages
    if "车主端" in title:
        draw.rectangle((0, size[1] - 48, size[0], size[1]), fill=(248, 248, 250))
        draw.line((0, size[1] - 48, size[0], size[1] - 48), fill=(210, 210, 215), width=1)
        tabs = ["首页", "预约", "订单", "我的"]
        tab_w = size[0] // 4
        for i, tab in enumerate(tabs):
            tx = i * tab_w + tab_w // 2 - 12
            draw.text((tx, size[1] - 36), tab, fill=(100, 100, 110), font=small_font)
    
    # Title at bottom
    title2_font = load_font(20)
    tw = draw.textlength(title, font=title2_font) if hasattr(draw, 'textlength') else len(title) * 20
    draw.text(((size[0] - tw) // 2, size[1] - 22), title, fill=(60, 60, 70), font=title2_font)
    
    out = DIAGRAM_DIR / f"{name}.png"
    img.save(out)
    return out


def create_diagrams():
    diagrams = {}
    diagrams["context"] = make_diagram(
        "01_context",
        "第0层：系统上下文图",
        [
            (80, 170, 210, 110, "车主\n搜索/预约/支付/寻车", (230, 242, 255)),
            (80, 420, 210, 110, "停车场管理员\n车场/车位/订单", (230, 242, 255)),
            (500, 250, 300, 190, "智慧停车管理与诱导系统\nVue3 + FastAPI + SQLite\n模拟AI/GIS/支付/设备", (235, 248, 238)),
            (990, 170, 210, 110, "交管监管人员\n态势/违停/政策", (255, 242, 224)),
            (990, 420, 210, 110, "外部能力\nAI/GIS/支付/设备", (255, 242, 224)),
        ],
        [
            (290, 225, 500, 315, "REST"),
            (290, 475, 500, 375, "管理"),
            (800, 315, 990, 225, "监管数据"),
            (800, 375, 990, 475, "模拟接口"),
        ],
    )
    diagrams["architecture"] = make_diagram(
        "02_architecture",
        "第1层：系统结构图",
        [
            (90, 130, 240, 110, "表现层\n车主端/管理端/监管端", (229, 241, 255)),
            (510, 130, 260, 110, "接口层\nRESTful JSON / Session", (238, 245, 255)),
            (940, 130, 250, 110, "业务服务层\n预约/计费/分析/运维", (237, 248, 239)),
            (300, 430, 260, 110, "数据访问层\nSQLAlchemy ORM", (255, 247, 220)),
            (740, 430, 260, 110, "数据层\nSQLite + 模拟数据", (255, 237, 229)),
        ],
        [
            (330, 185, 510, 185, "HTTP"),
            (770, 185, 940, 185, "路由"),
            (1020, 240, 560, 430, "服务调用"),
            (560, 485, 740, 485, "ORM"),
        ],
    )
    diagrams["activity"] = make_diagram(
        "03_activity",
        "业务流程活动图：车主停车闭环",
        [
            (80, 140, 180, 80, "搜索目的地", (230, 242, 255)),
            (310, 140, 190, 80, "查看一张图\n筛选车场", (230, 242, 255)),
            (550, 140, 190, 80, "预约车位\n锁定15分钟", (235, 248, 238)),
            (790, 140, 180, 80, "导航入场", (235, 248, 238)),
            (1030, 140, 190, 80, "车牌识别\n生成订单", (255, 247, 220)),
            (310, 430, 190, 80, "停车记录\n反向寻车", (255, 247, 220)),
            (550, 430, 190, 80, "出场计费", (255, 237, 229)),
            (790, 430, 180, 80, "模拟支付", (255, 237, 229)),
            (1030, 430, 190, 80, "释放车位\n统计沉淀", (235, 248, 238)),
        ],
        [
            (260, 180, 310, 180, ""),
            (500, 180, 550, 180, ""),
            (740, 180, 790, 180, ""),
            (970, 180, 1030, 180, ""),
            (1125, 220, 405, 430, "停车中"),
            (500, 470, 550, 470, ""),
            (740, 470, 790, 470, ""),
            (970, 470, 1030, 470, ""),
        ],
    )
    diagrams["sequence"] = make_diagram(
        "04_sequence",
        "核心时序图：预约与无感支付",
        [
            (70, 140, 180, 90, "车主端", (230, 242, 255)),
            (300, 140, 190, 90, "API网关", (238, 245, 255)),
            (540, 140, 220, 90, "预约服务", (235, 248, 238)),
            (810, 140, 220, 90, "订单/支付服务", (255, 247, 220)),
            (1080, 140, 170, 90, "数据库", (255, 237, 229)),
            (70, 420, 1180, 110, "1 查询车场 -> 2 创建预约 -> 3 锁定车位 -> 4 入场识别 -> 5 生成订单 -> 6 出场计费 -> 7 模拟支付 -> 8 释放车位", (250, 250, 250)),
        ],
        [
            (250, 185, 300, 185, "POST"),
            (490, 185, 540, 185, "reserve"),
            (760, 185, 1080, 185, "update spot"),
            (1030, 185, 1080, 185, "order"),
        ],
    )
    # ====== 界面设计图 (UI mockups) ======
    diagrams["ui_owner_search"] = make_ui_mockup(
        "ui_owner_01_search",
        "车主端-首页/搜索",
        "SmartPark",
        [
            UI_BAR("search", "搜索目的地、车场名称..."),
            UI_SECTION("附近车场"),
            UI_CARD("万达广场停车场", "空位: 32/200   ¥6/h   1.2km", "green"),
            UI_CARD("市中心医院停车场", "空位: 5/120   ¥4/h   0.8km", "orange"),
            UI_CARD("火车站P1停车场", "空位: 89/300   ¥3/h   2.5km", "green"),
            UI_CARD("商业街地下停车场", "空位: 0/80   ¥8/h   0.3km", "red"),
        ],
    )
    diagrams["ui_owner_detail"] = make_ui_mockup(
        "ui_owner_02_detail",
        "车主端-车场详情",
        "车场详情",
        [
            UI_BAR("back", "万达广场停车场"),
            UI_INFO("地址: 中山路100号  |  总车位: 200  |  当前空位: 32"),
            UI_INFO("收费: ¥6/h  |  免费15分钟  |  24h营业"),
            UI_SECTION("车位选择"),
            UI_GRID(4, 5, [0,0,1,0,2, 0,1,0,2,1, 1,1,0,1,0, 2,0,1,0,1]),
            UI_BUTTON("立即预约", "blue"),
        ],
    )
    diagrams["ui_owner_reserve"] = make_ui_mockup(
        "ui_owner_03_reserve",
        "车主端-预约确认",
        "预约确认",
        [
            UI_BAR("back", "预约确认"),
            UI_INFO("车场: 万达广场停车场"),
            UI_INFO("车位: B1-023  |  收费: ¥6/h"),
            UI_CARD("预约倒计时", "14:32 (剩余)", "orange"),
            UI_INFO("请在15分钟内到达车场确认预约"),
            UI_BUTTON("到场确认", "green"),
            UI_BUTTON("取消预约", "gray"),
        ],
    )
    diagrams["ui_owner_pay"] = make_ui_mockup(
        "ui_owner_04_pay",
        "车主端-出场支付",
        "出场支付",
        [
            UI_BAR("back", "支付"),
            UI_INFO("订单号: ORD202606050001"),
            UI_INFO("车场: 万达广场停车场  |  车位: B1-023"),
            UI_INFO("入场: 2026-06-05 14:00  |  出场: 2026-06-05 16:30"),
            UI_CARD("停车费用", "¥15.00 (2.5h x ¥6/h)", "blue"),
            UI_INFO("优惠券: -¥5.00  |  实付: ¥10.00"),
            UI_BUTTON("模拟支付", "green"),
        ],
    )
    diagrams["ui_owner_findcar"] = make_ui_mockup(
        "ui_owner_05_findcar",
        "车主端-反向寻车",
        "反向寻车",
        [
            UI_BAR("back", "反向寻车"),
            UI_BAR("input", "输入车牌号搜索"),
            UI_INFO("车辆: 粤B·12345  |  车场: 万达广场"),
            UI_INFO("楼层: B1  |  区域: C区  |  车位号: B1-023"),
            UI_CARD("寻车路线", "入口 → 电梯下行B1 → 左转C区 → 直行50m → B1-023", "blue"),
            UI_BUTTON("查看地图指引", "blue"),
        ],
    )
    # 管理端
    diagrams["ui_admin_overview"] = make_ui_mockup(
        "ui_admin_01_overview",
        "管理端-运营概览",
        "运营概览",
        [
            UI_BAR("dashboard", "智慧停车运营管理平台"),
            UI_SECTION("今日概览"),
            UI_STATS_ROW(["总车位: 800", "占用: 612", "利用率: 76.5%", "今日收入: ¥3,240"]),
            UI_SECTION("实时告警"),
            UI_CARD("⚠ 设备离线", "火车站P1-3号摄像头  14:22", "orange"),
            UI_CARD("⚠ 异常订单", "ORD06050187 缺少入场时间", "red"),
            UI_SECTION("快速入口"),
            UI_BUTTON("车场管理", "blue"),
            UI_BUTTON("订单管理", "blue"),
            UI_BUTTON("设备运维", "blue"),
        ],
    )
    diagrams["ui_admin_access"] = make_ui_mockup(
        "ui_admin_02_access",
        "管理端-数据接入",
        "停车场数据接入",
        [
            UI_BAR("dashboard", "数据接入管理"),
            UI_SECTION("接入车场列表"),
            UI_CARD("万达广场停车场", "状态: 在线  |  同步: 2026-06-05 16:30  |  空位: 32", "green"),
            UI_CARD("火车站P1停车场", "状态: 在线  |  同步: 2026-06-05 16:28  |  空位: 89", "green"),
            UI_CARD("商业街地下停车场", "状态: 离线  |  上次同步: 2026-06-05 14:00", "red"),
            UI_BUTTON("+ 新增车场接入", "blue"),
        ],
    )
    diagrams["ui_admin_guide"] = make_ui_mockup(
        "ui_admin_03_guide",
        "管理端-诱导发布",
        "停车诱导信息发布",
        [
            UI_BAR("dashboard", "诱导发布管理"),
            UI_SECTION("诱导屏列表"),
            UI_CARD("中山路-北段诱导屏", "内容: 万达广场空位32个  |  状态: 已发布", "green"),
            UI_CARD("解放路-东段诱导屏", "内容: 火车站P1空位89个  |  状态: 已发布", "green"),
            UI_CARD("人民路-南段诱导屏", "内容: 待生成  |  状态: 未发布", "gray"),
            UI_BUTTON("生成诱导内容", "blue"),
            UI_BUTTON("批量发布", "green"),
        ],
    )
    diagrams["ui_admin_pricing"] = make_ui_mockup(
        "ui_admin_04_pricing",
        "管理端-动态定价",
        "动态定价策略",
        [
            UI_BAR("dashboard", "定价策略管理"),
            UI_SECTION("定价规则"),
            UI_CARD("高峰时段加价", "时段: 17:00-20:00  |  系数: 1.5x  |  状态: 启用", "green"),
            UI_CARD("节假日浮动", "适用: 周末/法定假日  |  系数: 1.2x  |  状态: 启用", "green"),
            UI_CARD("夜间优惠", "时段: 22:00-08:00  |  系数: 0.5x  |  状态: 启用", "blue"),
            UI_BUTTON("+ 新增规则", "blue"),
            UI_BUTTON("价格预览", "gray"),
        ],
    )
    diagrams["ui_admin_device"] = make_ui_mockup(
        "ui_admin_05_device",
        "管理端-设备运维",
        "设备运维管理",
        [
            UI_BAR("dashboard", "设备运维"),
            UI_SECTION("设备状态"),
            UI_STATS_ROW(["总数: 48", "在线: 42", "离线: 3", "故障: 3"]),
            UI_CARD("3号摄像头", "位置: 万达B1-C区  |  状态: 离线  |  最后心跳: 14:20", "red"),
            UI_CARD("1号道闸", "位置: 万达入口  |  状态: 故障  |  故障: 电机异常", "red"),
            UI_CARD("2号诱导屏", "位置: 中山路北段  |  状态: 在线", "green"),
            UI_BUTTON("处理告警", "orange"),
        ],
    )
    # 监管端
    diagrams["ui_reg_dashboard"] = make_ui_mockup(
        "ui_reg_01_dashboard",
        "监管端-监管大屏",
        "城市停车监管大屏",
        [
            UI_BAR("dashboard", "城市停车监管平台"),
            UI_STATS_ROW(["接入车场: 156", "总车位: 45,200", "当前占用率: 73.8%", "违停待审核: 23"]),
            UI_SECTION("区域停车态势"),
            UI_CARD("中心城区", "占用率: 82.5%  |  车场: 48  |  空位: 1,230", "orange"),
            UI_CARD("城东新区", "占用率: 55.2%  |  车场: 32  |  空位: 3,450", "green"),
            UI_CARD("火车站片区", "占用率: 91.3%  |  车场: 18  |  空位: 420", "red"),
            UI_SECTION("违停趋势"),
            UI_INFO("本周违停: 156  |  环比: +12%  |  已处理: 133  |  处理率: 85.3%"),
        ],
    )
    diagrams["ui_reg_violation"] = make_ui_mockup(
        "ui_reg_02_violation",
        "监管端-违停审核",
        "违停审核管理",
        [
            UI_BAR("dashboard", "违停审核"),
            UI_SECTION("待审核记录"),
            UI_CARD("粤B·A5678", "位置: 中山路禁停区  |  时间: 2026-06-05 15:20  |  状态: 待审核", "orange"),
            UI_CARD("粤B·C9012", "位置: 人民路消防通道  |  时间: 2026-06-05 14:50  |  状态: 待审核", "orange"),
            UI_CARD("粤B·D3456", "位置: 解放路公交站  |  时间: 2026-06-05 14:10  |  状态: 已确认", "red"),
            UI_BUTTON("查看证据图片", "blue"),
            UI_BUTTON("确认违停", "red"),
            UI_BUTTON("驳回", "gray"),
        ],
    )
    diagrams["ui_reg_policy"] = make_ui_mockup(
        "ui_reg_03_policy",
        "监管端-政策评估",
        "停车政策效果评估",
        [
            UI_BAR("dashboard", "政策评估"),
            UI_SECTION("价格政策效果"),
            UI_CARD("中心城区高峰加价", "实施: 2026-05  |  占用率变化: -8.2%  |  周转率: +15.3%", "green"),
            UI_CARD("夜间停车优惠", "实施: 2026-04  |  夜间占用率: +22.5%  |  投诉: -5.1%", "green"),
            UI_SECTION("违停趋势分析"),
            UI_STATS_ROW(["本月违停: 623", "环比: +8.5%", "处罚率: 72.3%", "投诉率: 3.2%"]),
            UI_SECTION("投诉统计"),
            UI_INFO("停车费过高: 45件  |  车位不足: 32件  |  诱导信息不准: 18件"),
        ],
    )
    diagrams["er"] = make_diagram(
        "05_er",
        "概念模型：核心ER关系",
        [
            (90, 150, 190, 90, "users\n1:N vehicles/reservations/orders", (230, 242, 255)),
            (390, 150, 210, 90, "parking_lots\n1:N spots/devices/orders", (235, 248, 238)),
            (710, 150, 210, 90, "parking_spots\nstatus/type/location", (235, 248, 238)),
            (1010, 150, 190, 90, "reservations\nlock/confirm/cancel", (255, 247, 220)),
            (260, 430, 210, 90, "parking_orders\nentry/exit/pay", (255, 237, 229)),
            (570, 430, 210, 90, "violations/devices\nreview/maintenance", (255, 237, 229)),
            (880, 430, 220, 90, "pricing/marketing\nrules/coupons", (255, 247, 220)),
        ],
        [
            (280, 195, 390, 195, "管理"),
            (600, 195, 710, 195, "包含"),
            (920, 195, 1010, 195, "预约"),
            (1010, 240, 365, 430, "生成订单"),
            (600, 240, 675, 430, "设备/违停"),
            (470, 475, 880, 475, "计费/营销"),
        ],
    )
    return diagrams


def add_figure(doc, path, caption):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(6.5))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    set_run(run, size=9, font="宋体")


def fill_front_matter(doc):
    for paragraph in doc.paragraphs:
        if "XX 软件系统概要设计说明书" in paragraph.text:
            set_paragraph(paragraph, f"{PROJECT_NAME}概要设计说明书", bold=True, size=22, font="黑体")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            break

    cover = doc.tables[0]
    set_cell(cover.cell(0, 1), PROJECT_NAME, bold=True)
    set_cell(cover.cell(1, 0), "项目编号：SmartPark-19")
    set_cell(cover.cell(2, 0), "项目编号")
    set_cell(cover.cell(3, 0), "SmartPark-19")
    set_cell(cover.cell(3, 1), "V1.0")
    set_cell(cover.cell(3, 2), DOC_ID)

    approval = doc.tables[1]
    set_cell(approval.cell(0, 1), "程晓洋（SA）")
    set_cell(approval.cell(0, 3), DATE)
    set_cell(approval.cell(1, 1), "程子浩、丁梓钊")
    set_cell(approval.cell(1, 3), DATE)
    set_cell(approval.cell(2, 1), "第19组")
    set_cell(approval.cell(2, 3), DATE)

    rebuild_table(
        doc.tables[2],
        ["Date / 日期", "Revision Version / 修订版本", "CR ID /Defect ID / CR/ Defect号", "Sec No. / 修改章节", "Change Description / 修改描述", "Author / 作者"],
        [
            ("2026.06.03", "V0.1", "N/A", "1-2", "根据SRS V1.0完成HLD简介、系统上下文和系统结构设计", "程晓洋（SA）"),
            ("2026.06.03", "V0.2", "N/A", "2.2", "补充系统架构图、业务流程活动图和模块分解", "全员"),
            ("2026.06.04", "V0.3", "DR-01~DR-06", "2-4", "补充核心业务时序图、接口定义、ER图、表结构和核心界面说明", "程晓洋（SA）"),
            ("2026.06.05", "V1.0", "DR-07~DR-10", "5/全部", "补充出错处理设计，完成全员设计评审并形成正式提交版", "第19组"),
        ],
    )
    rebuild_table(
        doc.tables[3],
        ["Abbreviations缩略语", "Full spelling 英文全名", "Chinese explanation 中文解释"],
        [
            ("HLD", "High-Level Design", "概要设计"),
            ("SRS", "Software Requirement Specification", "需求规格说明书"),
            ("SA", "System Architect", "系统架构师"),
            ("REST", "Representational State Transfer", "表现层状态转移接口风格"),
            ("ORM", "Object Relational Mapping", "对象关系映射"),
            ("ERD", "Entity Relationship Diagram", "实体关系图"),
            ("SPA", "Single Page Application", "单页应用"),
            ("P0/P1/P2", "Priority Level", "需求优先级"),
        ],
    )


def update_toc(doc):
    toc_lines = [
        "目    录",
        "1 简介",
        "1.1 目的",
        "1.2 范围",
        "1.2.1 软件名称",
        "1.2.2 软件功能",
        "1.2.3 软件应用",
        "1.3 参考资料",
        "2 概要设计",
        "2.1 第0层设计描述",
        "2.1.1 软件系统上下文定义",
        "2.1.2 设计思路",
        "2.2 第1层设计描述",
        "2.2.1 系统结构",
        "2.2.1.1 系统结构描述",
        "2.2.1.2 业务流程说明",
        "2.2.2 分解描述",
        "2.2.3 接口描述",
        "3 数据结构/数据库设计",
        "3.1 概念模型",
        "3.2 数据库表设计",
        "3.3 存储过程设计",
        "3.4 视图设计",
        "3.5 触发器设计",
        "3.6 函数设计",
        "3.7 基础数据配置",
        "4 界面设计",
        "4.1 车主端核心界面",
        "4.2 管理端核心界面",
        "4.3 监管端核心界面",
        "5 出错处理设计",
        "5.1 出错处理目标",
        "5.2 错误分类与处理策略",
        "5.3 统一错误码",
        "5.4 关键业务异常流程",
        "5.5 日志、告警与恢复",
        "5.6 设计评审与发布结论",
    ]
    start = None
    end = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith("目"):
            start = i
        if start is not None and "Keywords 关键词" in para.text:
            end = i
            break
    if start is None:
        return
    if end is None:
        end = min(len(doc.paragraphs), start + 40)
    for offset, line in enumerate(toc_lines):
        index = start + offset
        if index >= end:
            break
        para = doc.paragraphs[index]
        set_paragraph(para, line, bold=(offset == 0), font="黑体" if offset == 0 else "宋体")
    for index in range(start + len(toc_lines), end):
        set_paragraph(doc.paragraphs[index], "")


def add_body(doc):
    diagrams = create_diagrams()

    add_heading(doc, "Keywords 关键词：", 1)
    add_para(doc, "智慧停车；概要设计；系统上下文；模块分解；时序图；接口设计；数据库设计；界面设计")
    add_heading(doc, "Abstract 摘    要：", 1)
    add_para(
        doc,
        "本文档按照《概要设计说明书》模板编写，基于《功能需求规格说明书 SRS V1.0》对AI-driven城市智慧停车管理与诱导系统进行高层设计。文档覆盖系统上下文、架构设计、业务流程、模块分解、接口、数据库结构、核心界面和出错处理，为后续编码实现与测试提供依据。"
    )
    add_heading(doc, "List of abbreviations 缩略语清单：", 1)
    add_para(doc, "缩略语见前置表格。")

    add_heading(doc, "1 简介", 1)
    add_heading(doc, "1.1 目的", 2)
    add_para(doc, "本文档描述AI-driven城市智慧停车管理与诱导系统的概要设计，明确系统边界、总体架构、模块拆分、核心业务流程、接口和数据结构。预期读者包括项目成员、系统架构师、开发人员、测试人员和课程评审人员。")
    add_heading(doc, "1.2 范围", 2)
    add_heading(doc, "1.2.1 软件名称", 3)
    add_para(doc, PROJECT_NAME)
    add_heading(doc, "1.2.2 软件功能", 3)
    add_para(doc, "系统覆盖SRS中的18个一级业务模块。P0模块形成车主停车闭环与运营管理闭环，P1模块增强监管、分析、诱导和共享能力，P2模块通过模拟数据与预留接口表达充电、违停、月卡、设备运维和联合营销等扩展场景。")
    add_heading(doc, "1.2.3 软件应用", 3)
    add_para(doc, "软件应用于城市停车资源查询、车位预约、停车场运营、违停监管、设备运维和商圈营销等教学实训场景。实训环境下不依赖真实AI模型、GIS地图、支付通道和硬件设备，而是通过模拟接口、状态字段、初始数据和可视化页面完成业务闭环。")
    add_heading(doc, "1.3 参考资料", 2)
    for item in ["《功能需求规格说明书 SRS V1.0》", "《软件开发计划 V1.0》", "《项目立项报告 V1.0》", "教师发布的第三周项目任务作业与每日任务截图（2026.06.01-2026.06.05）"]:
        add_bullet(doc, item)

    add_heading(doc, "2 概要设计", 1)
    add_heading(doc, "2.1 第0层设计描述", 2)
    add_heading(doc, "2.1.1 软件系统上下文定义", 3)
    add_para(doc, "系统外部实体包括车主、停车场管理员、交管监管人员和外部模拟能力。车主通过车主端完成查询、预约、支付和寻车；管理员通过管理端维护车场、车位、订单、设备、价格和营销；监管人员查看汇总态势、违停和政策评估；AI识别、GIS、支付、道闸、诱导屏和充电桩在实训阶段以模拟服务接入。")
    add_figure(doc, diagrams["context"], "图2-1 软件系统上下文图")
    add_heading(doc, "2.1.2 设计思路", 3)
    add_para(doc, "系统采用前后端分离B/S结构。前端使用Vue3、Element Plus、Vite、ECharts组织车主端、管理端和监管端页面；后端使用FastAPI提供RESTful JSON接口；数据库使用SQLite和SQLAlchemy ORM；外部能力采用适配器模式预留服务接口，当前实现使用模拟数据和状态面板。")
    add_para(doc, "设计原则包括：模块边界对应SRS的18个一级业务模块；注册登录、基础CRUD、订单和统计图表作为支撑能力归入相关模块；核心闭环优先实现；高风险外部依赖采用可替换模拟服务。")

    add_heading(doc, "2.2 第1层设计描述", 2)
    add_heading(doc, "2.2.1 系统结构", 3)
    add_heading(doc, "2.2.1.1 系统结构描述", 4)
    add_figure(doc, diagrams["architecture"], "图2-2 系统结构图")
    add_table(
        doc,
        ["层次", "组成", "职责"],
        [
            ("表现层", "车主端、管理端、监管端", "提供搜索、预约、支付、运营、监管等页面"),
            ("接口层", "FastAPI Router、Session中间件", "统一鉴权、参数校验、错误码和JSON响应"),
            ("业务服务层", "预约、订单、分析、定价、运维等服务", "承载状态机、计费规则、统计聚合和模拟服务编排"),
            ("数据访问层", "SQLAlchemy ORM", "封装模型、事务、查询和数据一致性约束"),
            ("数据层", "SQLite、初始化数据、日志表", "保存业务状态、配置、订单和审计数据"),
        ],
    )
    add_heading(doc, "2.2.1.2 业务流程说明", 4)
    add_para(doc, "核心业务流程以车主停车闭环为主线：查询停车资源、预约锁定车位、导航入场、车牌识别生成订单、停车期间支持寻车、出场计费、模拟支付、释放车位并沉淀运营统计。管理端与监管端围绕该闭环读取和维护车场、订单、设备、违停、价格和分析数据。")
    add_figure(doc, diagrams["activity"], "图2-3 车主停车闭环活动图")
    add_heading(doc, "2.2.2 分解描述", 3)
    add_para(doc, "系统模块分解严格对应需求规格说明书中的18个一级业务模块，按P0/P1/P2安排实现优先级。")
    add_table(doc, ["模块ID", "模块名称", "优先级", "功能摘要", "主要数据对象"], MODULES)
    add_heading(doc, "2.2.2.1 P0核心闭环模块", 4)
    add_para(doc, "P0模块包括M2、M3、M5、M6、M13、M14，负责停车场数据接入、一张图查询、预约、车牌识别计费、运营后台和车主移动端，是课程演示必须稳定闭合的主路径。")
    add_heading(doc, "2.2.2.1.1 智能导航与车位预约", 4)
    add_para(doc, "类设计：ReservationService负责预约创建、取消、确认和过期释放；SpotService负责车位状态校验与锁定；LotService负责车场查询和推荐；Reservation模型保存预约状态、锁定到期时间和关联车位。")
    add_para(doc, "功能实现说明：车主端提交预约请求后，接口层校验登录态与参数，预约服务检查车位可用性，写入reservations并将parking_spots.status置为locked。超时任务扫描expire_at释放未确认预约。")
    add_figure(doc, diagrams["sequence"], "图2-4 预约与无感支付核心时序图")
    add_heading(doc, "2.2.2.1.2 车牌识别与无感支付", 4)
    add_para(doc, "类设计：PlateRecognitionAdapter模拟车牌识别结果；OrderService负责订单生成、计费和状态更新；PaymentAdapter模拟扣款；ParkingOrder模型记录入场、出场、金额和支付状态。")
    add_para(doc, "功能实现说明：车辆入场时系统根据车牌与预约状态生成订单；出场时按费率、免费时长和动态定价规则计算金额；模拟支付成功后订单完成并释放车位。")
    add_heading(doc, "2.2.2.2 P1/P2增强与扩展模块", 4)
    add_para(doc, "P1模块重点服务监管、分析和运营提效，P2模块采用模拟数据、状态字段和预留接口完成扩展表达。所有模块都通过统一API、统一错误码和统一数据字典与核心闭环联动。")
    add_module_expansion(doc)
    add_heading(doc, "2.2.3 接口描述", 3)
    add_para(doc, "接口统一采用HTTP/1.1 + RESTful JSON，统一响应格式为 {code, msg, data}。code=0表示成功，400表示参数错误，401表示未登录，403表示无权限，404表示资源不存在，500表示服务端异常。")
    add_table(
        doc,
        ["接口组", "方法与路径", "说明", "关联模块"],
        [
            ("停车资源", "GET /api/lots/search", "按位置、区域、价格、空位查询停车资源", "M2/M3"),
            ("预约", "POST /api/reservations", "创建预约并锁定车位15分钟", "M5"),
            ("预约确认", "PUT /api/reservations/{id}/confirm", "车辆到场后确认预约", "M5/M6"),
            ("车牌事件", "POST /api/license-plate/events", "模拟入出场车牌识别事件", "M6"),
            ("支付", "POST /api/payments/mock", "模拟无感支付并更新订单", "M6"),
            ("诱导屏", "POST /api/guide-screens/publish", "生成并发布诱导内容", "M4"),
            ("动态定价", "POST /api/pricing-rules/evaluate", "按规则计算价格系数", "M8"),
            ("监管态势", "GET /api/regulation/dashboard", "查询城市停车监管汇总", "M17"),
            ("设备运维", "PUT /api/devices/{id}/status", "更新设备状态和告警处理结果", "M16"),
            ("营销核销", "POST /api/coupons/redeem", "核销商圈停车优惠", "M18"),
        ],
    )

    add_heading(doc, "3 数据结构/数据库设计", 1)
    add_heading(doc, "3.1 概念模型", 2)
    add_para(doc, "核心概念模型围绕用户、车辆、停车场、车位、预约、订单、设备、违停证据、定价规则、营销活动和统计快照展开。用户与车辆、预约、订单存在一对多关系；停车场与车位、设备、订单存在一对多关系；预约确认后生成订单；订单、违停、定价和营销数据进入统计快照供运营和监管使用。")
    add_figure(doc, diagrams["er"], "图3-1 核心实体关系图")
    add_heading(doc, "3.2 数据库表设计", 2)
    add_table(
        doc,
        ["表名", "用途", "关键字段"],
        [
            ("users", "用户与角色权限", "id, username, password_hash, phone, role, created_at"),
            ("vehicles", "车牌与车辆信息", "id, user_id, plate_number, vehicle_type, default_flag"),
            ("parking_lots", "停车场基础信息和接入状态", "id, name, address, region, total_spots, rate_per_hour, status"),
            ("parking_spots", "车位状态、类型和位置", "id, lot_id, spot_number, floor, zone, status, spot_type"),
            ("reservations", "预约锁定、确认、取消和过期", "id, user_id, spot_id, status, expire_at, confirm_at"),
            ("parking_orders", "入出场订单、计费和支付", "id, user_id, spot_id, plate_number, entry_time, exit_time, amount, status"),
            ("shared_spots", "共享车位发布与收益", "id, owner_id, spot_id, available_start, available_end, price, status"),
            ("charging_piles", "充电桩和充电车位状态", "id, spot_id, power_kw, status, fault_type, last_update"),
            ("violations", "违停证据与审核", "id, plate_number, location, evidence_url, review_status, violation_at"),
            ("monthly_cards", "月卡申请、审批和续费", "id, user_id, plate_number, lot_id, valid_to, status"),
            ("devices", "设备台账、在线状态和运维日志", "id, lot_id, device_type, status, firmware_version, last_heartbeat"),
            ("guide_screens", "诱导屏位置和发布内容", "id, location, related_lot_id, content, publish_status"),
            ("pricing_rules", "动态定价规则", "id, lot_id, period, factor, reason, enabled"),
            ("marketing_campaigns", "商圈营销活动", "id, name, merchant, rule_desc, valid_to, status"),
            ("analytics_snapshots", "统计指标快照", "id, lot_id, metric_type, metric_value, snapshot_at"),
        ],
    )
    add_database_field_expansion(doc)
    add_heading(doc, "3.3 存储过程设计", 2)
    add_para(doc, "实训阶段使用SQLite，不设计数据库存储过程。预约超时释放、统计聚合和价格计算由后端服务层定时任务或接口逻辑实现。")
    add_heading(doc, "3.4 视图设计", 2)
    add_table(
        doc,
        ["视图名", "用途", "字段口径"],
        [
            ("v_lot_realtime_status", "停车一张图实时展示", "车场、区域、总车位、空余车位、价格、更新时间"),
            ("v_order_revenue_daily", "运营收入日报", "日期、车场、订单数、收入、平均停车时长"),
            ("v_regulation_summary", "监管汇总", "区域、车场数、总车位、占用率、违停数"),
        ],
    )
    add_heading(doc, "3.5 触发器设计", 2)
    add_para(doc, "不使用数据库触发器。订单完成、预约取消、支付成功等状态变更由服务层显式更新，并写入操作日志，避免演示环境中隐式逻辑难以追踪。")
    add_heading(doc, "3.6 函数设计", 2)
    add_para(doc, "主要函数包括calculate_fee(order, pricing_rule)、release_expired_reservations(now)、aggregate_lot_status(lot_id)、evaluate_pricing_factor(lot_id, time_range)和mask_regulation_data(records)。")
    add_heading(doc, "3.7 基础数据配置", 2)
    add_para(doc, "基础数据包括示例停车场、车位、车辆、管理员账号、设备、诱导屏、价格规则、营销活动和历史订单。初始化脚本应保证车主端、运营端和监管端都有可演示数据。")

    add_heading(doc, "4 界面设计", 1)
    add_para(doc, "本节展示系统的核心界面设计，包含车主端、管理端和监管端三大类界面。每个界面均配有说明表格和界面模拟图。")
    add_heading(doc, "4.1 车主端核心界面", 2)
    add_para(doc, "车主端采用移动端优先的卡片式布局，底部导航栏包含首页、预约、订单、我的四个入口。")
    add_table(
        doc,
        ["界面", "核心元素", "交互说明"],
        [
            ("首页/搜索", "目的地搜索框、附近车场卡片、空位/价格标签", "输入目的地后查询停车一张图并按距离与空位排序"),
            ("车场详情", "车场信息、车位网格、预约按钮、收费规则", "选择可预约车位后进入预约确认"),
            ("预约确认", "预约车位、倒计时、取消/到场确认按钮", "锁定15分钟，超时自动释放"),
            ("出场支付", "订单金额、停车时长、模拟支付按钮", "支付成功后完成订单并释放车位"),
            ("反向寻车", "车牌输入、楼层区域、路线文本", "按进行中订单生成寻车指引"),
        ],
    )
    add_figure(doc, diagrams["ui_owner_search"], "图4-1 车主端-首页/搜索")
    add_figure(doc, diagrams["ui_owner_detail"], "图4-2 车主端-车场详情")
    add_figure(doc, diagrams["ui_owner_reserve"], "图4-3 车主端-预约确认")
    add_figure(doc, diagrams["ui_owner_pay"], "图4-4 车主端-出场支付")
    add_figure(doc, diagrams["ui_owner_findcar"], "图4-5 车主端-反向寻车")
    add_heading(doc, "4.2 管理端核心界面", 2)
    add_para(doc, "管理端采用桌面端布局，左侧菜单导航，右侧内容区域展示数据表格、统计卡片和操作入口。")
    add_table(
        doc,
        ["界面", "核心元素", "交互说明"],
        [
            ("运营概览", "车位利用率、收入、订单、异常卡片", "查看车场经营状态并进入明细"),
            ("数据接入", "停车场接入配置、同步状态、离线告警", "维护路外停车场模拟接口"),
            ("诱导发布", "诱导屏列表、发布内容、发布状态", "生成并模拟发布诱导信息"),
            ("动态定价", "价格规则表、时段系数、发布记录", "配置并审计价格策略"),
            ("设备运维", "设备列表、故障等级、处理记录", "处理摄像头、道闸、诱导屏和充电桩状态"),
        ],
    )
    add_figure(doc, diagrams["ui_admin_overview"], "图4-6 管理端-运营概览")
    add_figure(doc, diagrams["ui_admin_access"], "图4-7 管理端-数据接入")
    add_figure(doc, diagrams["ui_admin_guide"], "图4-8 管理端-诱导发布")
    add_figure(doc, diagrams["ui_admin_pricing"], "图4-9 管理端-动态定价")
    add_figure(doc, diagrams["ui_admin_device"], "图4-10 管理端-设备运维")
    add_heading(doc, "4.3 监管端核心界面", 2)
    add_para(doc, "监管端以大屏展示和数据汇总为主，支持区域态势、违停审核和政策评估功能。")
    add_table(
        doc,
        ["界面", "核心元素", "交互说明"],
        [
            ("监管大屏", "区域占用率、供需趋势、违停数量", "展示城市级汇总数据"),
            ("违停审核", "证据图片、车牌、位置、复核状态", "审核模拟违停取证记录"),
            ("政策评估", "价格、饱和度、投诉/违停趋势", "辅助评估停车政策效果"),
        ],
    )
    add_figure(doc, diagrams["ui_reg_dashboard"], "图4-11 监管端-监管大屏")
    add_figure(doc, diagrams["ui_reg_violation"], "图4-12 监管端-违停审核")
    add_figure(doc, diagrams["ui_reg_policy"], "图4-13 监管端-政策评估")

    add_heading(doc, "5 出错处理设计", 1)
    add_heading(doc, "5.1 出错处理目标", 2)
    add_para(doc, "出错处理设计用于保证系统在参数错误、权限不足、数据冲突、设备离线、模拟接口失败和支付异常等情况下仍能给出明确反馈，并尽量保持核心业务状态一致。系统采用“前端提示 + 后端校验 + 数据状态保护 + 操作日志记录”的组合策略，使异常可发现、可定位、可恢复、可复盘。")
    add_para(doc, "本系统的出错处理范围覆盖车主端、管理端、监管端、后端服务、数据库访问和外部模拟能力。车主端重点保证用户知道下一步如何处理；管理端重点保证异常可进入人工处理；后端重点保证接口返回统一、状态不被错误写入；数据库重点保证预约、订单、车位和支付状态不会相互矛盾。")
    add_heading(doc, "5.2 错误分类与处理策略", 2)
    add_table(
        doc,
        ["错误场景", "触发条件", "处理策略", "用户/管理员提示"],
        [
            ("输入参数错误", "必填字段为空、格式不合法、分页参数越界", "后端Pydantic校验失败后返回400；前端表单同步提示错误字段", "请检查输入内容"),
            ("未登录或权限不足", "Session失效、普通用户访问管理接口、监管数据越权", "接口返回401/403，前端跳转登录或提示无权限；管理操作不落库", "请先登录/无权执行该操作"),
            ("预约车位已被占用", "创建预约前车位状态由free变为locked/occupied", "事务内重新校验车位状态，不创建预约；刷新车位列表", "车位已不可预约，请重新选择"),
            ("预约超时", "锁定时间超过15分钟且未确认到场", "定时任务释放车位并更新预约状态为expired；保留预约记录", "预约已超时，车位已释放"),
            ("模拟支付失败", "支付适配器返回失败、金额异常、订单已关闭", "订单保持pending_pay状态，允许重试或人工处理；不释放车位直到流程确认", "支付失败，请重试或联系管理员"),
            ("设备/诱导屏离线", "设备心跳超时、发布诱导内容失败", "保留最近一次有效状态并生成告警；发布失败时不覆盖原内容", "设备离线，已记录告警"),
            ("外部模拟接口异常", "AI/GIS/支付/硬件模拟服务不可用或返回空数据", "降级使用缓存/初始数据，记录日志；页面标注更新时间", "数据暂不可用，展示最近一次结果"),
            ("监管数据敏感", "监管端请求包含个人手机号、完整车牌或用户明细", "汇总脱敏后展示，不暴露个人手机号和完整车牌", "仅展示汇总统计"),
        ],
    )
    add_heading(doc, "5.3 统一错误码", 2)
    add_para(doc, "所有接口返回结构统一为 {code, msg, data}。成功时code=0，失败时code使用固定枚举，msg给出可读提示，data可携带错误字段、重试建议或异常追踪编号。前端Axios拦截器根据错误码统一处理登录失效、权限不足、参数错误和服务异常。")
    add_table(
        doc,
        ["错误码", "含义", "适用场景", "前端处理"],
        [
            ("0", "成功", "接口正常完成", "更新页面状态"),
            ("400", "参数错误", "表单字段为空、格式错误、范围越界", "展示字段级提示"),
            ("401", "未登录/登录失效", "Session不存在或过期", "跳转登录页"),
            ("403", "无权限", "角色不匹配、越权访问", "提示无权操作"),
            ("404", "资源不存在", "车位、订单、预约、设备不存在", "提示重新刷新"),
            ("409", "状态冲突", "车位被占、预约已取消、订单已支付", "刷新业务状态并提示重试"),
            ("429", "操作过于频繁", "短时间重复预约、重复支付、重复提交", "限制按钮并提示稍后再试"),
            ("500", "服务端异常", "数据库异常、服务逻辑异常", "展示通用错误并记录追踪号"),
            ("503", "外部模拟服务不可用", "AI/GIS/支付/设备模拟接口失败", "展示缓存数据或降级提示"),
        ],
    )
    add_heading(doc, "5.4 关键业务异常流程", 2)
    add_table(
        doc,
        ["业务流程", "异常点", "处理步骤", "状态结果"],
        [
            ("车主预约车位", "车位被并发预约", "重新读取车位状态；若非free则拒绝创建预约；返回409；前端刷新车位网格", "无新预约，车位状态以数据库为准"),
            ("预约到场确认", "预约已超时", "校验expire_at；超时则更新预约为expired并释放车位；提示重新预约", "reservation=expired, spot=free"),
            ("车牌入场", "车牌无法识别或无预约", "生成待人工确认记录；允许管理员补录车牌或绑定订单", "order=pending_confirm"),
            ("出场计费", "订单缺少入场时间或费率异常", "拒绝结算；写入异常日志；管理员修正订单数据后重试", "order=exception"),
            ("模拟支付", "支付失败或重复支付", "失败保持待支付；重复支付返回409并读取已有支付结果", "order=pending_pay或paid"),
            ("诱导信息发布", "诱导屏离线", "不覆盖上一条有效发布内容；生成设备告警；记录发布时间和失败原因", "publish_log=failed"),
            ("违停取证审核", "证据不足或车牌冲突", "进入人工复核状态；不进入监管确认列表", "violation=pending_review"),
        ],
    )
    add_heading(doc, "5.5 日志、告警与恢复", 2)
    add_para(doc, "日志策略：关键操作记录操作者、时间、接口路径、对象ID、原状态、新状态、处理结果和错误码；异常日志记录请求参数摘要、异常类型、追踪编号和处理方式，便于测试阶段复盘。日志不保存明文密码、完整支付信息和不必要的个人敏感信息。")
    add_para(doc, "告警策略：设备离线、接口同步失败、支付失败、预约状态冲突、统计数据异常和诱导发布失败均进入告警列表。告警等级分为提示、重要、严重三级；管理端支持筛选、处理、备注和关闭。")
    add_para(doc, "恢复策略：系统优先通过重新读取数据库状态恢复页面；对外部模拟能力失败采用最近一次有效数据或初始化数据降级；对预约、订单、支付等核心状态冲突，不做静默覆盖，由服务层返回错误码并要求用户或管理员重新确认。")
    add_heading(doc, "5.6 设计评审与发布结论", 2)
    add_para(doc, "2026年6月5日由系统架构师程晓洋组织全员设计评审，程子浩以RA/QA视角检查需求到设计的追踪关系，丁梓钊以PM/前端视角检查页面、交互和异常提示是否满足用户场景。评审重点围绕HLD完整性、出错处理、接口一致性、数据库状态一致性和第4周编码可执行性展开。")
    add_table(
        doc,
        ["编号", "评审项", "结论", "处理结果"],
        [
            ("DR-07", "HLD是否覆盖6月3日和6月4日任务要求", "通过", "上下文图、架构图、活动图、时序图、接口、ER/表结构和界面设计已整合"),
            ("DR-08", "第5章出错处理是否支撑核心业务闭环", "通过", "补充错误分类、错误码、预约/支付/设备等关键异常流程"),
            ("DR-09", "SRS 18个一级模块是否可追踪到HLD模块、接口和数据表", "通过", "模块分解表、接口清单和数据库表设计保持一致"),
            ("DR-10", "第4周编码是否具备明确输入", "通过", "已明确前后端分工、P0优先级和异常处理约束"),
        ],
    )
    add_para(doc, "评审结论：HLD V1.0通过小组设计评审，可作为第4周详细设计、编码实现和测试用例编写依据；系统设计提交至QQ群“04_系统设计”目录。")


def build():
    if not TEMPLATE.exists():
        raise FileNotFoundError(f"Missing template base: {TEMPLATE}")
    doc = Document(str(TEMPLATE))
    fill_front_matter(doc)
    update_toc(doc)
    strip_template_body(doc)
    add_body(doc)
    doc.save(OUTPUT)
    print(f"[保存] {OUTPUT}")


if __name__ == "__main__":
    build()
