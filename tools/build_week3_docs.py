# -*- coding: utf-8 -*-
"""生成第三周交付物Word文档：SRS需求规格说明书 + HLD概要设计说明书"""
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DOC_DIR = ROOT / "文档（项目相关的作业所需提交文档生成在这里）"
SRS_DIR = DOC_DIR / "03_需求分析"
HLD_DIR = DOC_DIR / "04_系统设计"

GROUP_NAME = "19组"
LEADER = "程晓洋"
PROJECT_NAME = "AI-driven城市智慧停车管理与诱导系统"
DATE_TEXT = "2026年6月2日"

BLUE = RGBColor(36, 93, 150)
DARK_BLUE = RGBColor(28, 67, 110)
TEXT = RGBColor(35, 35, 35)
MUTED = RGBColor(88, 96, 105)
LIGHT_BLUE = "EAF2FB"
LIGHT_GRAY = "F4F6F8"
BORDER = "B9C6D3"


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_cell_text(cell, text, bold=False, color=TEXT, size=10):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def configure_document(doc, footer_text):
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.88)
    section.bottom_margin = Inches(0.88)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in [
        ("Heading 1", 15, BLUE, 14, 7),
        ("Heading 2", 12.5, DARK_BLUE, 9, 5),
        ("Heading 3", 11.5, DARK_BLUE, 7, 3),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(footer_text)
    set_run_font(run, size=8.5, color=MUTED)


def add_title(doc, title, subtitle):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(18)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(PROJECT_NAME)
    set_run_font(run, size=20, bold=True, color=RGBColor(0, 0, 0))

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(14)
    run = paragraph.add_run(subtitle)
    set_run_font(run, size=14, bold=True, color=DARK_BLUE)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(22)
    paragraph.paragraph_format.line_spacing = 1.18
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color=TEXT)


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.15
        run = paragraph.add_run(item)
        set_run_font(run, size=10.5, color=TEXT)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    set_table_borders(table)
    for idx, header in enumerate(headers):
        set_cell_text(table.cell(0, idx), header, bold=True, color=DARK_BLUE, size=9.8)
        shade_cell(table.cell(0, idx), LIGHT_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=9.4)
    doc.add_paragraph()
    return table


def add_meta_table(doc, doc_type, submit_dir):
    rows = [
        ("项目名称", PROJECT_NAME),
        ("小组编号", GROUP_NAME),
        ("组长", LEADER),
        ("小组成员", "程晓洋、丁梓钊、程子浩"),
        ("文档类型", doc_type),
        ("版本号", "V1.0"),
        ("编写日期", DATE_TEXT),
        ("提交目录", submit_dir),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_width(table, [1700, 7900])
    set_table_borders(table)
    for idx, (label, value) in enumerate(rows):
        set_cell_text(table.cell(idx, 0), label, bold=True, color=DARK_BLUE)
        set_cell_text(table.cell(idx, 1), value)
        shade_cell(table.cell(idx, 0), LIGHT_GRAY)
    doc.add_paragraph()


def build_srs():
    """生成功能需求规格说明书 SRS V1.0 Word文档"""
    SRS_DIR.mkdir(parents=True, exist_ok=True)
    out_path = SRS_DIR / f"{GROUP_NAME}-{LEADER}-需求分析.docx"

    doc = Document()
    configure_document(doc, f"{GROUP_NAME}-{LEADER}-需求分析 | SRS V1.0")
    add_title(doc, "功能需求规格说明书", "Software Requirement Specification V1.0")
    add_meta_table(doc, "Software Requirement Specification / 需求规格说明书", "03_需求分析")

    # 第一章 引言
    add_heading(doc, "一、引言")
    add_heading(doc, "1.1 编写目的", 2)
    add_body(doc, '本文档旨在详细描述\u201cAI-driven城市智慧停车管理与诱导系统\u201d的全部功能需求与非功能性需求，作为后续系统设计、编码实现、测试验收的依据。')
    add_heading(doc, "1.2 项目背景", 2)
    add_body(doc, "全国停车位缺口超8000万个，城市核心区车位利用率不足60%，车主找车位平均耗时15-30分钟。本项目构建一个集车位查询、在线预约、智能诱导、计费支付、数据看板于一体的智慧停车管理平台。")
    add_heading(doc, "1.3 用户角色", 2)
    add_table(doc,
        ["角色", "标识", "描述"],
        [
            ("私家车主", "USER", "使用停车服务的主要用户，需要找车位、预约、支付、寻车"),
            ("停车场管理员", "ADMIN", "管理单个或多个停车场的运营人员"),
            ("车位主（共享）", "OWNER", "将闲置车位出租的个人或单位"),
            ("城市交管部门", "GOV", "监管全市停车态势的管理部门"),
        ], [2000, 1200, 6400])

    # 第二章 功能需求（P0核心模块）
    add_heading(doc, "二、功能需求详述（P0核心模块）")

    # M1
    add_heading(doc, "2.1 用户注册与登录（M1）", 2)
    add_body(doc, "【优先级】P0核心 | 【负责人】程晓洋")
    add_body(doc, "提供用户注册、登录、个人信息管理功能，是整个系统的入口。支持三种角色注册：私家车主、车位主、停车场管理员。")
    add_table(doc,
        ["编号", "功能点", "描述"],
        [
            ("M1-F1", "用户注册", "新用户填写注册信息创建账号"),
            ("M1-F2", "用户登录", "已注册用户凭用户名密码登录"),
            ("M1-F3", "个人信息查看", "查看当前登录用户信息"),
            ("M1-F4", "个人信息修改", "修改手机号等个人信息"),
            ("M1-F5", "密码修改", "修改登录密码"),
            ("M1-F6", "退出登录", "清除登录状态"),
        ], [1200, 2400, 6000])
    add_body(doc, "业务规则：用户名3-20字符、密码6-20字符含字母数字、手机号11位格式、密码bcrypt哈希存储、Session管理24小时有效。")
    add_body(doc, "验收标准：注册/登录/密码修改功能正常，错误密码3次锁定30秒，Session过期需重新登录。")

    # M2
    add_heading(doc, "2.2 停车场信息管理（M2）", 2)
    add_body(doc, "【优先级】P0核心 | 【负责人】程晓洋")
    add_body(doc, "停车场管理员可对停车场基本信息进行增删改查，包括车场名称、地址、坐标、总车位数、收费标准等，是系统的数据基础。")
    add_table(doc,
        ["编号", "功能点", "描述"],
        [
            ("M2-F1", "停车场列表", "分页展示所有停车场基本信息"),
            ("M2-F2", "新增停车场", "管理员添加新停车场"),
            ("M2-F3", "编辑停车场", "修改停车场信息"),
            ("M2-F4", "删除停车场", "删除停车场（级联删除车位数据）"),
            ("M2-F5", "停车场详情", "查看单个停车场完整信息"),
            ("M2-F6", "车位配置", "为停车场批量配置车位"),
        ], [1200, 2400, 6000])

    # M3
    add_heading(doc, "2.3 车位状态实时查询（M3）", 2)
    add_body(doc, "【优先级】P0核心 | 【负责人】程子浩")
    add_body(doc, "车主可搜索目的地周边停车场，查看各车场实时空余车位数量，按距离、价格、空余数等多维度筛选排序。")
    add_table(doc,
        ["编号", "功能点", "描述"],
        [
            ("M3-F1", "目的地搜索", "输入目的地搜索周边停车场"),
            ("M3-F2", "停车场列表展示", "展示车场卡片（名称、距离、空余、价格）"),
            ("M3-F3", "多维度筛选", "按距离/价格/空余数筛选排序"),
            ("M3-F4", "车位详情查看", "进入车场查看具体车位分布"),
            ("M3-F5", "实时状态刷新", "自动/手动刷新车位占用状态（30秒间隔）"),
            ("M3-F6", "城市停车一张图", "模拟GIS地图展示全市车场分布"),
        ], [1200, 2400, 6000])

    # M4
    add_heading(doc, "2.4 车位在线预约（M4）", 2)
    add_body(doc, "【优先级】P0核心 | 【负责人】程子浩")
    add_body(doc, "车主选定停车场后，可查看具体车位并在线预约锁定。车位锁定后保留15分钟，超时未到场确认则自动释放。")
    add_body(doc, "状态机设计：free → locked（预约锁定）→ occupied（到场确认）→ free（出场释放）")
    add_body(doc, "业务规则：同一用户同时只能有1个活跃预约、锁定时间15分钟可配置、超时自动释放。")

    # M5
    add_heading(doc, "2.5 停车计费与模拟支付（M5）", 2)
    add_body(doc, "【优先级】P0核心 | 【负责人】程子浩")
    add_body(doc, "车辆出场时系统自动计算停车时长和费用，车主确认后进行模拟支付。计费公式：费用 = 停车时长(小时) × 每小时费率。")
    add_body(doc, "业务规则：不足1小时按1小时计费、支持不同车场不同费率、支持免费时长（首15分钟）、支付为模拟操作。")

    # M6
    add_heading(doc, "2.6 停车记录与反向寻车（M6）", 2)
    add_body(doc, "【优先级】P1重要 | 【负责人】丁梓钊")
    add_body(doc, "车主可查看历史停车记录。在大型停车场内，输入车牌号查询当前停车位置，系统提供步行导航指引帮助车主快速找到车辆。")

    # M7
    add_heading(doc, "2.7 车主服务端（M7）", 2)
    add_body(doc, "【优先级】P0核心 | 【负责人】丁梓钊")
    add_body(doc, "车主端的统一首页入口，集成搜索、推荐、预约、支付等功能入口，是车主用户的操作门户。")

    # M8
    add_heading(doc, "2.8 车场运营管理后台（M8）", 2)
    add_body(doc, "【优先级】P1重要 | 【负责人】程晓洋")
    add_body(doc, "停车场管理员通过Web后台查看所属车场的实时车位状态、进出流量、收入统计、设备状态，支持远程操作和报表导出。")

    # M9
    add_heading(doc, "2.9 数据可视化看板（M9）", 2)
    add_body(doc, "【优先级】P1重要 | 【负责人】程晓洋")
    add_body(doc, "使用ECharts图表库，以可视化方式展示车位周转率、高峰饱和度、收入趋势、车位占用热力图等核心运营指标。")

    # 第三章 扩展模块
    add_heading(doc, "三、扩展功能模块（M10-M22）")
    add_body(doc, '以下为P2增强模块（M10-M11完整实现）和P2*模拟扩展模块（M12-M22简化实现）。P2*模块因实训环境限制（无法部署真实AI视觉、GIS地图、支付系统），采用\u201c需求完整定义+编码简化实现（CRUD+模拟数据展示）\u201d策略。')

    add_table(doc,
        ["模块ID", "模块名称", "优先级", "实现策略"],
        [
            ("M10", "共享车位管理", "P2增强", "完整CRUD实现"),
            ("M11", "违停记录管理", "P2增强", "完整CRUD实现"),
            ("M12", "路内停车位视觉检测", "P2*模拟", "模拟数据展示面板"),
            ("M13", "停车场数据联网接入", "P2*模拟", "接入状态面板"),
            ("M14", "停车诱导信息发布", "P2*模拟", "诱导屏管理页面"),
            ("M15", "车牌识别与无感支付", "P2*模拟", "模拟识别+支付"),
            ("M16", "停车行为大数据分析", "P2*模拟", "M9扩展分析页"),
            ("M17", "动态定价策略引擎", "P2*模拟", "费率系数配置"),
            ("M18", "充电车位智能管理", "P2*模拟", "充电类型+告警"),
            ("M19", "长期停车月卡管理", "P2*模拟", "月卡CRUD+续费"),
            ("M20", "设备运维管理平台", "P2*模拟", "设备列表+监控"),
            ("M21", "城市停车监管平台", "P2*模拟", "监管大屏页面"),
            ("M22", "商圈停车联合营销", "P2*模拟", "营销活动管理"),
        ], [1200, 2800, 1200, 4400])

    # 第四章 非功能性需求
    add_heading(doc, "四、非功能性需求")
    add_table(doc,
        ["类型", "要求"],
        [
            ("性能", "页面加载≤2s，API响应≤500ms"),
            ("可用性", "支持Chrome/Edge/Firefox，移动端响应式适配"),
            ("安全", "密码bcrypt哈希存储，Session认证，参数化查询防SQL注入"),
            ("可维护性", "PEP8 + ESLint规范，Git提交信息规范，Swagger API文档"),
            ("数据", "模拟50+停车场、500+车位，6张核心数据库表+扩展表"),
        ], [2000, 7600])

    # 第五章 接口需求
    add_heading(doc, "五、接口需求")
    add_body(doc, "前后端通信采用RESTful JSON格式，统一返回格式：{code: int, msg: str, data: object}。")
    add_body(doc, "状态码：200成功、400参数错误、401未登录、403无权限、404未找到、500服务器错误。")
    add_body(doc, "API端点总计40+个，覆盖认证、车场、车位、预约、订单、记录、管理、分析8个模块。")

    # 第六章 需求评审结论
    add_heading(doc, "六、需求评审结论")
    add_body(doc, "评审日期：2026年6月2日 | 评审方式：小组会议（全员参与）")
    add_body(doc, "评审发现8项问题（5项遗漏、1项矛盾、2项不明确），全部修订闭合。评审结论：✅ 通过。22个功能模块全部覆盖（P0核心7个+P1重要2个+P2增强2个+P2*模拟扩展11个），需求完整性、可测试性、可追溯性均已满足。")
    add_body(doc, "评审角色：PM（程晓洋）→ 全员参与评审 → RA（程子浩）修订 → PM（丁梓钊）确认")

    # 附录 需求追踪矩阵
    add_heading(doc, "附录：需求追踪矩阵（RTM）")
    rtm_rows = [
        ("M1", "用户注册与登录", "P0", "程晓洋", "6个功能点"),
        ("M2", "停车场信息管理", "P0", "程晓洋", "6个功能点"),
        ("M3", "车位状态实时查询", "P0", "程子浩", "6个功能点"),
        ("M4", "车位在线预约", "P0", "程子浩", "6个功能点"),
        ("M5", "停车计费与模拟支付", "P0", "程子浩", "5个功能点"),
        ("M6", "停车记录与反向寻车", "P1", "丁梓钊", "5个功能点"),
        ("M7", "车主服务端(统一入口)", "P0", "丁梓钊", "5个功能点"),
        ("M8", "车场运营管理后台", "P1", "程晓洋", "6个功能点"),
        ("M9", "数据可视化看板", "P1", "程晓洋", "6个功能点"),
        ("M10", "共享车位管理", "P2", "丁梓钊", "5个功能点"),
        ("M11", "违停记录管理", "P2", "丁梓钊", "4个功能点"),
        ("M12", "路内停车位视觉检测", "P2*", "丁梓钊", "2个功能点"),
        ("M13", "停车场数据联网接入", "P2*", "程晓洋", "2个功能点"),
        ("M14", "停车诱导信息发布", "P2*", "程子浩", "2个功能点"),
        ("M15", "车牌识别与无感支付", "P2*", "程子浩", "2个功能点"),
        ("M16", "停车行为大数据分析", "P2*", "程晓洋", "2个功能点"),
        ("M17", "动态定价策略引擎", "P2*", "程子浩", "2个功能点"),
        ("M18", "充电车位智能管理", "P2*", "丁梓钊", "2个功能点"),
        ("M19", "长期停车月卡管理", "P2*", "丁梓钊", "2个功能点"),
        ("M20", "设备运维管理平台", "P2*", "丁梓钊", "2个功能点"),
        ("M21", "城市停车监管平台", "P2*", "程晓洋", "2个功能点"),
        ("M22", "商圈停车联合营销", "P2*", "丁梓钊", "2个功能点"),
    ]
    add_table(doc, ["模块ID", "模块名称", "优先级", "负责人", "功能点数"], rtm_rows, [1200, 3000, 1000, 1400, 1400])

    add_body(doc, f"编制：第{GROUP_NAME} | {DATE_TEXT} | SRS V1.0 正式版")
    doc.save(out_path)
    print(f"SRS: {out_path}")


def build_hld():
    """生成概要设计说明书 HLD V1.0 Word文档"""
    HLD_DIR.mkdir(parents=True, exist_ok=True)
    out_path = HLD_DIR / f"{GROUP_NAME}-{LEADER}-系统设计.docx"

    doc = Document()
    configure_document(doc, f"{GROUP_NAME}-{LEADER}-系统设计 | HLD V1.0")
    add_title(doc, "概要设计说明书", "High-Level Design V1.0")
    add_meta_table(doc, "System Design HLD / 概要设计说明书", "04_系统设计")

    # 第一章
    add_heading(doc, "一、引言")
    add_body(doc, '本文档是\u201cAI-driven城市智慧停车管理与诱导系统\u201d的概要设计说明书（HLD），描述系统的整体架构、模块划分、技术选型、数据库设计和接口规范，作为后续详细设计和编码实现的依据。')

    # 第二章 系统总体设计
    add_heading(doc, "二、系统总体设计")
    add_heading(doc, "2.1 系统架构", 2)
    add_body(doc, "采用前后端分离的B/S架构：前端Vue3 SPA（车主端+管理端）→ HTTP RESTful JSON → FastAPI后端 → SQLAlchemy ORM → SQLite数据库。")
    add_body(doc, "设计原则：单一职责、接口隔离、前后端分离、数据安全、约定优于配置。")

    add_heading(doc, "2.2 技术选型", 2)
    add_table(doc,
        ["层级", "技术", "版本", "说明"],
        [
            ("前端框架", "Vue 3", "3.x", "组件化开发"),
            ("UI组件库", "Element Plus", "2.x", "Vue3官方推荐"),
            ("图表库", "ECharts", "5.x", "数据可视化"),
            ("构建工具", "Vite", "5.x", "快速冷启动"),
            ("后端框架", "FastAPI", "0.100+", "异步高性能"),
            ("ORM", "SQLAlchemy", "2.x", "Python最成熟ORM"),
            ("数据库", "SQLite", "3.x", "零部署文件存储"),
            ("密码哈希", "bcrypt", "4.x", "安全哈希算法"),
            ("认证", "Session(Starlette)", "内置", "简单可靠"),
        ], [1800, 2400, 1200, 4200])

    # 第三章 模块设计
    add_heading(doc, "三、模块设计")
    add_body(doc, "系统共22个功能模块，按四级优先级分层：P0核心7个（M1-M7完整实现）、P1重要2个（M8-M9完整实现）、P2增强2个（M10-M11完整实现）、P2*模拟11个（M12-M22简化实现）。")

    add_table(doc,
        ["模块", "前端路由示例", "后端API端点", "数据表"],
        [
            ("M1 用户注册登录", "/login, /register, /profile", "POST /api/auth/*", "users"),
            ("M2 车场管理", "/admin/lots", "GET/POST/PUT/DELETE /api/lots", "parking_lots"),
            ("M3 车位查询", "/search, /lots/:id/spots", "GET /api/lots/search", "parking_lots, parking_spots"),
            ("M4 车位预约", "/spots/:id/reserve", "POST /api/reservations", "reservations"),
            ("M5 计费支付", "/orders/:id/checkout", "POST /api/orders/*", "parking_orders"),
            ("M6 记录寻车", "/my-records, /find-car", "GET /api/records/*", "parking_orders"),
            ("M7 车主端首页", "/", "GET /api/home/stats", "聚合多表"),
            ("M8 运营后台", "/admin/dashboard", "GET /api/admin/*", "多表统计"),
            ("M9 数据看板", "/admin/analytics", "GET /api/analytics/*", "多表统计"),
        ], [1800, 2600, 2800, 2400])

    # 第四章 数据库设计
    add_heading(doc, "四、数据库设计")
    add_body(doc, "数据库使用SQLite，共8张核心表（含扩展预留）：")
    add_table(doc,
        ["表名", "用途", "核心字段"],
        [
            ("users", "用户表", "id, username, password_hash, phone, role"),
            ("parking_lots", "停车场表", "id, name, address, total_spots, rate_per_hour"),
            ("parking_spots", "车位表", "id, lot_id, spot_number, floor, zone, status"),
            ("reservations", "预约记录表", "id, user_id, spot_id, status, expire_at"),
            ("parking_orders", "停车订单表", "id, user_id, spot_id, entry_time, exit_time, amount"),
            ("violations", "违停记录表", "id, lot_id, plate_number, description"),
            ("shared_spots", "共享车位表", "id, owner_id, price_per_hour, available_days"),
            ("devices", "设备表(预留)", "id, name, type, status, lot_id"),
        ], [1800, 2200, 5600])

    add_body(doc, "索引设计：username唯一索引、(lot_id, status)复合索引、(user_id, status)复合索引、expire_at定时扫描索引等。")

    # 第五章 接口设计
    add_heading(doc, "五、接口设计")
    add_body(doc, 'API统一返回格式：{code: 200, msg: "success", data: {}}。分页格式：{items, total, page, page_size}。')
    add_body(doc, "API端点总计40+个，分为8个模块：认证(6个)、车场(7个)、车位(3个)、预约(5个)、订单(4个)、记录(3个)、管理(5个)、分析(6个)。")

    # 第六章 前端设计
    add_heading(doc, "六、前端设计")
    add_body(doc, "前端路由：车主端13条路由 + 管理端9条路由。组件树分为layouts/（布局层）、components/（公共组件层）、views/（页面层）、stores/（Pinia状态管理）、api/（Axios请求层）。")
    add_body(doc, "状态管理（Pinia）：auth（用户登录状态）、parking（停车场和车位数据）、admin（管理端统计数据）。")

    # 第七章 后端设计
    add_heading(doc, "七、后端设计")
    add_body(doc, "后端目录结构：main.py（入口）、config.py（配置）、database.py（数据库连接）、models/（SQLAlchemy模型）、schemas/（Pydantic模型）、routers/（API路由）、services/（业务逻辑层）、middleware/（认证中间件）、utils/（工具函数）、seed_data.py（模拟数据初始化）。")

    # 第八章 部署架构
    add_heading(doc, "八、部署架构")
    add_body(doc, "开发环境：前端Vite Dev Server（localhost:5173）+ 后端Uvicorn（localhost:8000）+ SQLite文件数据库（smartpark.db）。")
    add_body(doc, "启动方式：后端 pip install → python seed_data.py → uvicorn main:app --reload；前端 npm install → npm run dev。")

    # 第九章 安全设计
    add_heading(doc, "九、安全设计")
    add_bullets(doc, [
        "密码安全：bcrypt哈希存储，不可逆",
        "认证安全：Session管理，24小时过期",
        "接口安全：所有API校验登录态，管理接口校验角色",
        "数据安全：SQLAlchemy参数化查询，防SQL注入",
        "输入校验：Pydantic模型校验 + Element Plus表单校验",
    ])

    # 第十章 设计评审结论
    add_heading(doc, "十、设计评审结论")
    add_body(doc, "评审日期：2026年6月2日 | 评审方式：小组设计评审会议（全员参与）")
    add_body(doc, "评审发现6项建议/确认项，全部处理闭合。评审结论：✅ 通过。架构合理性、技术可行性、可扩展性、安全性均已满足。")
    add_body(doc, "评审角色：SA（程晓洋）→ RA（程子浩）→ PM（丁梓钊）→ 全员评审")

    add_body(doc, f"编制：第{GROUP_NAME} | {DATE_TEXT} | HLD V1.0")
    doc.save(out_path)
    print(f"HLD: {out_path}")


if __name__ == "__main__":
    build_srs()
    build_hld()
    print("Done! Both SRS and HLD Word documents generated.")
