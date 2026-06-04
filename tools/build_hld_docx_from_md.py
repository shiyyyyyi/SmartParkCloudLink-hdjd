# -*- coding: utf-8 -*-
"""Build the HLD Word deliverable from the checked Markdown HLD draft.

The source Markdown already follows the HLD template sections. This script
converts it into a submit-ready docx without touching the original template
directory.
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
HLD_DIR = ROOT / "文档（项目相关的作业所需提交文档生成在这里）" / "04_系统设计"
SOURCE = HLD_DIR / "概要设计说明书_HLD_V1.0.md"
OUTPUT = HLD_DIR / "19组-程晓洋-系统设计.docx"


def set_run(run, size=10.5, bold=False, color=None, font="宋体"):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.2
    normal.paragraph_format.space_after = Pt(5)

    for name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12), ("Heading 4", 11)]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(5)


def parse_table(lines, start):
    rows = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        raw = lines[index].strip().strip("|")
        rows.append([cell.strip() for cell in raw.split("|")])
        index += 1
    if len(rows) >= 2 and all(set(cell.replace(":", "").strip()) <= {"-"} for cell in rows[1]):
        return rows[:1] + rows[2:], index
    return rows, index


def add_table(doc, rows):
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = doc.add_table(rows=1, cols=width)
    table.style = "Table Grid"
    for col in range(width):
        value = rows[0][col] if col < len(rows[0]) else ""
        run = table.cell(0, col).paragraphs[0].add_run(value)
        set_run(run, size=9.5, bold=True, font="黑体")
    for row in rows[1:]:
        cells = table.add_row().cells
        for col in range(width):
            value = row[col] if col < len(row) else ""
            run = cells[col].paragraphs[0].add_run(value)
            set_run(run, size=9)
    doc.add_paragraph()


def add_code_block(doc, block):
    for line in block:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line if line else " ")
        set_run(run, size=8.5, font="Consolas")


def build():
    md_lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_doc(doc)

    index = 0
    in_code = False
    code_block = []
    while index < len(md_lines):
        line = md_lines[index]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_block)
                code_block = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue

        if in_code:
            code_block.append(line.rstrip())
            index += 1
            continue

        if not stripped or stripped == "---":
            index += 1
            continue

        if stripped.startswith("|"):
            rows, index = parse_table(md_lines, index)
            add_table(doc, rows)
            continue

        if stripped.startswith("#"):
            level = min(len(stripped) - len(stripped.lstrip("#")), 4)
            text = stripped[level:].strip()
            paragraph = doc.add_paragraph(style=f"Heading {level}")
            paragraph.add_run(text)
            if level == 1:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            index += 1
            continue

        if stripped.startswith(">"):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(stripped.lstrip(">").strip())
            set_run(run, size=10, bold=False)
            index += 1
            continue

        if stripped.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            run = paragraph.add_run(stripped[2:].strip())
            set_run(run, size=10.5)
            index += 1
            continue

        paragraph = doc.add_paragraph()
        run = paragraph.add_run(stripped)
        set_run(run, size=10.5)
        index += 1

    doc.save(OUTPUT)
    print(f"[保存] {OUTPUT}")


if __name__ == "__main__":
    build()
