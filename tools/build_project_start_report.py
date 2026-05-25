# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOC_DIR = ROOT / "文档（项目相关的作业所需提交文档生成在这里）"
OUT_DIR = DOC_DIR / "01_项目立项"
OUT_PATH = OUT_DIR / "10组-程晓洋-项目立项.docx"

PROJECT_NAME = "AI-driven 城市智慧停车管理与诱导系统"
GROUP_NAME = "10组"
LEADER = "程晓洋"
DATE_TEXT = "2026年5月25日"

BLUE = RGBColor(36, 93, 150)
DARK_BLUE = RGBColor(28, 67, 110)
TEXT = RGBColor(35, 35, 35)
MUTED = RGBColor(88, 96, 105)
LIGHT_BLUE = "EAF2FB"
LIGHT_GRAY = "F4F6F8"
BORDER = "B9C6D3"


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def style_paragraph(paragraph, size=11, color=TEXT, bold=None):
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color, bold=bold)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_cell_text(cell, text, bold=False, color=TEXT, size=10):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def configure_document(doc):
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.88)
    section.bottom_margin = Inches(0.88)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in [
        ("Heading 1", 15, BLUE, 14, 7),
        ("Heading 2", 12.5, DARK_BLUE, 9, 5),
        ("Heading 3", 11.5, DARK_BLUE, 7, 3),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"{GROUP_NAME}-{LEADER}-项目立项 | {PROJECT_NAME}")
    set_run_font(run, size=8.5, color=MUTED)


def add_title(doc):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(18)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(PROJECT_NAME)
    set_run_font(run, size=22, bold=True, color=RGBColor(0, 0, 0))

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(14)
    run = paragraph.add_run("项目立项报告（项目启动文档）")
    set_run_font(run, size=14, bold=True, color=DARK_BLUE)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)
    style_paragraph(paragraph, size=15 if level == 1 else 12.5, color=BLUE if level == 1 else DARK_BLUE, bold=True)


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(22)
    paragraph.paragraph_format.line_spacing = 1.18
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color=TEXT)


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.15
        run = paragraph.add_run(item)
        set_run_font(run, size=10.5, color=TEXT)


def add_numbered(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.15
        run = paragraph.add_run(item)
        set_run_font(run, size=10.5, color=TEXT)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    set_table_borders(table)
    for idx, header in enumerate(headers):
        set_cell_text(table.cell(0, idx), header, bold=True, color=DARK_BLUE, size=9.8)
        shade_cell(table.cell(0, idx), LIGHT_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=9.4)
    doc.add_paragraph()
    return table


def add_meta_table(doc):
    rows = [
        ("项目名称", PROJECT_NAME),
        ("小组编号", GROUP_NAME),
        ("组长", LEADER),
        ("小组成员", "程晓洋、丁梓钊、程子浩"),
        ("项目行业", "智慧停车 / 智慧城市 / 城市交通管理"),
        ("文档类型", "Project Start Report / 项目启动文档"),
        ("编写日期", DATE_TEXT),
        ("提交目录", "01_项目立项"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_width(table, [1700, 7900])
    set_table_borders(table)
    for idx, (label, value) in enumerate(rows):
        set_cell_text(table.cell(idx, 0), label, bold=True, color=DARK_BLUE)
        set_cell_text(table.cell(idx, 1), value)
        shade_cell(table.cell(idx, 0), LIGHT_GRAY)
    doc.add_paragraph()


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    configure_document(doc)
    add_title(doc)
    add_meta_table(doc)

    add_heading(doc, "一、项目背景")
    add_body(
        doc,
        "随着城市化进程加速和汽车保有量持续增长，城市核心区停车位供需矛盾突出。车主寻找车位耗时长、"
        "路内停车状态难以及时感知、停车场数据分散、收费与诱导方式粗放，共同造成无效交通流增加、"
        "道路拥堵加剧和停车资源利用率不足。本项目拟以 AI 视觉感知、物联网接入、大数据分析和城市级停车诱导为基础，"
        "建设覆盖路内停车、路外停车场、共享停车和充电车位等场景的智慧停车管理与诱导平台。"
    )
    add_body(
        doc,
        "从实训项目角度看，该选题业务边界清晰、用户角色明确、功能模块完整，既能支撑需求分析和系统设计文档撰写，"
        "也便于后续分阶段完成前后端、数据库、统计分析和模拟算法接口等编码实践。"
    )

    add_heading(doc, "二、第一日任务完成情况")
    add_table(
        doc,
        ["任务项", "完成内容", "输出结果"],
        [
            ("项目需求调研", "梳理停车难、车位信息孤岛、停车诱导缺失、收费规则粗放、无感支付覆盖不足等痛点。", "形成项目背景、用户诉求和核心功能范围。"),
            ("完成分组", "确定每组三人，成员为程晓洋、丁梓钊、程子浩。", "明确组长和初步角色分工。"),
            ("小组会议与选题", f"围绕候选项目进行讨论，最终确定“{PROJECT_NAME}”。", "形成选题结论和首期建设边界。"),
            ("项目启动文档", "基于模板完成项目立项报告，覆盖背景、目标、范围、功能、分工、计划、风险和预期成果。", "生成可提交 Word 文档。"),
        ],
        [1800, 4700, 3100],
    )

    add_heading(doc, "三、项目建设目标")
    add_bullets(
        doc,
        [
            "实现停车场和路内车位状态的实时采集、汇聚与可视化展示。",
            "为车主提供停车场搜索、智能推荐、车位预约、导航诱导、停车记录和无感支付服务。",
            "为停车场运营方提供车位、订单、收入、设备和异常事件的一体化管理能力。",
            "为城市管理部门提供停车态势监控、供需分析、周转率分析和政策评估数据支撑。",
            "通过动态定价和共享停车机制提升存量车位利用率，缓解热点区域停车压力。",
        ],
    )

    add_heading(doc, "四、项目范围")
    add_heading(doc, "4.1 建设范围", level=2)
    add_bullets(
        doc,
        [
            "车主移动端或小程序：停车场搜索、车位预约、停车记录、在线支付、反向寻车等。",
            "车场运营管理后台：停车场信息维护、车位状态管理、设备状态监控、订单与收入统计等。",
            "城市停车监管平台：全市停车一张图、供需态势、收费统计、投诉热力和政策评估。",
            "数据与算法模块：车位检测结果接入、车牌识别结果接入、动态定价规则、停车行为数据分析。",
        ],
    )
    add_heading(doc, "4.2 首期实现边界", level=2)
    add_body(
        doc,
        "考虑实训周期与实现难度，首期编码建议聚焦“车位查询 - 推荐预约 - 入出场记录 - 订单支付 - 管理统计”的核心闭环。"
        "AI 视觉识别、GIS 地图和支付通道可先以模拟数据和接口预留方式完成演示，后续迭代再接入真实服务。"
    )

    add_heading(doc, "五、用户与业务需求")
    add_table(
        doc,
        ["用户角色", "核心诉求", "主要功能"],
        [
            ("车主", "快速找到合适车位、少绕路、少排队、支付方便。", "停车场搜索、智能推荐、预约、导航、无感支付、停车记录、电子发票。"),
            ("停车场管理员", "提升车位周转率，降低人工管理成本，掌握收入和设备状态。", "车位状态管理、订单管理、收入统计、设备监控、异常处理。"),
            ("城市管理部门", "掌握全市停车态势，辅助规划和监管，减少违停与拥堵。", "停车一张图、供需分析、周转率排行、违停取证、监管大屏。"),
            ("车位资源方", "盘活闲置车位，支持错时共享和收益管理。", "共享时段发布、价格配置、预约审核、收益统计。"),
        ],
        [1650, 3600, 4350],
    )

    add_heading(doc, "六、核心功能模块")
    add_table(
        doc,
        ["模块", "功能说明", "首期优先级"],
        [
            ("车位状态感知", "通过 AI 视频桩、高位相机或模拟接口获取空闲、占用、违停和车牌信息。", "高"),
            ("停车场数据接入", "对接停车场管理系统，汇聚实时空余车位、价格和运营状态。", "高"),
            ("城市停车一张图", "在地图或模拟地图中展示停车场位置、空位、价格、距离和区域热度。", "高"),
            ("智能推荐与预约", "按目的地、距离、价格、空位、预约可用性推荐停车方案。", "高"),
            ("车牌识别与支付", "支持入出场识别、自动计费、绑定账户扣款或模拟支付。", "中"),
            ("运营管理后台", "管理停车场、车位、订单、收入、设备状态和异常告警。", "高"),
            ("动态定价策略", "基于区域热度、时段和利用率调整价格，并控制价格浮动范围。", "中"),
            ("共享停车管理", "发布错时共享车位，支持预约和收益统计。", "中"),
            ("充电车位管理", "监测充电桩状态并识别燃油车占用充电位等异常。", "低"),
            ("大数据分析", "统计周转率、平均停车时长、峰值饱和度、区域热度等指标。", "高"),
        ],
        [1900, 5950, 1750],
    )

    add_heading(doc, "七、技术可行性分析")
    add_numbered(
        doc,
        [
            "数据层可采用关系型数据库管理用户、车辆、停车场、车位、预约、订单、设备和统计数据，结构清晰且便于实训开发。",
            "后端可采用 RESTful API 组织业务服务，模块之间通过标准接口解耦，便于前后端并行开发。",
            "前端可采用 Web 管理后台与移动端原型实现主要交互，GIS 能力可先以地图组件或坐标模拟方式完成。",
            "AI 识别和动态定价模块在首期可采用规则引擎和模拟数据，保证业务闭环可运行，同时保留真实算法接入接口。",
            "系统具备分阶段迭代条件，可先完成核心停车服务闭环，再扩展监管大屏、共享停车、充电车位和违停取证等功能。",
        ],
    )

    add_heading(doc, "八、组织分工")
    add_table(
        doc,
        ["成员", "角色", "职责"],
        [
            ("程晓洋", "组长/项目经理、数据库管理员、主力程序员", "负责项目统筹、进度协调、数据库设计、核心接口开发和文档汇总。"),
            ("丁梓钊", "产品经理、主力程序员、测试工程师", "负责需求梳理、功能原型、车主端/管理端关键业务开发和测试用例设计。"),
            ("程子浩", "主力程序员、需求分析师、QA", "负责需求分析、后台模块开发、质量检查、接口联调和文档校对。"),
        ],
        [1500, 3100, 5000],
    )

    add_heading(doc, "九、阶段计划")
    add_table(
        doc,
        ["阶段", "时间", "主要任务", "交付物"],
        [
            ("项目立项", "5月25日", "完成项目调研、分组、选题会议和项目启动文档。", "项目立项报告、会议纪要"),
            ("项目计划", "5月26日-5月27日", "制定开发计划、里程碑、分工、风险和进度安排。", "项目计划报告"),
            ("需求分析", "5月27日-5月30日", "完成用户需求、功能需求、非功能需求、用例和数据需求分析。", "需求规格说明书"),
            ("设计与实现", "后续阶段", "完成系统设计、数据库设计、接口设计、编码实现和测试。", "设计文档、源代码、测试报告"),
        ],
        [1500, 1700, 4200, 2200],
    )

    add_heading(doc, "十、风险与应对")
    add_table(
        doc,
        ["风险项", "影响", "应对措施"],
        [
            ("复杂场景下识别准确率难以保证", "影响车位状态和车牌识别可靠性。", "首期使用模拟识别数据，设计接口预留；后续引入多帧校验和人工复核机制。"),
            ("功能模块较多导致范围失控", "影响进度和交付质量。", "采用优先级管理，先完成核心闭环，再扩展中低优先级功能。"),
            ("动态定价用户接受度不足", "可能引发价格争议。", "价格规则透明化，设置浮动上限和免费短停等保障策略。"),
            ("外部地图、支付和停车场接口依赖", "联调成本较高。", "先使用 mock 数据和标准接口协议，降低外部依赖对实训进度的影响。"),
        ],
        [2400, 2800, 4400],
    )

    add_heading(doc, "十一、预期成果")
    add_bullets(
        doc,
        [
            "形成项目立项报告、项目计划报告、需求规格说明书和后续设计测试文档。",
            "完成可演示的城市智慧停车管理与诱导系统原型或核心功能模块。",
            "实现停车场信息、车位状态、预约订单、支付记录和统计分析等核心数据闭环。",
            "为后续系统设计、编码实现和验收演示奠定清晰范围与交付基础。",
        ],
    )

    add_heading(doc, "十二、立项结论")
    add_body(
        doc,
        f"经第一日需求调研和小组会议讨论，小组确定选择“{PROJECT_NAME}”作为实训项目。"
        "该项目具有明确的现实问题、完整的业务链路、可分阶段落地的技术方案和较好的展示价值，"
        "符合本周项目立项、项目计划和需求分析阶段的交付要求，建议正式立项并进入项目计划与需求规格说明书编写阶段。"
    )

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
